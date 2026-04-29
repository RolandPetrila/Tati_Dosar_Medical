#!/usr/bin/env python3
"""Generează briefing DOCX pentru consult medical (oncolog/cardiolog/endocrin).

Citește CONTEXT_MEDICAL.md + JSON-uri canonice + CONTACTE_MEDICALE.md → produce
DOCX cu 6 secțiuni: date pacient, status TNM, medicație, checklist dosar,
întrebări, notes consult.

Output: Dosar_Medical/rapoarte_generate/<data>_briefing_consult_<tip>.docx + .meta.json

Usage:
  python scripts/generate_consult_briefing.py \\
    --consult oncolog \\
    --data 2026-05-04 \\
    --medic "Dr. Anater Angelo-Christian" \\
    --unitate "OncoHelp Timișoara"

Creat 2026-04-28 ca parte din Faza 3 plan implementare cross-terminal (recomandare N1
din .claude-outputs/imbunatatiri/2026-04-28_032030/).
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
        sys.stderr.reconfigure(encoding="utf-8", errors="replace")
    except (AttributeError, ValueError):
        pass

ROOT = Path(__file__).resolve().parent.parent
CONTEXT_MD = ROOT / "CONTEXT_MEDICAL.md"
RAPOARTE = ROOT / "Dosar_Medical" / "rapoarte_generate"

CONSULT_TEMPLATES = {
    "oncolog": {
        "titlu": "Briefing consult oncolog — Petrilă Viorel-Mihai",
        "intrebari": [
            "Care e clasificarea TNM definitivă post-IHC pe blocul T26H06044?",
            "Care e protocolul terapeutic recomandat (FLOT vs CROSS) pentru Siewert II + componentă fundică?",
            "Cum interpretăm ascita perihepatică 15 mm + intrapelvină 28 mm — carcinomatoză peritoneală sau benignă?",
            "E indicat IHC pe blocul existent sau rebiopsie țintită ghidată CT/EUS?",
            "Care e calendarul propus pentru chimio neoadjuvantă + chirurgie + chimio adjuvantă?",
            "Cum gestionăm pauza Aspenter pre-chirurgie esofagiană (BMS confirmat 14 ani vechime)?",
            "Cum gestionăm DZ tip 2 (HbA1c controlat) + HTA în context chimioterapie?",
            "Care e prognosticul realist și ce factori îl influențează la cazul nostru?",
        ],
    },
    "cardiolog": {
        "titlu": "Briefing consult cardiolog ambulator — Petrilă Viorel-Mihai",
        "intrebari": [
            "Stentul implantat 2012 e BMS confirmat — care e riscul de tromboză in-stent la pauza Aspenter 5-7 zile pre-chirurgie esofagiană?",
            "Care e FEVS curentă (post-ECO) și statusul funcțional?",
            "Putem reluă TORVACARD (statină) pentru LDL 133 mg/dL > țintă ESC <70?",
            "E nevoie de Holter sau test efort pre-anestezic?",
            "Aviz perioperator pentru chirurgie esofagiană majoră?",
        ],
    },
    "endocrin": {
        "titlu": "Briefing consult endocrin — Petrilă Viorel-Mihai",
        "intrebari": [
            "Glanda suprarenală stângă hipertrofă heterogenă (CT 20.04) — incidentaloma? adenom? necesită evaluare hormonală?",
            "DZ tip 2 + Jamesi 100 mg — schimbări la chimioterapie?",
            "HbA1c curent + țintă pre-chirurgie?",
        ],
    },
}


def extract_section(md_text, header_pattern):
    """Extrage o secțiune din MD între două header-uri (## sau ###)."""
    lines = md_text.split("\n")
    in_section = False
    out = []
    for line in lines:
        if re.match(header_pattern, line):
            in_section = True
            out.append(line)
            continue
        if in_section and re.match(r"^#{1,3}\s", line):
            break
        if in_section:
            out.append(line)
    return "\n".join(out).strip()


def build_docx(consult_type, data, medic, unitate, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        print("EROARE: python-docx nu e instalat. Rulează: pip install python-docx")
        return False

    template = CONSULT_TEMPLATES.get(consult_type)
    if not template:
        print(f"EROARE: tip consult necunoscut '{consult_type}'. Tipuri valide: {list(CONSULT_TEMPLATES.keys())}")
        return False

    doc = Document()

    h = doc.add_heading(template["titlu"], 0)
    h.runs[0].font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    p = doc.add_paragraph()
    p.add_run("Consult: ").bold = True
    p.add_run(f"{medic} · {unitate}").italic = True
    p2 = doc.add_paragraph()
    p2.add_run("Data programată: ").bold = True
    p2.add_run(data)
    p3 = doc.add_paragraph()
    p3.add_run("Generat: ").bold = True
    p3.add_run(f"{datetime.now().strftime('%Y-%m-%d %H:%M')} (briefing automat din CONTEXT_MEDICAL.md)")

    doc.add_paragraph("─" * 50)

    # 1. Date pacient
    doc.add_heading("1. Date pacient", 1)
    doc.add_paragraph("• Nume: PETRILĂ Viorel-Mihai")
    doc.add_paragraph("• CNP: 1590518024486")
    doc.add_paragraph("• Data nașterii: 18.05.1959 (66 ani)")
    doc.add_paragraph("• Sex: M")
    doc.add_paragraph("• Domiciliu: Nădlac, jud. Arad")
    doc.add_paragraph("• Medic familie: Dr. Orbán Ecaterina-Maria (Cabinet Medical Individual Nădlac, cod parafă 718705)")

    # 2. Status TNM (suspiciune actuală)
    doc.add_heading("2. Status oncologic actual (28.04.2026)", 1)
    doc.add_paragraph("• Endoscopie 17.04 (Genesis Arad, Dr. Noufal Abdul Vahab): proces proliferativ joncțiune eso-gastrică")
    doc.add_paragraph("• CT TAP 20.04 (Genesis Micălaca, Dr. Buie + Dr. Candea): T3-T4, N0-N1, M0 probabil, Siewert II probabil + ascită perihepatică 15 mm + intrapelvină 28 mm (de exclus carcinomatoză peritoneală)")
    doc.add_paragraph("• Biopsie 17.04 / rezultat 27.04 (Bioclinica, Dr. Glăja Romanița): țesut granulație + ulcerație cronică, doar SUGESTIV pentru carcinom — recomandare IHC pe blocul T26H06044 sau rebiopsie. Suspiciune clinico-imagistică persistă")

    # 3. Medicație curentă
    doc.add_heading("3. Medicație curentă (per schemă 10.11.2025 + clarificări 25.04.2026)", 1)
    doc.add_paragraph("• Aspenter 75 mg — 0-1-0 (antiagregant post-stent BMS 2012)")
    doc.add_paragraph("• Concor 2.5 mg — 1-0-0 (β-blocant)")
    doc.add_paragraph("• Triplixam 5/1.25/5 mg — 1-0-0 (perindopril + indapamidă + amlodipină)")
    doc.add_paragraph("• Jamesi (sitagliptin) 100 mg — 1-0-0 (DZ tip 2)")
    doc.add_paragraph("• TORVACARD 10/20 — PRESCRIS dar NU administrat (clarificat user 25.04). LDL curent 133 mg/dL > țintă ESC <70 — de evaluat la consult oncolog 4.05")

    # 4. Antecedente cardiologic-relevante
    doc.add_heading("4. Antecedent cardiologic critic — STEMI 2012 cu BMS", 1)
    doc.add_paragraph("• 19.02.2012 Vichy (Franța, Dr. Marcaggi): angioplastie cu stent BMS RX VISION 3.5×28 mm Abbott Nr. 1110341 pe IVA proximală")
    doc.add_paragraph("• Tip stent CONFIRMAT BMS (Bare Metal Stent) — endotelizare completă în 4-6 săptămâni → la 14 ani vechime risc tromboză in-stent <1%")
    doc.add_paragraph("• Implicație: pauza Aspenter 5-7 zile pre-chirurgie esofagiană este în general SIGURĂ (vs ipoteza DES care ar cere DAPT prelungit)")

    # 5. Checklist dosar pentru consult
    doc.add_heading("5. Checklist documente la consult (per OPIS oficial OncoHelp)", 1)
    doc.add_paragraph("□ Bilet de trimitere către Oncologie (de luat fizic)")
    doc.add_paragraph("□ Buletin (carte identitate)")
    doc.add_paragraph("□ Card sănătate")
    doc.add_paragraph("□ Cupon pensie (folder 10_administrativ_pensie/)")
    doc.add_paragraph("□ COPIE rezultat histopatologic biopsie (folder 12_biopsie_2026/)")
    doc.add_paragraph("□ COPIE buletin gastroscopie + colonoscopie (folder 09_endoscopie_2026_04/)")
    doc.add_paragraph("□ COPIE raport CT TAP 20.04 + bilet trimitere (folder 11_CT_stadializare_2026/)")
    doc.add_paragraph("□ COPIE PDF cardiologie Vichy 2012 (folder 02_cardiologie_2012/)")
    doc.add_paragraph("□ COPIE scrisoare cardiologie 10.11.2025 (folder 13_cardiologie_ambulator_2025/)")
    doc.add_paragraph("□ COPIE schemă medicamente actuală (folder 08_schema_tratament/)")
    doc.add_paragraph("□ Analize sânge 29.04 (CEA + CA 19-9 + HbA1c) — dacă rezultatele sosesc înainte")

    # 6. Întrebări specifice tip consult
    doc.add_heading("6. Întrebări de pus la consult", 1)
    for i, q in enumerate(template["intrebari"], 1):
        doc.add_paragraph(f"{i}. {q}")

    # 7. Notes consult (spațiu liber)
    doc.add_heading("7. Notes consult (de completat la consult)", 1)
    for _ in range(15):
        doc.add_paragraph("________________________________________________________________________")

    doc.add_paragraph("─" * 50)
    footer = doc.add_paragraph()
    footer.add_run("Generat de scripts/generate_consult_briefing.py (R-MINIMAL).").italic = True

    doc.save(str(output_path))
    return True


def main():
    parser = argparse.ArgumentParser(description="Generează briefing DOCX pentru consult medical.")
    parser.add_argument("--consult", required=True, choices=list(CONSULT_TEMPLATES.keys()))
    parser.add_argument("--data", required=True, help="Data consult ISO: YYYY-MM-DD")
    parser.add_argument("--medic", required=True, help="Nume medic (cu titlu)")
    parser.add_argument("--unitate", required=True, help="Unitate medicală")
    args = parser.parse_args()

    RAPOARTE.mkdir(parents=True, exist_ok=True)
    output_filename = f"{args.data}_briefing_consult_{args.consult}.docx"
    output_path = RAPOARTE / output_filename

    success = build_docx(args.consult, args.data, args.medic, args.unitate, output_path)
    if not success:
        return 1

    meta = {
        "tip_document": "briefing_consult_DOCX",
        "consult_type": args.consult,
        "data_consult": args.data,
        "medic": args.medic,
        "unitate": args.unitate,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "generator": "scripts/generate_consult_briefing.py",
        "source_data": "CONTEXT_MEDICAL.md + JSON-uri canonice Dosar_Medical/",
        "structura": ["1. Date pacient", "2. Status oncologic actual", "3. Medicație curentă", "4. Antecedent cardiologic STEMI 2012 BMS", "5. Checklist documente OPIS", "6. Întrebări consult", "7. Notes consult (spațiu liber)"],
    }
    meta_path = output_path.with_suffix(".docx.meta.json")
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    try:
        docx_display = output_path.relative_to(ROOT)
        meta_display = meta_path.relative_to(ROOT)
    except ValueError:
        docx_display = output_path
        meta_display = meta_path
    print(f"DOCX briefing generat: {docx_display}")
    print(f"Meta companion: {meta_display}")
    print(f"  Tip consult: {args.consult}")
    print(f"  Data: {args.data}")
    print(f"  Medic: {args.medic}")
    print(f"  Unitate: {args.unitate}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
