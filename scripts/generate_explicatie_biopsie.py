#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generator DOCX — Explicatie rezultat biopsie 17.04.2026 pentru familie (Roland).

R17 (marcaje certitudine) + R18 (sync DASHBOARD) + R20 (mod de lucru).
Document destinat: Roland Petrila — limbaj accesibil, cu analogii.

Rulare: python scripts/generate_explicatie_biopsie.py
Output: Documente_Informative/EXPLICATIE_REZULTAT_BIOPSIE_2026-04-28.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Documente_Informative" / "EXPLICATIE_REZULTAT_BIOPSIE_2026-04-28.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x38, 0x76, 0x1D)
ORANGE = RGBColor(0xC0, 0x55, 0x04)
RED = RGBColor(0xC0, 0x00, 0x00)
GRAY_DARK = RGBColor(0x40, 0x40, 0x40)
TEXT = RGBColor(0x20, 0x20, 0x20)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement("w:" + edge)
            border.set(qn("w:val"), kwargs[edge].get("val", "single"))
            border.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            border.set(qn("w:color"), kwargs[edge].get("color", "auto"))
            tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_h1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_h2(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_h3(doc, text, color=GRAY_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_para(doc, text, bold=False, italic=False, size=11, color=TEXT, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_callout(doc, title, body, kind="info"):
    palette = {
        "info": ("D9E2F3", "2E75B6"),
        "warn": ("FCE4D6", "C05504"),
        "ok": ("E2EFDA", "38761D"),
        "urgent": ("FADBD8", "C00000"),
        "analogy": ("F2F2F2", "595959"),
    }
    fill, border = palette.get(kind, palette["info"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_border(
        cell,
        top={"sz": 8, "color": border},
        left={"sz": 24, "color": border},
        bottom={"sz": 8, "color": border},
        right={"sz": 8, "color": border},
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    r1 = p_title.add_run(title)
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(border)
    r1.font.name = "Calibri"
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(4)
    r2 = p_body.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEXT
    r2.font.name = "Calibri"
    doc.add_paragraph()


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def add_table_2col(doc, header_left, header_right, rows, color_zebra="F2F2F2"):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.autofit = False
    hdr = table.rows[0]
    for idx, txt in enumerate([header_left, header_right]):
        cell = hdr.cells[idx]
        set_cell_bg(cell, "1F4E79")
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    for ri, (left, right) in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        if ri % 2 == 0:
            for c in row_cells:
                set_cell_bg(c, color_zebra)
        for c, txt in zip(row_cells, [left, right]):
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(10.5)
            r.font.name = "Calibri"
            r.font.color.rgb = TEXT


# Use only ASCII quotes inside Python string literals; smart quotes via unicode escapes.
LDQ = "„"  # double low-9 quote ("low double quote")
RDQ = "”"  # right double quotation mark
EM = "—"   # em dash


def q(text):
    """Wrap text with Romanian-style smart quotes."""
    return LDQ + text + RDQ


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ============ COVER ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("EXPLICATIE\nREZULTAT BIOPSIE\n17 aprilie 2026")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Document pentru Roland " + EM + " explicatie simpla, fara jargon medical")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("Pacient: PETRILA VIOREL-MIHAI (66 ani)")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = TEXT
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Buletin Bioclinica nr. 26417A0362 / cod T26H06044\n"
        "Semnat: Dr. Glaja Romanita (medic primar anatomopatolog)\n"
        "Generat: 28 aprilie 2026"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run(
        "Document generat: 28 aprilie 2026\n"
        "Pentru consultul oncologic de luni 4 mai 2026 " + EM + " OncoHelp Timisoara"
    )
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    doc.add_page_break()

    # ============ TL;DR ============
    add_h1(doc, "În două rânduri")

    add_callout(
        doc,
        "Concluzia raportului în limbaj clar:",
        "Biopsia NU spune " + q("cancer confirmat") + " " + EM
        + " dar NU spune nici " + q("nu e cancer") + ". "
        "Spune " + q("nu pot decide din proba asta — e prea puțin țesut tumoral în ea") + ". "
        "Suspiciunea de cancer rămâne ridicată din CT și endoscopie. "
        "Următorul pas (decis de oncolog luni 4 mai): un test suplimentar pe blocul de țesut "
        "deja prelevat (IHC), sau o nouă biopsie țintită mai în profunzime.",
        kind="info",
    )

    # ============ 1. CITATUL TEXTUAL ============
    add_h1(doc, "1. Citatul textual al raportului")

    add_para(
        doc,
        "Anatomopatologul (Dr. Glăja Romanița, medic primar la Bioclinica) a analizat la microscop "
        "cele 2 fragmente de țesut prelevate de Dr. Noufal la endoscopia din 17 aprilie. "
        "Iată concluzia ei textuală, citată din raport:",
    )

    add_callout(
        doc,
        "Concluzia anatomopatologului (text original):",
        q("Ansamblul histologic, pe secțiuni seriate și în colorația uzuală, "
          "pledează pentru ȚESUT DE GRANULAȚIE pe fond de ULCERAȚIE CRONICĂ, "
          "fiind doar SUGESTIV pentru infiltrat carcinomatos.")
        + "\n\n"
        + q("De corelat cu datele endoscopice/imagistice (diagnostic histologic tumoral mult limitat de "
            "numărul mic al celulelor epiteliale atipice); eventuală evaluare imunohistochimică pentru "
            "diagnostic histologic de certitudine și conduită terapeutică."),
        kind="info",
    )

    add_para(
        doc,
        "În capitolul următor explic FIECARE termen folosit în raport — antet, descrieri, concluzie — "
        "prin firul narativ unic " + q("casa cu o pată suspectă pe perete") + ", așa încât să înțelegi "
        "fără jargon medical exact ce a văzut Dr. Glăja la microscop și DE CE concluzia ei e " + q("sugestiv")
        + " în loc de " + q("confirmat") + ".",
    )

    # ============ 2. ANATOMIA RAPORTULUI ============
    add_h1(doc, "2. Anatomia raportului — fiecare mențiune, explicată")

    add_h2(doc, "Firul narativ: casa cu o pată suspectă pe perete")

    add_callout(
        doc,
        "Povestea de bază pe care o vom folosi pe parcurs:",
        "Imaginează-ți pacientul ca pe o CASĂ. Esofagul tatălui tău e UN PERETE din această casă. "
        "Pe perete a apărut o PATĂ NEREGULATĂ — o crustă/leziune care s-a îngroșat circumferențial.\n\n"
        "Dr. Noufal (gastroenterologul) a fost INSPECTORUL care a venit cu lanterna (endoscopul) și a "
        "verificat peretele. Văzând pata, a luat cu o " + q("pensă") + " 2 firimituri MINUSCULE de tencuială "
        "(biopsia) și le-a trimis la laboratorul central din Timișoara.\n\n"
        "La laborator, EXPERTUL ANATOMOPATOLOG (Dr. Glăja) e ca un specialist care examinează "
        "firimiturile cu o " + q("lupă cu putere extremă") + " (microscopul) după ce le-a tratat cu "
        "" + q("vopsele speciale") + " (coloranți). Apoi scrie un RAPORT despre ce vede în firimituri — "
        "dar nu despre întreaga pată, ci doar despre fragmentele pe care le are în față.\n\n"
        "Pe acest fir narativ vom merge mai departe — termen cu termen, citat cu citat.",
        kind="analogy",
    )

    # --- 2.1 Antetul ---
    add_h2(doc, "2.1 Antetul — cine, ce, când, unde")

    add_para(doc, "Raportul începe cu un antet care identifică pacientul, proba și laboratorul:")
    add_bullet(doc, "PETRILA VIOREL MIHAI, M, 66 ani — pacientul.", bold_prefix="Pacient:")
    add_bullet(doc, "1590518024486 — codul numeric personal.", bold_prefix="CNP:")
    add_bullet(doc, "18.05.1959 (66 ani împliniți).", bold_prefix="Data nașterii:")
    add_bullet(doc, "Str. Vasile Goldiș 42, Nădlac, Arad.", bold_prefix="Adresa:")
    add_bullet(doc, "00003 Bioclinica Vlaicu (Arad) — locul unde au ajuns firimiturile imediat după prelevare.", bold_prefix="Punct recoltare:")
    add_bullet(doc, "26417A0362 — codul administrativ al întregii investigații.", bold_prefix="Buletin nr.:")
    add_bullet(doc, "17.04.2026 14:21 (în timpul endoscopiei la Genesis Arad).", bold_prefix="Recoltat:")
    add_bullet(doc, "Bioclinica SA, Calea Torontalului 1, Timișoara — laboratorul central de anatomopatologie.", bold_prefix="Lucrat:")
    add_bullet(doc, "28.04.2026 06:46 (când raportul a fost publicat pe portalul Bioclinica).", bold_prefix="Generat PDF:")

    add_callout(
        doc,
        "De ce e recoltat la Arad și lucrat la Timișoara?",
        "Bioclinica are puncte de recoltare în mai multe orașe (Vlaicu = punctul din Arad), "
        "dar examinarea histopatologică se face într-un laborator central cu echipamente specializate "
        "și experți medici primari (Dr. Glăja, Dr. Teoran).\n\n"
        "Pe firul nostru: Dr. Noufal a luat firimiturile la Arad, le-a împachetat și le-a trimis "
        "(tipic prin curier) la laboratorul central din Timișoara unde sunt experții cu lupele puternice. "
        "După analiză, raportul a fost publicat pe portal pe 28.04.",
        kind="analogy",
    )

    # --- 2.2 Cod diagnostic ---
    add_h2(doc, "2.2 " + q("Cod diagnostic T26H06044"))

    add_para(
        doc,
        "Codul T26H06044 e un IDENTIFICATOR UNIC al biopsiei tatălui tău în baza de date Bioclinica. "
        "El rămâne legat permanent de blocul de parafină (cubul de ceară cu firimiturile incluse) — "
        "așa anatomopatologul poate cere oricând " + q("felii noi tăiate din T26H06044") + " pentru analize "
        "suplimentare (cum ar fi IHC), fără să mai fie nevoie de o biopsie nouă.",
    )

    add_callout(
        doc,
        "Pe firul nostru:",
        "Codul T26H06044 e ca " + q("numărul de inventar") + " al unui obiect dintr-un muzeu — "
        "oricine din lume poate cere acel obiect anume folosind codul. Dacă oncologul de luni 4 mai "
        "decide IHC, va trimite la Bioclinica o solicitare cu codul T26H06044 și laboratorul va găsi "
        "exact CUBUL DE CEARĂ al tatălui tău (nu altul).",
        kind="analogy",
    )

    # --- 2.3 Examen histopatologic + colorație H&E ---
    add_h2(doc, "2.3 " + q("Examen histopatologic (colorația hematoxilină-eozină)"))

    add_h3(doc, "» Ce înseamnă " + q("histopatologic"))
    add_para(
        doc,
        "" + q("Histo") + " = țesut. " + q("Pato") + " = boală. " + q("Logic") + " = studiu. "
        "Histopatologic = examinarea ȚESUTURILOR la microscop pentru a căuta semne de boală. E diferit "
        "de " + q("citologic") + " (care studiază celule individuale, ex. test Babeș-Papanicolau), "
        "pentru că se uită la cum sunt celulele ORGANIZATE între ele și cu structurile din jur "
        "(vase de sânge, țesut conjunctiv, glande).",
    )

    add_h3(doc, "» Ce înseamnă " + q("colorația hematoxilină-eozină") + " (H&E)")
    add_para(
        doc,
        "E STANDARDUL UNIVERSAL în anatomopatologie. Două vopsele aplicate pe felii ultrasubțiri de țesut:",
    )
    add_bullet(doc, "colorează în ALBASTRU/VIOLET nucleii celulelor (partea cu ADN-ul).", bold_prefix="Hematoxilina")
    add_bullet(doc, "colorează în ROZ/ROȘU citoplasma și fibrele de țesut conjunctiv din jur.", bold_prefix="Eozina")

    add_callout(
        doc,
        "Pe firul nostru:",
        "Fără coloranți, firimiturile ar fi gri-transparente la microscop și ar fi imposibil să distingi "
        "structurile. Cu H&E, totul devine ca o HARTĂ CU DOUĂ CULORI: orașele albastre (nucleii celulelor — "
        "unde " + q("locuiește") + " ADN-ul) și șoselele/țesutul roz (citoplasma + fibrele). Toate "
        "biopsiile încep cu această colorație de bază; coloranții speciali (IHC) sunt o etapă suplimentară "
        "ulterioară.",
        kind="analogy",
    )

    # --- 2.4 Diagnostic clinic ---
    add_h2(doc, "2.4 " + q("Diagnostic clinic: Proces proliferativ"))

    add_para(
        doc,
        "Aceasta NU e concluzia anatomopatologului. E ce a SCRIS Dr. Noufal pe biletul de trimitere "
        "către laborator, ÎNAINTE de analiză:",
    )
    add_bullet(doc, "= se înmulțește activ, crește.", bold_prefix="Proliferativ")
    add_bullet(doc, "= ceva care se petrece, în desfășurare.", bold_prefix="Proces")
    add_para(
        doc,
        "Adică Dr. Noufal a scris: " + q("Bănuiesc că pe peretele esofagian se petrece ceva care creste "
        "activ. Vă rog să-mi spuneți ce e.") + " Nu spune " + q("cancer") + " — spune " + q("ceva care "
        "se înmulțește, analizați și spuneți-mi ce e") + ". E începutul investigației, nu concluzia.",
    )

    # --- 2.5 Piesa A + bloc parafină ---
    add_h2(doc, "2.5 " + q("Piesa A — Biopsie. Număr blocuri de parafină procesate: 1"))

    add_h3(doc, "» " + q("Piesa A"))
    add_para(
        doc,
        "Primul (și singurul) lot de fragmente prelevate. Dacă Dr. Noufal ar fi luat biopsii din 2 zone "
        "diferite (ex. esofag + stomac), ar fi avut " + q("Piesa A") + " și " + q("Piesa B") + ". Aici "
        "există o singură zonă biopsiată — leziunea esofagiană.",
    )

    add_h3(doc, "» " + q("Biopsie"))
    add_para(doc, "Prelevare de țesut viu pentru analiză la microscop.")

    add_h3(doc, "» " + q("Bloc de parafină") + " (procesare standard)")
    add_para(doc, "După ce firimiturile ajung la laborator, urmează o procedură standard:")
    add_bullet(doc, "Sunt " + q("fixate") + " într-o soluție (formol) — ca să nu se descompună.")
    add_bullet(doc, "Sunt incluse în PARAFINĂ topită (un fel de ceară), care se solidifică într-un MIC CUB.")
    add_bullet(doc, "Cubul e tăiat cu un cuțit special (microtom) în " + q("felii ultrasubțiri") + " — 3-5 micrometri (miimi de milimetru) — cum ai tăia un salam la felii foarte subțiri.")
    add_bullet(doc, "Feliile se pun pe LAME DE STICLĂ.")
    add_bullet(doc, "Lamele se colorează cu H&E (vezi 2.3) și se examinează la microscop.")

    add_h3(doc, "» " + q("Procesate: 1") + " (un singur bloc)")
    add_para(
        doc,
        "Cele DOUĂ firimituri prelevate de Dr. Noufal au fost atât de mici (sub-milimetrice) încât au "
        "încăput împreună într-un singur bloc de parafină. Asta e prima limitare materială — material "
        "puțin pentru analiză.",
    )

    add_callout(
        doc,
        "Pe firul nostru:",
        "Firimiturile de tencuială au fost băgate într-un cub de ceară. Cubul a fost tăiat în felii ca "
        "salamul, fiecare felie a fost lipită pe o " + q("oglindă") + " (lama de sticlă), vopsită cu "
        "albastru și roșu (H&E), și pusă sub lupa expertului. Pentru că firimiturile au fost minuscule, "
        "s-a făcut DOAR UN SINGUR cub — restul depinde de ce vede expertul în feliile rezultate.",
        kind="analogy",
    )

    # --- 2.6 Macroscopic ---
    add_h2(doc, "2.6 MACROSCOPIC — descrierea Dr. Teoran (medic specialist)")

    add_para(
        doc,
        "" + q("Macroscopic") + " = cum arată firimiturile la OCHIUL LIBER (sau lupă mică), ÎNAINTE de a "
        "fi tăiate în felii și colorate. E primul pas: o descriere fizică simplă.",
    )

    add_callout(
        doc,
        "Textul exact din raport:",
        q("Două piese biopsice cu aspect neregulat, cafenii, elastice, cu dimensiuni de 0,2/0,1/0,1 cm "
          "ambele. Se orientează în totalitate, fără secționare (T26H06044-A1)."),
        kind="info",
    )

    add_h3(doc, "» " + q("Două piese biopsice"))
    add_para(doc, "Dr. Noufal a prelevat DOI fragmente separate — nu unul singur.")

    add_h3(doc, "» " + q("aspect neregulat"))
    add_para(
        doc,
        "Nu sunt rotunde sau ovale — au formă neregulată, ruptă. E NORMAL pentru o biopsie cu pensa "
        "endoscopică (pensa rupe, nu taie precis).",
    )

    add_h3(doc, "» " + q("cafenii"))
    add_para(
        doc,
        "Culoare maro-deschis, brună. E culoarea normală a țesutului viu fixat în formol (proteinele se "
        "colorează ușor brun după fixare).",
    )

    add_h3(doc, "» " + q("elastice"))
    add_para(
        doc,
        "Țesutul are oarecare elasticitate când e atins — semn că e țesut moale, viu, NU necrozat dur, "
        "NU calcificat, NU putrid. Nimic alarmant aici; descrierea e standard pentru o biopsie de mucoasă.",
    )

    add_h3(doc, "» " + q("dimensiuni de 0,2/0,1/0,1 cm ambele") + " — ATENȚIE, cifră CHEIE")
    add_para(
        doc,
        "0,2 cm = 2 milimetri (cea mai mare dimensiune a firimiturii). 0,1 cm = 1 milimetru (celelalte 2 "
        "dimensiuni). Volum aproximativ: ~2 mm³ per firimitură. AMBELE firimituri au exact aceleași "
        "dimensiuni (2 × 1 × 1 mm).",
    )

    add_callout(
        doc,
        "Pe firul nostru — limitarea cheie:",
        "Fiecare firimitură e MAI MICĂ decât un bob de orez. Inspectorul a luat 2 grăunțe minuscule de "
        "tencuială — fiecare cât un grăunte de nisip ușor mai mare. Cu așa de puțin material, ai foarte "
        "puține " + q("camere") + " de explorat la microscop.\n\n"
        "ASTA E LIMITAREA principală care a făcut concluzia să fie " + q("doar sugestiv") + " — nu suficient "
        "pentru " + q("certitudine") + ".",
        kind="warn",
    )

    add_h3(doc, "» " + q("Se orientează în totalitate, fără secționare (T26H06044-A1)"))
    add_para(doc, "Trei elemente tehnice de procesare:")
    add_bullet(doc, "= se așează cu o anumită orientare în blocul de parafină, ca să poată fi tăiate corect (perpendicular pe suprafața lor — să se vadă " + q("toată grosimea") + ").", bold_prefix="se orientează")
    add_bullet(doc, "= se folosesc TOATE cele 2 firimituri (nu se aruncă nimic, nu se păstrează rezerve — pentru că sunt prea mici).", bold_prefix="în totalitate")
    add_bullet(doc, "= nu sunt tăiate în prealabil în bucățele și mai mici — sunt incluse așa cum au fost prelevate.", bold_prefix="fără secționare")
    add_bullet(doc, "= cod subtotal (T26H06044 = codul biopsiei, A1 = primul bloc din " + q("Piesa A") + ").", bold_prefix="T26H06044-A1")

    # --- 2.7 Microscopic ---
    add_h2(doc, "2.7 MICROSCOPIC — descrierea Dr. Glăja (medic primar)")

    add_para(
        doc,
        "Aici vine analiza la microscop, după ce feliile au fost tăiate, lipite pe lame și colorate H&E. "
        "Descrierea e ÎMPĂRȚITĂ pe cele DOUĂ firimituri prelevate — pentru că au conținut țesut diferit.",
    )

    add_h3(doc, "FRAGMENTUL 1 — descriere detaliată")

    add_callout(
        doc,
        "Textul exact (fragmentul 1):",
        q("Un fragment tisular biopsic constituit din numeroase structuri vasculare cu endoteliul "
          "tumefiat, unele hiperemiate / cu marginație leucocitară și manșon leucocitar, cu orientare "
          "perpendiculară pe suprafața acoperită parțial de detritus și necroză fibrinoidă; prezente "
          "deopotrivă elemente celulare inflamatorii mononucleate; aparent se disting celule epitelioide "
          "de talie medie, cu nucleul nucleolat, cu nucleol eozinofil și citoplasma cantitativ moderată, "
          "palid colorată / slab eozinofilă, singulare/grupate."),
        kind="info",
    )

    # --- termen 1: structuri vasculare cu endoteliu tumefiat ---
    add_h3(doc, "» " + q("structuri vasculare cu endoteliul tumefiat"))
    add_bullet(doc, "= vase de sânge mici (capilare).", bold_prefix="Structuri vasculare")
    add_bullet(doc, "= stratul subțire de celule care căptușește pe interior peretele unui vas de sânge — ca o tapiserie pe interiorul unui tub.", bold_prefix="Endoteliul")
    add_bullet(doc, "= umflat, mărit de volum.", bold_prefix="Tumefiat")
    add_callout(
        doc,
        "Pe firul nostru:",
        "În firimitura 1, expertul vede ȚEVI DE CUPRU SUBȚIRI (capilare) cu pereții lor interiori umflați "
        "— ca o țeavă veche cu rugină pe dinăuntru: pereții devin mai groși, lumenul se îngustează. Asta "
        "arată INFLAMAȚIE locală — vasele reacționează la o agresiune (o rană, o iritație, un infiltrat).",
        kind="analogy",
    )

    # --- termen 2: hiperemiate, marginație, manșon leucocitar ---
    add_h3(doc, "» " + q("unele hiperemiate / cu marginație leucocitară și manșon leucocitar"))
    add_bullet(doc, "" + EM + " = pline cu sânge (HIPER = mult, EMIE = sânge). Vasele sunt dilatate, congestionate.", bold_prefix="Hiperemiate")
    add_bullet(doc, "= globulele albe (leucocite) se așază LA MARGINEA peretelui interior al vasului, gata să iasă în țesut. E etapa 1 a unei reacții inflamatorii — globulele albe vin la " + q("locul faptei") + ".", bold_prefix="Marginație leucocitară")
    add_bullet(doc, "= globulele albe deja au IEȘIT din vase și formează un INEL de jur-împrejurul vasului, ca o gardă strânsă în jurul țevii. " + q("Manșon") + " = tubul îngust, gulerul.", bold_prefix="Manșon leucocitar")
    add_callout(
        doc,
        "Pe firul nostru:",
        "Țevile sunt PLINE cu apă (hiperemie). Pe pereții lor s-au adunat ZECI de polițiști gata să sară "
        "(marginație). Alți polițiști deja au coborât din țeavă și formează un CERC STRÂNS de jur-împrejurul "
        "ei (manșon). E imaginea CLASICĂ a unei zone unde a sosit " + q("echipa de intervenție inflamatorie")
        + " — semn de inflamație activă în jurul leziunii.",
        kind="analogy",
    )

    # --- termen 3: orientare perpendiculară ---
    add_h3(doc, "» " + q("cu orientare perpendiculară pe suprafață"))
    add_para(
        doc,
        "Vasele sunt așezate VERTICAL pe suprafața mucoasei (urcă din profund spre superficial). "
        "E orientarea NORMALĂ pentru capilarele care hrănesc mucoasa — nu o anomalie. E un detaliu de "
        "context spațial care îi arată anatomopatologului că secțiunea a fost tăiată corect.",
    )

    # --- termen 4: detritus + necroza fibrinoida ---
    add_h3(doc, "» " + q("acoperită parțial de detritus și necroză fibrinoidă"))
    add_bullet(doc, "= resturi celulare moarte, fragmente de țesut degradat — ca " + q("murdăria") + " care se acumulează la suprafața unei răni.", bold_prefix="Detritus")
    add_bullet(doc, "= un tip particular de moarte celulară unde țesutul mort e îmbibat cu FIBRINĂ (proteina din cheaguri de sânge). Apare ca o pojghiță roșie-roz, omogenă, fără structură celulară clară.", bold_prefix="Necroză fibrinoidă")
    add_callout(
        doc,
        "Pe firul nostru:",
        "SUPRAFAȚA peretelui (mucoasa) e acoperită PARȚIAL cu o pojghiță maro-roșiatică formată din: "
        "resturi de tencuială căzută (detritus) + un STRAT DE PANSAMENT NATURAL cu cheaguri "
        "(necroză fibrinoidă). E exact ce ai vedea pe o RANĂ DESCHISĂ VECHE care a tot încercat să se "
        "vindece, dar n-a reușit complet — crustă peste crustă, vie și moartă amestecate.",
        kind="analogy",
    )

    # --- termen 5: elemente celulare inflamatorii mononucleate ---
    add_h3(doc, "» " + q("prezente deopotrivă elemente celulare inflamatorii mononucleate"))
    add_bullet(doc, "= globule albe (leucocite) responsabile de inflamație.", bold_prefix="Elemente celulare inflamatorii")
    add_bullet(doc, "= au UN SINGUR nucleu (spre deosebire de neutrofile, care au mai mulți lobi nucleari). Mononuclearele includ: limfocite (memorie imunitară), plasmocite (produc anticorpi), macrofage (curăță detritusul), monocite (precursorii macrofagelor).", bold_prefix="Mononucleate")
    add_callout(
        doc,
        "Pe firul nostru:",
        "Pe lângă " + q("echipa de intervenție rapidă") + " (neutrofile, care vin în inflamațiile acute), "
        "au sosit și ECHIPELE DE SPECIALIȘTI PE TERMEN LUNG — limfocite, macrofage, plasmocite. "
        "E semnul unei INFLAMAȚII CRONICE (de durată), nu doar acute. Asta confirmă " + q("ulcerația cronică")
        + " din concluzie.",
        kind="analogy",
    )

    # --- termen 6: aparent se disting celule epitelioide ---
    add_h3(doc, "» " + q("aparent se disting celule epitelioide de talie medie"))
    add_bullet(doc, "— anatomopatologul folosește acest cuvânt CU PRECAUȚIE. Înseamnă " + q("par să fie") + ", " + q("cred că văd") + ", NU " + q("sunt sigur") + ".", bold_prefix="Aparent")
    add_bullet(doc, "= celule cu aspect de celulă EPITELIALĂ (din epiteliu — stratul de celule care căptușește mucoasele). E TIPUL de celulă care formează ADENOCARCINOAMELE (cancere care apar pornind din epiteliu).", bold_prefix="Celule epitelioide")
    add_bullet(doc, "— nu sunt nici foarte mari, nici foarte mici (dimensiune intermediară).", bold_prefix="De talie medie")
    add_callout(
        doc,
        "Pe firul nostru — momentul-cheie:",
        "Aici e exact PUNCTUL DE INCERTITUDINE. Expertul vede CÂTEVA celule care " + q("seamănă") + " cu "
        "cele care formează căptușeala peretelui (epiteliu). Dacă ar fi fost MULTE și CLAR ATIPICE, ar fi "
        "spus " + q("CANCER") + ". Dar fiind PUȚINE și " + q("aparent") + ", spune doar " + q("par să fie")
        + ". E motivul pentru care concluzia nu poate fi clară.",
        kind="warn",
    )

    # --- termen 7: nucleu nucleolat, nucleol eozinofil ---
    add_h3(doc, "» " + q("cu nucleul nucleolat, cu nucleol eozinofil"))
    add_bullet(doc, "= " + q("creierul") + " celulei, conține ADN-ul.", bold_prefix="Nucleu")
    add_bullet(doc, "= are NUCLEOL (un punct mai DENS din interiorul nucleului — fabrica de ribozomi, locul unde celula " + q("scrie") + " proteinele).", bold_prefix="Nucleolat")
    add_bullet(doc, "= nucleolul se colorează în ROȘU/ROZ cu eozina — semn că e plin cu proteine și " + q("lucrează intens") + ".", bold_prefix="Nucleol eozinofil")
    add_callout(
        doc,
        "Pe firul nostru — atenție la nuanță:",
        "Celulele nu doar EXISTĂ în firimitură, ci au și " + q("motorașe interne mari și active") + " "
        "(nucleoli prominenți eozinofili) — semn că celulele sunt în PROLIFERARE rapidă, exact așa cum apar "
        "celulele tumorale.\n\n"
        "DAR " + EM + " atenție: nucleolul prominent NU e patognomonic (specific cancerului). Apare și în "
        "celule NORMALE care " + q("lucrează tare") + " — celule reactive în inflamație, celule de "
        "regenerare. De aceea Dr. Glăja nu poate spune " + q("cancer") + " doar pe baza acestui detaliu.",
        kind="analogy",
    )

    # --- termen 8: citoplasma moderată, palid colorată/slab eozinofilă ---
    add_h3(doc, "» " + q("citoplasma cantitativ moderată, palid colorată / slab eozinofilă"))
    add_bullet(doc, "= " + q("interiorul") + " celulei (în afara nucleului), unde au loc procesele metabolice.", bold_prefix="Citoplasma")
    add_bullet(doc, "= celula are nici puțină, nici multă citoplasma.", bold_prefix="Cantitativ moderată")
    add_bullet(doc, "= citoplasma se colorează SLAB cu eozina (apare roz pal, nu intens roz/roșu).", bold_prefix="Palid colorată / slab eozinofilă")
    add_callout(
        doc,
        "Pe firul nostru:",
        "Celulele suspecte au " + q("corpul") + " de mărime medie, ușor decolorat — nici intens roșu "
        "(ca celulele musculare), nici transparent. E o caracteristică INTERMEDIARĂ, ambiguă, care nu trage "
        "spre o concluzie clară (nici tipic tumoral, nici tipic reactiv).",
        kind="analogy",
    )

    # --- termen 9: singulare/grupate ---
    add_h3(doc, "» " + q("singulare / grupate"))
    add_para(doc, "Celulele atipice sunt văzute uneori una câte una izolate (SINGULARE), uneori în mici aglomerări de 2-3 celule împreună (GRUPATE).")
    add_callout(
        doc,
        "Pe firul nostru — al doilea criteriu lipsă:",
        "Dacă expertul ar fi văzut GRUPURI MARI de celule atipice, formând " + q("cuiburi") + " sau "
        "structuri organizate (glande tumorale, formațiuni infiltrative), ar fi putut spune cu mai multă "
        "certitudine " + q("ADENOCARCINOM") + ". Aici sunt prea răzlețe pentru a putea identifica un "
        + q("model arhitectural tumoral") + ".",
        kind="warn",
    )

    add_h3(doc, "FRAGMENTUL 2 — descriere scurtă")

    add_callout(
        doc,
        "Textul exact (fragmentul 2):",
        q("Un fragment tisular mic reprezentat în exclusivitate de epiteliu stratificat scuamos "
          "necheratinizat, cu aspecte de exocitoză (granulocite neutrofile) și de extravazate hematice, "
          "fără suport conjunctiv."),
        kind="info",
    )

    # --- termen 10: epiteliu stratificat scuamos necheratinizat ---
    add_h3(doc, "» " + q("epiteliu stratificat scuamos necheratinizat"))
    add_para(doc, "STRATUL NORMAL care căptușește esofagul. Patru cuvinte explicate:")
    add_bullet(doc, "= strat de celule de suprafață.", bold_prefix="Epiteliu")
    add_bullet(doc, "= mai multe straturi de celule SUPRAPUSE (nu un singur rând).", bold_prefix="Stratificat")
    add_bullet(doc, "= forma celulelor e turtită, ca SOLZII DE PEȘTE (squamous = solz în latină).", bold_prefix="Scuamos")
    add_bullet(doc, "= stratul de la suprafață NU produce cheratină (proteina dură care face pielea uscată). Mucoasele interne sunt umede, moi.", bold_prefix="Necheratinizat")
    add_callout(
        doc,
        "Pe firul nostru:",
        "Al doilea firimitur (cel mai mic) e DOAR TAPISERIA NORMALĂ a peretelui — fără pată, fără "
        "leziune, fără celule suspecte. Doar mucoasa esofagiană sănătoasă, cum trebuie să arate. Asta e "
        "un lucru BUN " + EM + " înseamnă că zona din jurul leziunii e structural normală.",
        kind="ok",
    )

    # --- termen 11: exocitoză (granulocite neutrofile) ---
    add_h3(doc, "» " + q("aspecte de exocitoză (granulocite neutrofile)"))
    add_bullet(doc, "= IEȘIREA celulelor (în cazul de față, globule albe) DIN vase ÎN epiteliu/țesut. " + q("Exo") + " = în afara, " + q("cytosis") + " = mișcare celulară.", bold_prefix="Exocitoză")
    add_bullet(doc, "= un tip de globule albe cu NUCLEU MULTI-LOBAT și granule fine în citoplasma. Sunt " + q("pompierii") + " sistemului imun — primii care vin la o inflamație acută.", bold_prefix="Granulocite neutrofile")
    add_callout(
        doc,
        "Pe firul nostru:",
        "În al 2-lea firimitur (mucoasa normală), se văd globule albe care INFILTREAZĂ țesutul — semn de "
        "inflamație activă chiar și aici. E ca și cum chiar și în zona FĂRĂ pată a peretelui se găsesc "
        "câțiva pompieri în alertă, semn că leziunea iritează zone vecine.",
        kind="analogy",
    )

    # --- termen 12: extravazate hematice ---
    add_h3(doc, "» " + q("extravazate hematice"))
    add_bullet(doc, "= ieșit DIN vasele de sânge (extra = în afara, vas = vas).", bold_prefix="Extravazate")
    add_bullet(doc, "= sânge.", bold_prefix="Hematice")
    add_para(
        doc,
        "Vase mici care s-au " + q("spart") + " și au lăsat să curgă sânge în țesutul din jur (mici "
        "hemoragii punctiforme). Pe biopsia esofagiană e FRECVENT — pensa de biopsie poate produce mici "
        "extravazate la prelevare. NU e un semn alarmant.",
    )

    # --- termen 13: fără suport conjunctiv ---
    add_h3(doc, "» " + q("fără suport conjunctiv") + " — limitarea tehnică-cheie")
    add_bullet(doc, "= stratul de țesut conjunctiv (lamina propria + submucoasa) care se află SUB epiteliu și îl SUSȚINE.", bold_prefix="Suport conjunctiv")
    add_para(
        doc,
        "În fragmentul 2 NU e prezent — adică firimitura conține DOAR EPITELIUL SUPERFICIAL, nu și "
        "straturile profunde.",
    )
    add_callout(
        doc,
        "Pe firul nostru — DE CE biopsia a fost insuficientă:",
        "Al 2-lea firimitur e DOAR " + q("tapiseria de pe perete") + " (epiteliul) — FĂRĂ tencuiala "
        "suportoare de dedesubt (țesut conjunctiv). Pensa de biopsie a luat doar SUPRAFAȚA — nu a ajuns mai "
        "adânc.\n\n"
        "ASTA e motivul tehnic pentru care biopsia n-a putut surprinde STRATUL TUMORAL PROFUND — "
        "fragmentele au fost prea superficiale. Tumorile esofagiene infiltrează ÎN PROFUNZIME (lamina "
        "propria, submucoasa, musculara), iar pensa endoscopică ajunge tipic doar la suprafață.",
        kind="warn",
    )

    # --- 2.8 Concluzia ---
    add_h2(doc, "2.8 Concluzia — explicat termen cu termen")

    add_callout(
        doc,
        "Textul exact al concluziei:",
        q("Ansamblul histologic, pe secțiuni seriate și în colorația uzuală, pledează pentru ȚESUT DE "
          "GRANULAȚIE pe fond de ULCERAȚIE CRONICĂ, fiind doar SUGESTIV pentru infiltrat carcinomatos."),
        kind="info",
    )

    add_h3(doc, "» " + q("Ansamblul histologic, pe secțiuni seriate și în colorația uzuală"))
    add_bullet(doc, "= imaginea de ansamblu, cum arată întregul țesut analizat.", bold_prefix="Ansamblul histologic")
    add_bullet(doc, "= mai multe felii TĂIATE LA RÂND din același bloc de parafină, la diferite niveluri (3-5 felii diferite). Ai astfel o imagine din mai multe " + q("planuri") + ", nu doar dintr-unul.", bold_prefix="Secțiuni seriate")
    add_bullet(doc, "= colorația H&E (vezi 2.3) — standardul universal.", bold_prefix="Colorația uzuală")
    add_para(doc, "Tradus: " + q("Pe baza tuturor feliilor pe care le-am examinat, cu coloranții obișnuiți, văd...") + ".")

    add_h3(doc, "» " + q("pledează pentru"))
    add_para(
        doc,
        "Termen anatomopatologic pentru " + q("seamănă cel mai mult cu") + ", " + q("aspectul dominant "
        "este") + ". E o formulare PRUDENTĂ — nu spune " + q("este") + ", ci " + q("argumentează în "
        "favoarea") + ". Lasă deschisă posibilitatea altor interpretări.",
    )

    add_h3(doc, "» " + q("țesut de granulație pe fond de ulcerație cronică"))
    add_bullet(doc, "= țesut de REPARAȚIE, ca o " + q("crustă vie") + " care apare când o suprafață e rănită și organismul încearcă să se vindece. Conține multe vase mici noi (capilare), fibre de țesut și celule inflamatorii.", bold_prefix="Țesut de granulație")
    add_bullet(doc, "= o RANĂ care durează DE MULT timp (săptămâni-luni). " + q("Ulcerație") + " = pierdere de țesut superficial; " + q("cronică") + " = de durată.", bold_prefix="Ulcerație cronică")
    add_para(doc, "Acest tip de țesut apare în multe situații:")
    add_bullet(doc, "ulcer peptic (rană de stomac/esofag fără cancer)")
    add_bullet(doc, "esofagită severă cronică")
    add_bullet(doc, "PESTE / DEASUPRA unei tumori (zona de crustă care acoperă tumora propriu-zisă)")
    add_para(
        doc,
        "Cu alte cuvinte: țesutul de granulație SINGUR nu spune nici cancer, nici nu e cancer. E doar un "
        "STRAT SUPERFICIAL INFLAMAT — care POATE acoperi o tumoare, sau POATE fi o ulcerație non-tumorală.",
    )

    add_h3(doc, "» " + q("doar SUGESTIV pentru infiltrat carcinomatos"))
    add_bullet(doc, "= termen ANATOMOPATOLOGIC. Înseamnă " + q("indicii da, certitudine nu") + ". Sunt câteva semne, dar nu suficiente pentru un diagnostic ferm.", bold_prefix="Sugestiv")
    add_bullet(doc, "= INFILTRAT CARCINOMATOS = celule de carcinom (un tip de cancer pornit din epiteliu) care infiltrează (= pătrund) în țesuturile din jur. " + q("Adenocarcinom") + " = un sub-tip specific de carcinom care provine din celule glandulare, frecvent în esofag-stomac.", bold_prefix="Infiltrat carcinomatos")
    add_callout(
        doc,
        "Pe firul nostru — sinteza expertului:",
        "Ce a văzut Dr. Glăja, în limba ta:\n\n"
        "" + q("Am examinat cu atenție firimiturile, am tăiat felii la mai multe nivele, am colorat "
        "uzual. Pe baza a ce am văzut, IMPRESIA DOMINANTĂ e CRUSTĂ VIE PE O RANĂ VECHE care încearcă "
        "să se vindece (granulație + ulcerație cronică). Am văzut și CÂTEVA CELULE care POT fi de "
        "cancer — sunt indicii. Dar firimiturile au fost prea mici, am avut prea puține celule "
        "atipice, nu pot să spun cu certitudine că e cancer. Sunt SUGESTII, nu DOVEZI.") + "",
        kind="analogy",
    )

    # --- 2.9 Nota laboratorului ---
    add_h2(doc, "2.9 Nota laboratorului — recomandarea finală")

    add_callout(
        doc,
        "Textul exact al notei:",
        q("De corelat cu datele endoscopice/imagistice (diagnostic histologic tumoral mult limitat de "
          "numărul mic al celulelor epiteliale atipice); eventuala evaluare imunohistochimică pentru "
          "diagnostic histologic de certitudine și conduită terapeutică."),
        kind="info",
    )

    add_h3(doc, "» " + q("De corelat cu datele endoscopice/imagistice"))
    add_para(
        doc,
        "= " + q("Vă rog să luați în considerare ce s-a văzut la endoscopie și la CT/RMN, NU doar "
        "raportul meu") + ". Anatomopatologul își pune concluzia în CONTEXTUL clinic global. Singur, "
        "raportul nu e suficient pentru o decizie.",
    )

    add_h3(doc, "» " + q("diagnostic histologic tumoral mult limitat de numărul mic al celulelor epiteliale atipice"))
    add_para(
        doc,
        "Tradus exact: " + q("Nu pot pune diagnosticul de tumoră cu certitudine, pentru că am avut prea "
        "puține celule cu aspect atipic în firimiturile primite") + ". E o RECUNOAȘTERE EXPLICITĂ a "
        "limitării materialului. NU e o eroare a Dr. Glăja — e adevărul tehnic al situației.",
    )

    add_h3(doc, "» " + q("eventuala evaluare imunohistochimică pentru diagnostic histologic de certitudine"))
    add_bullet(doc, "= " + q("ar fi util să faceți IHC") + " (NU obligatoriu, dar puternic recomandat).", bold_prefix="Eventuala evaluare imunohistochimică")
    add_bullet(doc, "= ca să avem confirmarea HISTOLOGICĂ (din țesut, prin microscop) că E sau NU cancer.", bold_prefix="Pentru diagnostic histologic de certitudine")
    add_bullet(doc, "= și ca să știm ce TRATAMENT să aplicăm (FLOT, imunoterapie, terapie țintită — depinde de markeri).", bold_prefix="Și conduită terapeutică")
    add_callout(
        doc,
        "Pe firul nostru — recomandarea expertului:",
        "Expertul scrie, în concluzia lui finală: " + q("Asta-i tot ce pot să spun cu firimiturile astea. "
        "Dacă vreți răspuns clar (cancer DA / cancer NU), faceți teste suplimentare cu coloranți speciali "
        "(IHC) pe ce mai am eu aici (blocul T26H06044). De aceste teste depinde ce tratament veți alege.")
        + "\n\nEXACT acesta e motivul pentru care la consultul de luni 4 mai cu Dr. Anater, prima decizie "
        "va fi: IHC pe blocul existent, sau rebiopsie țintită nouă.",
        kind="ok",
    )

    # --- 2.10 Cei doi medici ---
    add_h2(doc, "2.10 Cei doi medici care au semnat raportul")

    add_h3(doc, "Dr. Teoran Samy Ștefan — medic specialist (cod parafă G70575)")
    add_para(
        doc,
        "A făcut DESCRIEREA MACROSCOPICĂ (cum arătau firimiturile la ochiul liber). " + q("Specialist") + " "
        "= nivel de competență 2 (după rezidențiat de 4 ani + examen specialitate). Are dreptul să "
        "examineze biopsii și să facă diagnostice.",
    )

    add_h3(doc, "Dr. Glăja Romanița — medic PRIMAR anatomopatolog (cod parafă 367427)")
    add_para(
        doc,
        "A făcut DESCRIEREA MICROSCOPICĂ + a SEMNAT raportul cu CONCLUZIA. " + q("Primar") + " = nivel "
        "de competență 3 (cel mai înalt: specialist + minim 5 ani experiență + examen primariat). E cea mai "
        "înaltă autoritate medicală în domeniul anatomopatologiei. Codul 367427 e codul unic în Registrul "
        "Național al Medicilor.",
    )

    add_callout(
        doc,
        "Pe firul nostru:",
        "La inspecția peretelui sunt 2 oameni:\n\n"
        "• " + q("Inspectorul JUNIOR") + " (medic specialist Teoran) face descrierea " + q("cum arată "
        "firimiturile") + " la primă vedere.\n\n"
        "• " + q("Inspectorul ȘEF") + " (medic primar Glăja, expert recunoscut) ia firimiturile la "
        "microscop, scrie raportul oficial și pune SEMNĂTURA OFICIALĂ. Concluzia raportului e "
        "RESPONSABILITATEA Dr. Glăja, nu a Dr. Teoran. Semnătura ei + ștampila + codul 367427 dau "
        "raportului GREUTATE LEGALĂ ȘI MEDICALĂ.",
        kind="analogy",
    )

    add_para(
        doc,
        "Dr. Glăja a semnat raportul cu data " + q("27.04.2026") + " (când a finalizat analiza). "
        "PDF-ul a fost generat automat de sistem pe " + q("28.04.2026 06:46") + " (când a fost publicat "
        "pe portalul Bioclinica).",
        italic=True,
    )

    # --- 2.11 Sinteza microscopică ---
    add_h2(doc, "2.11 Pe scurt — ce înseamnă ÎMPREUNĂ toate aceste detalii")

    add_callout(
        doc,
        "Sinteza descrierii microscopice, în limba ta:",
        "Pe firimiturile prelevate, expertul a văzut:\n\n"
        "• o SUPRAFAȚĂ ULCERATĂ acoperită cu o crustă vie (granulație + necroză fibrinoidă);\n"
        "• vase de sânge inflamate, cu globule albe care invadează țesutul (mononucleare cronice + neutrofile acute);\n"
        "• mici hemoragii punctiforme (extravazate);\n"
        "• CÂTEVA celule cu aspect care POATE indica adenocarcinom (epitelioide, cu nucleoli prominenți), "
        "dar PUȚINE și răzlețe;\n"
        "• un al doilea fragment a fost doar mucoasă superficială fără strat profund.\n\n"
        "În concluzie: NU a putut vedea miezul tumorii din cauza dimensiunii mici a probelor și a faptului "
        "că ele au fost prelevate din zona de crustă superficială, nu din profunzime. A văzut INDICII de "
        "cancer, dar nu DOVADĂ histologică.",
        kind="info",
    )

    # ============ CE NU ÎNSEAMNĂ ============
    add_h1(doc, "3. Ce NU înseamnă acest rezultat")

    add_callout(
        doc,
        "X NU înseamnă " + q("nu e cancer, e doar o ulcerație"),
        "Suspiciunea de cancer rămâne ridicată — biopsia n-a putut DOVEDI cancerul, "
        "dar n-a putut nici să-l EXCLUDĂ. CT-ul a arătat un proces tumoral T3-T4 infiltrativ, "
        "endoscopia a arătat o stenoză " + q("nedepășibilă") + ". Aceste lucruri NU s-au schimbat după biopsie.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("e confirmat cancer"),
        "Sunt indicii sugestive, dar nu certitudine histologică. "
        "Un diagnostic oficial de cancer cere ca anatomopatologul să spună " + q("cancer confirmat")
        + " — iar acest raport spune doar " + q("sugestiv pentru cancer")
        + ". E o nuanță importantă legal și medical.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("trebuie să luăm tot procesul de la zero"),
        "Probabil oncologul va decide să facă un test suplimentar (IHC) "
        "PE BLOCUL DEJA EXISTENT (T26H06044) — fără reintervenție. Asta poate da răspunsul în 3-7 zile.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("întârziem tratamentul"),
        "Consultul de luni 4 mai rămâne valid și esențial. "
        "Dr. Anater va decide acolo opțiunea (IHC vs rebiopsie) — iar oricare dintre ele "
        "se mișcă în paralel cu pregătirea dosarului oncologic + opinia chirurgului.",
        kind="warn",
    )

    # ============ CE URMEAZA ============
    add_h1(doc, "4. Ce urmează — 3 opțiuni standard")

    add_para(
        doc,
        "Anatomopatologul recomandă explicit în nota laboratorului: "
        + q("eventuală evaluare imunohistochimică pentru diagnostic histologic de certitudine") + ".",
    )
    add_para(
        doc,
        "Concret, oncologul de luni 4 mai (Dr. Anater) va alege una dintre aceste 3 căi:",
    )

    add_h2(doc, "Opțiunea 1 — IHC pe blocul existent (recomandare anatomopatolog)")

    add_callout(
        doc,
        "Ce e IHC (imunohistochimie)?",
        "Anatomopatologul ia blocul de parafină T26H06044 (deja la Bioclinica Timișoara) "
        "și aplică niște coloranți speciali numiți markeri pe felii noi din același țesut. "
        "Acești markeri se atașează doar de proteinele specifice celulelor canceroase. "
        "Dacă celulele atipice se colorează cu markerii potriviți → confirmare cancer + tip exact.",
        kind="ok",
    )

    add_para(doc, "Markerii tipici care se cer pentru o leziune eso-gastrică:")
    add_bullet(doc, "Pan-CK / CK7 / CK20 — confirmă originea epitelială a celulelor atipice")
    add_bullet(doc, "CDX-2 — pozitiv în adenocarcinoame de origine intestinală/eso-gastrică")
    add_bullet(doc, "p53 — markerul mutațiilor oncogene clasice")
    add_bullet(doc, "Dacă se confirmă cancer: HER2, PD-L1, MSI — ghidează decizia terapeutică (FLOT, imunoterapie, trastuzumab)")

    add_table_2col(
        doc,
        "AVANTAJE",
        "DEZAVANTAJE",
        [
            ("Rapid (3-7 zile rezultat)", "Dacă atipia e prea redusă, IHC poate fi tot inconcluziv"),
            ("Fără reintervenție pentru pacient", "Se face pe puținele celule atipice deja existente"),
            ("Recomandat explicit de Dr. Glăja", "Necesită trimitere de la oncolog la Bioclinica"),
            ("Cost moderat (~300-600 lei)", "Decizia depinde de Dr. Anater"),
        ],
    )

    add_h2(doc, "Opțiunea 2 — Rebiopsie țintită endoscopică")

    add_para(
        doc,
        "O nouă endoscopie cu prelevare de fragmente mai numeroase și mai profunde, eventual sub ghidaj "
        "EUS (ecoendoscopie — un endoscop cu o sondă de ecografie integrată care poate vedea stratul tumoral "
        "profund înainte de a preleva, ca un GPS pentru biopsie).",
    )

    add_table_2col(
        doc,
        "AVANTAJE",
        "DEZAVANTAJE",
        [
            ("Randament diagnostic superior", "Necesită reintervenție (anestezie, programare)"),
            ("Poate preleva din stratul tumoral profund", "Adaugă 1-3 săptămâni la timeline"),
            ("EUS vede stadiul T cu precizie", "Cost suplimentar (poate fi în CAS)"),
            ("Standard NCCN/ESMO la biopsie inconcluzivă", "Risc minor de sângerare/perforație"),
        ],
    )

    add_h2(doc, "Opțiunea 3 — Combinare IHC + rebiopsie în paralel")

    add_para(
        doc,
        "Cea mai eficientă în timp dacă oncologul vrea răspuns rapid + plan B în caz de IHC inconcluziv. "
        "Decizie tipică în comisiile tumor board când rezultatul histologic ține în loc planul terapeutic.",
    )

    # ============ DE CE E IMPORTANT ============
    add_h1(doc, "5. De ce e important să avem răspunsul în câteva săptămâni")

    add_para(
        doc,
        "Diagnosticul histologic de certitudine e poarta de intrare în orice protocol oncologic. Fără el:",
    )
    add_bullet(doc, "Nu se poate iniția chemoterapie (FLOT) — riscul de a trata o leziune non-tumorală cu medicamente toxice e inacceptabil")
    add_bullet(doc, "Nu se poate planifica chirurgia oncologică")
    add_bullet(doc, "Nu se pot accesa terapii țintite (HER2, PD-L1) — care depind de markeri moleculari obținuți doar prin IHC")

    add_callout(
        doc,
        "În cifre concrete (timeline):",
        "Dacă Dr. Anater alege IHC luni 4.05 → rezultat ~10-12 mai → comisie + plan terapeutic ~mijlocul lunii mai.\n"
        "Dacă alege rebiopsie → endoscopie ~7-14 mai → rezultat ~14-21 mai → plan terapeutic ~sfârșitul lunii mai.\n"
        "Dacă alege combinare → câștigi paralel câteva zile.\n"
        "Aceste 1-2 săptămâni NU sunt o pierdere — sunt timpul minim pentru un diagnostic CORECT, "
        "fără care planul de tratament n-ar fi sigur.",
        kind="info",
    )

    # ============ CE FACEM CONCRET ============
    add_h1(doc, "6. Ce facem concret până luni 4 mai")

    add_h2(doc, "Acțiuni rămase pentru pregătirea consultului")

    add_bullet(doc, "analize sânge la Bioclinica Arad (CEA + CA 19-9 + HbA1c) — programat", bold_prefix="Mâine 29 aprilie:")
    add_bullet(doc, "consult cardiologic ambulator → solicită ECG + ECO + scrisoare medicală cu FEVS și aviz perioperator", bold_prefix="Joi 30 aprilie:")
    add_bullet(doc, "asamblare dosar fizic — printează: buletinul Bioclinica al biopsiei + acest document explicativ + scrisoarea cardio + analizele 29.04", bold_prefix="30.04 - 03.05:")
    add_bullet(doc, "consult OncoHelp Timișoara cu Dr. Anater + comisie oncologică + opinie chirurg eso", bold_prefix="Luni 4 mai:")

    add_h2(doc, "Întrebarea cheie pentru Dr. Anater (luni 4 mai)")

    add_callout(
        doc,
        "Întrebare prioritate 1:",
        q("Domnule doctor, având în vedere că biopsia 17.04 e inconcluzivă (sugestiv carcinomatos, "
          "dar limitat de numărul mic de celule atipice), care e următorul pas — IHC pe blocul T26H06044 "
          "la Bioclinica (cu ce panel de markeri?), rebiopsie țintită endoscopică (eventual EUS), "
          "sau combinare? Și cât timp adaugă fiecare opțiune în plan-ul terapeutic?"),
        kind="urgent",
    )

    add_para(
        doc,
        "Lista completă de întrebări pentru consultul oncologic e în "
        "TODO.md secțiunea pentru viitorul oncolog digestiv (reprioritizată azi cu IHC ca PRIORITATE 1).",
    )

    # ============ MORAL ============
    add_h1(doc, "7. Pe scurt, pentru moral")

    add_callout(
        doc,
        "Lucruri obiective FAVORABILE",
        "• CT-ul a arătat M0 PROBABIL (fără metastaze la distanță vizibile) — neschimbat.\n"
        "• Funcția renală e bună (creatinină 0.83), suportă chemoterapie și contrast.\n"
        "• Nu are scădere ponderală — semn că starea generală nutrițională e ok.\n"
        "• Nu are disfagie progresivă clasică — încă se alimentează rezonabil.\n"
        "• Cardiologic e stabil (post-stent 2012, controlat farmacologic).\n"
        "• Are o echipă medicală formată și un consult la o clinică oncologică reputabilă programat.",
        kind="ok",
    )

    add_callout(
        doc,
        "Lucruri obiective DE URMĂRIT",
        "• Suspiciunea oncologică nu s-a infirmat — biopsia a lăsat-o deschisă.\n"
        "• Ascita observată la CT necesită clarificare (Anater zice probabil cardiacă, dar de exclus cu certitudine).\n"
        "• Diagnosticul de certitudine va veni la 1-3 săptămâni după consult — nu mai devreme.\n"
        "• Plan-ul terapeutic real (FLOT vs altă schemă vs paliativ) se decide după IHC sau rebiopsie.",
        kind="warn",
    )

    add_callout(
        doc,
        "Mesaj final",
        "Nu tragem concluzii pe biopsia singură. Așteptăm consultul de luni cu Dr. Anater "
        "(care a fost foarte echilibrat în mailuri — separând clar oncologic de cardiac) "
        "și comisia tumor board care va decide pasul următor cu echipa completă.",
        kind="info",
    )

    # ============ FOOTER ============
    doc.add_page_break()
    add_h2(doc, "Surse și referințe")

    add_para(
        doc,
        "Acest document a fost generat de Claude Opus 4.7 (Claude Code) pe 28 aprilie 2026, pe baza:",
    )
    add_bullet(doc, "Buletinul Bioclinica nr. 26417A0362 / cod T26H06044 (semnat Dr. Glăja Romanița 27.04)")
    add_bullet(doc, "JSON canonic: Dosar_Medical/2026-04-17_biopsie_esofagiana_histopatologic.json")
    add_bullet(doc, "Interpretare clinică completă: CONTEXT_MEDICAL.md §7.4")
    add_bullet(doc, "Ghidul ESMO pentru cancer eso-gastric (referință generală pentru protocoalele FLOT/CROSS)")
    add_bullet(doc, "Standardele NCCN privind imunohistochimia în cancere eso-gastrice")

    add_h2(doc, "Atenționare obligatorie")
    add_callout(
        doc,
        "Acest document NU înlocuiește consultul medical",
        "Conform Regulii 17 din regulamentul dosarului, această explicație are scop educativ pentru familie. "
        "Diagnosticul, tratamentul și deciziile clinice aparțin EXCLUSIV medicilor curanți "
        "(Dr. Anater - oncolog OncoHelp, Dr. Glăja - anatomopatolog Bioclinica, comisia tumor board OncoHelp).",
        kind="warn",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Document generat: 28 aprilie 2026 · Claude Opus 4.7 (1M context) · Pentru Roland Petrila")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK

    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"OK saved: {OUTPUT} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
