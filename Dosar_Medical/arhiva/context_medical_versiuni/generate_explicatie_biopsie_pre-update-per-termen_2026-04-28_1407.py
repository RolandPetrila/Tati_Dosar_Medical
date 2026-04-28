#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generator DOCX — Explicatie rezultat biopsie 17.04.2026 pentru familie (Roland).

R17 (marcaje certitudine) + R18 (sync DASHBOARD) + R20 (mod de lucru).
Document destinat: Roland Petrila — limbaj accesibil, cu analogii.

Rulare: python scripts/generate_explicatie_biopsie.py
Output: Documente_Informative/EXPLICATIE_REZULTAT_BIOPSIE_2026-04-28.docx
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "Documente_Informative" / "EXPLICATIE_REZULTAT_BIOPSIE_2026-04-28.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREEN = RGBColor(0x38, 0x76, 0x1D)
ORANGE = RGBColor(0xC0, 0x55, 0x04)
RED = RGBColor(0xC0, 0x00, 0x00)
GRAY_DARK = RGBColor(0x40, 0x40, 0x40)
TEXT = RGBColor(0x20, 0x20, 0x20)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in kwargs:
            border = OxmlElement("w:" + edge)
            border.set(qn("w:val"), kwargs[edge].get("val", "single"))
            border.set(qn("w:sz"), str(kwargs[edge].get("sz", 4)))
            border.set(qn("w:color"), kwargs[edge].get("color", "auto"))
            tc_borders.append(border)
    tc_pr.append(tc_borders)


def add_h1(doc, text, color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_h2(doc, text, color=BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_h3(doc, text, color=GRAY_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = color
    r.font.name = "Calibri"


def add_para(doc, text, bold=False, italic=False, size=11, color=TEXT, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return p


def add_callout(doc, title, body, kind="info"):
    palette = {
        "info": ("D9E2F3", "2E75B6"),
        "warn": ("FCE4D6", "C05504"),
        "ok": ("E2EFDA", "38761D"),
        "urgent": ("FADBD8", "C00000"),
        "analogy": ("F2F2F2", "595959"),
    }
    fill, border = palette.get(kind, palette["info"])
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, fill)
    set_cell_border(
        cell,
        top={"sz": 8, "color": border},
        left={"sz": 24, "color": border},
        bottom={"sz": 8, "color": border},
        right={"sz": 8, "color": border},
    )
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    r1 = p_title.add_run(title)
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor.from_string(border)
    r1.font.name = "Calibri"
    p_body = cell.add_paragraph()
    p_body.paragraph_format.space_after = Pt(4)
    r2 = p_body.add_run(body)
    r2.font.size = Pt(11)
    r2.font.color.rgb = TEXT
    r2.font.name = "Calibri"
    doc.add_paragraph()


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.font.bold = True
        r1.font.size = Pt(11)
        r1.font.name = "Calibri"
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
        r.font.name = "Calibri"


def add_table_2col(doc, header_left, header_right, rows, color_zebra="F2F2F2"):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.autofit = False
    hdr = table.rows[0]
    for idx, txt in enumerate([header_left, header_right]):
        cell = hdr.cells[idx]
        set_cell_bg(cell, "1F4E79")
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(11)
        r.font.name = "Calibri"
    for ri, (left, right) in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        if ri % 2 == 0:
            for c in row_cells:
                set_cell_bg(c, color_zebra)
        for c, txt in zip(row_cells, [left, right]):
            c.text = ""
            p = c.paragraphs[0]
            r = p.add_run(txt)
            r.font.size = Pt(10.5)
            r.font.name = "Calibri"
            r.font.color.rgb = TEXT


# Use only ASCII quotes inside Python string literals; smart quotes via unicode escapes.
LDQ = "„"  # double low-9 quote ("low double quote")
RDQ = "”"  # right double quotation mark
EM = "—"   # em dash


def q(text):
    """Wrap text with Romanian-style smart quotes."""
    return LDQ + text + RDQ


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ============ COVER ============
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    r = p.add_run("EXPLICATIE\nREZULTAT BIOPSIE\n17 aprilie 2026")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Document pentru Roland " + EM + " explicatie simpla, fara jargon medical")
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("Pacient: PETRILA VIOREL-MIHAI (66 ani)")
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = TEXT
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Buletin Bioclinica nr. 26417A0362 / cod T26H06044\n"
        "Semnat: Dr. Glaja Romanita (medic primar anatomopatolog)\n"
        "Generat: 28 aprilie 2026"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run(
        "Document generat: 28 aprilie 2026\n"
        "Pentru consultul oncologic de luni 4 mai 2026 " + EM + " OncoHelp Timisoara"
    )
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK
    r.font.name = "Calibri"

    doc.add_page_break()

    # ============ TL;DR ============
    add_h1(doc, "În două rânduri")

    add_callout(
        doc,
        "Concluzia raportului în limbaj clar:",
        "Biopsia NU spune " + q("cancer confirmat") + " " + EM
        + " dar NU spune nici " + q("nu e cancer") + ". "
        "Spune " + q("nu pot decide din proba asta — e prea puțin țesut tumoral în ea") + ". "
        "Suspiciunea de cancer rămâne ridicată din CT și endoscopie. "
        "Următorul pas (decis de oncolog luni 4 mai): un test suplimentar pe blocul de țesut "
        "deja prelevat (IHC), sau o nouă biopsie țintită mai în profunzime.",
        kind="info",
    )

    # ============ CE SCRIE EXACT ============
    add_h1(doc, "1. Ce scrie exact în raport")

    add_para(
        doc,
        "Anatomopatologul (Dr. Glăja Romanița, medic primar la Bioclinica) a analizat la microscop "
        "cele 2 fragmente de țesut prelevate de Dr. Noufal la endoscopia din 17 aprilie. "
        "Iată concluzia ei textuală, citată din raport:",
    )

    add_callout(
        doc,
        "Concluzia anatomopatologului (text original):",
        q("Ansamblul histologic, pe secțiuni seriate și în colorația uzuală, "
          "pledează pentru ȚESUT DE GRANULAȚIE pe fond de ULCERAȚIE CRONICĂ, "
          "fiind doar SUGESTIV pentru infiltrat carcinomatos.")
        + "\n\n"
        + q("De corelat cu datele endoscopice/imagistice (diagnostic histologic tumoral mult limitat de "
            "numărul mic al celulelor epiteliale atipice); eventuală evaluare imunohistochimică pentru "
            "diagnostic histologic de certitudine și conduită terapeutică."),
        kind="info",
    )

    add_h2(doc, "Ce înseamnă fiecare bucată — pe rând")

    add_h3(doc, "» " + q("țesut de granulație pe fond de ulcerație cronică"))
    add_para(
        doc,
        "Țesut de granulație = țesut de reparație, ca o crustă vie care apare când o suprafață e rănită "
        "și organismul încearcă să se vindece. Ulcerație cronică = o rană care durează de mult timp.",
    )
    add_para(doc, "Acest tip de țesut apare în multe situații:")
    add_bullet(doc, "ulcer peptic (rană de stomac/esofag fără cancer)")
    add_bullet(doc, "esofagită severă cronică")
    add_bullet(doc, "PESTE / DEASUPRA unei tumori (zona de crustă care acoperă tumora propriu-zisă)")

    add_para(
        doc,
        "Cu alte cuvinte: țesutul de granulație SINGUR nu spune nici cancer, nici nu e cancer. "
        "E doar un strat superficial inflamat.",
    )

    add_h3(doc, "» " + q("doar SUGESTIV pentru infiltrat carcinomatos"))
    add_para(
        doc,
        "În proba prelevată, anatomopatologul a văzut și CÂTEVA celule cu aspect care POATE indica un cancer "
        "(celule epitelioide cu nucleu nucleolat, cu nucleol eozinofil — caracteristici tipice pentru adenocarcinom). "
        "DAR " + EM + " și aici e cuvântul-cheie " + EM + " au fost prea puține celule de tipul ăsta pentru a "
        "putea pune un diagnostic ferm. Sugestiv e un termen anatomopatologic care înseamnă "
        "indicii da, certitudine nu.",
    )

    add_h3(doc, "» De ce nu a fost suficient?")
    add_para(
        doc,
        "Răspunsul e în nota laboratorului: " + q("diagnostic histologic tumoral mult limitat de numărul "
        "mic al celulelor epiteliale atipice") + ".",
    )
    add_para(
        doc,
        "Cu alte cuvinte: cele 2 fragmente prelevate au fost foarte mici (sub-milimetrice — 2 piese de "
        "0,2 / 0,1 / 0,1 cm fiecare) și au conținut în principal stratul superficial (granulație + necroză), "
        "nu stratul tumoral profund. Asta NU e o eroare a Dr. Noufal " + EM + " e foarte greu să prelevezi adânc "
        "dintr-o leziune circumferențială stenozantă fără să riști sângerare sau perforație.",
    )

    add_callout(
        doc,
        "Analogie pentru a înțelege:",
        "E ca și cum ai vrea să afli ce e într-o casă, dar te uiți doar pe geamul de la intrare. "
        "Vezi clar holul (stratul superficial = granulație), poate vezi un colț din salon "
        "(câteva celule atipice), dar nu poți să spui sigur câte camere are casa, ce mobilier e, "
        "cine locuiește acolo. Ai nevoie să intri și să te uiți cu atenție în fiecare cameră.",
        kind="analogy",
    )

    # ============ CE NU ÎNSEAMNĂ ============
    add_h1(doc, "2. Ce NU înseamnă acest rezultat")

    add_callout(
        doc,
        "X NU înseamnă " + q("nu e cancer, e doar o ulcerație"),
        "Suspiciunea de cancer rămâne ridicată — biopsia n-a putut DOVEDI cancerul, "
        "dar n-a putut nici să-l EXCLUDĂ. CT-ul a arătat un proces tumoral T3-T4 infiltrativ, "
        "endoscopia a arătat o stenoză " + q("nedepășibilă") + ". Aceste lucruri NU s-au schimbat după biopsie.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("e confirmat cancer"),
        "Sunt indicii sugestive, dar nu certitudine histologică. "
        "Un diagnostic oficial de cancer cere ca anatomopatologul să spună " + q("cancer confirmat")
        + " — iar acest raport spune doar " + q("sugestiv pentru cancer")
        + ". E o nuanță importantă legal și medical.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("trebuie să luăm tot procesul de la zero"),
        "Probabil oncologul va decide să facă un test suplimentar (IHC) "
        "PE BLOCUL DEJA EXISTENT (T26H06044) — fără reintervenție. Asta poate da răspunsul în 3-7 zile.",
        kind="warn",
    )

    add_callout(
        doc,
        "X NU înseamnă " + q("întârziem tratamentul"),
        "Consultul de luni 4 mai rămâne valid și esențial. "
        "Dr. Anater va decide acolo opțiunea (IHC vs rebiopsie) — iar oricare dintre ele "
        "se mișcă în paralel cu pregătirea dosarului oncologic + opinia chirurgului.",
        kind="warn",
    )

    # ============ CE URMEAZA ============
    add_h1(doc, "3. Ce urmează — 3 opțiuni standard")

    add_para(
        doc,
        "Anatomopatologul recomandă explicit în nota laboratorului: "
        + q("eventuală evaluare imunohistochimică pentru diagnostic histologic de certitudine") + ".",
    )
    add_para(
        doc,
        "Concret, oncologul de luni 4 mai (Dr. Anater) va alege una dintre aceste 3 căi:",
    )

    add_h2(doc, "Opțiunea 1 — IHC pe blocul existent (recomandare anatomopatolog)")

    add_callout(
        doc,
        "Ce e IHC (imunohistochimie)?",
        "Anatomopatologul ia blocul de parafină T26H06044 (deja la Bioclinica Timișoara) "
        "și aplică niște coloranți speciali numiți markeri pe felii noi din același țesut. "
        "Acești markeri se atașează doar de proteinele specifice celulelor canceroase. "
        "Dacă celulele atipice se colorează cu markerii potriviți → confirmare cancer + tip exact.",
        kind="ok",
    )

    add_para(doc, "Markerii tipici care se cer pentru o leziune eso-gastrică:")
    add_bullet(doc, "Pan-CK / CK7 / CK20 — confirmă originea epitelială a celulelor atipice")
    add_bullet(doc, "CDX-2 — pozitiv în adenocarcinoame de origine intestinală/eso-gastrică")
    add_bullet(doc, "p53 — markerul mutațiilor oncogene clasice")
    add_bullet(doc, "Dacă se confirmă cancer: HER2, PD-L1, MSI — ghidează decizia terapeutică (FLOT, imunoterapie, trastuzumab)")

    add_table_2col(
        doc,
        "AVANTAJE",
        "DEZAVANTAJE",
        [
            ("Rapid (3-7 zile rezultat)", "Dacă atipia e prea redusă, IHC poate fi tot inconcluziv"),
            ("Fără reintervenție pentru pacient", "Se face pe puținele celule atipice deja existente"),
            ("Recomandat explicit de Dr. Glăja", "Necesită trimitere de la oncolog la Bioclinica"),
            ("Cost moderat (~300-600 lei)", "Decizia depinde de Dr. Anater"),
        ],
    )

    add_h2(doc, "Opțiunea 2 — Rebiopsie țintită endoscopică")

    add_para(
        doc,
        "O nouă endoscopie cu prelevare de fragmente mai numeroase și mai profunde, eventual sub ghidaj "
        "EUS (ecoendoscopie — un endoscop cu o sondă de ecografie integrată care poate vedea stratul tumoral "
        "profund înainte de a preleva, ca un GPS pentru biopsie).",
    )

    add_table_2col(
        doc,
        "AVANTAJE",
        "DEZAVANTAJE",
        [
            ("Randament diagnostic superior", "Necesită reintervenție (anestezie, programare)"),
            ("Poate preleva din stratul tumoral profund", "Adaugă 1-3 săptămâni la timeline"),
            ("EUS vede stadiul T cu precizie", "Cost suplimentar (poate fi în CAS)"),
            ("Standard NCCN/ESMO la biopsie inconcluzivă", "Risc minor de sângerare/perforație"),
        ],
    )

    add_h2(doc, "Opțiunea 3 — Combinare IHC + rebiopsie în paralel")

    add_para(
        doc,
        "Cea mai eficientă în timp dacă oncologul vrea răspuns rapid + plan B în caz de IHC inconcluziv. "
        "Decizie tipică în comisiile tumor board când rezultatul histologic ține în loc planul terapeutic.",
    )

    # ============ DE CE E IMPORTANT ============
    add_h1(doc, "4. De ce e important să avem răspunsul în câteva săptămâni")

    add_para(
        doc,
        "Diagnosticul histologic de certitudine e poarta de intrare în orice protocol oncologic. Fără el:",
    )
    add_bullet(doc, "Nu se poate iniția chemoterapie (FLOT) — riscul de a trata o leziune non-tumorală cu medicamente toxice e inacceptabil")
    add_bullet(doc, "Nu se poate planifica chirurgia oncologică")
    add_bullet(doc, "Nu se pot accesa terapii țintite (HER2, PD-L1) — care depind de markeri moleculari obținuți doar prin IHC")

    add_callout(
        doc,
        "În cifre concrete (timeline):",
        "Dacă Dr. Anater alege IHC luni 4.05 → rezultat ~10-12 mai → comisie + plan terapeutic ~mijlocul lunii mai.\n"
        "Dacă alege rebiopsie → endoscopie ~7-14 mai → rezultat ~14-21 mai → plan terapeutic ~sfârșitul lunii mai.\n"
        "Dacă alege combinare → câștigi paralel câteva zile.\n"
        "Aceste 1-2 săptămâni NU sunt o pierdere — sunt timpul minim pentru un diagnostic CORECT, "
        "fără care planul de tratament n-ar fi sigur.",
        kind="info",
    )

    # ============ CE FACEM CONCRET ============
    add_h1(doc, "5. Ce facem concret până luni 4 mai")

    add_h2(doc, "Acțiuni rămase pentru pregătirea consultului")

    add_bullet(doc, "analize sânge la Bioclinica Arad (CEA + CA 19-9 + HbA1c) — programat", bold_prefix="Mâine 29 aprilie:")
    add_bullet(doc, "consult cardiologic ambulator → solicită ECG + ECO + scrisoare medicală cu FEVS și aviz perioperator", bold_prefix="Joi 30 aprilie:")
    add_bullet(doc, "asamblare dosar fizic — printează: buletinul Bioclinica al biopsiei + acest document explicativ + scrisoarea cardio + analizele 29.04", bold_prefix="30.04 - 03.05:")
    add_bullet(doc, "consult OncoHelp Timișoara cu Dr. Anater + comisie oncologică + opinie chirurg eso", bold_prefix="Luni 4 mai:")

    add_h2(doc, "Întrebarea cheie pentru Dr. Anater (luni 4 mai)")

    add_callout(
        doc,
        "Întrebare prioritate 1:",
        q("Domnule doctor, având în vedere că biopsia 17.04 e inconcluzivă (sugestiv carcinomatos, "
          "dar limitat de numărul mic de celule atipice), care e următorul pas — IHC pe blocul T26H06044 "
          "la Bioclinica (cu ce panel de markeri?), rebiopsie țintită endoscopică (eventual EUS), "
          "sau combinare? Și cât timp adaugă fiecare opțiune în plan-ul terapeutic?"),
        kind="urgent",
    )

    add_para(
        doc,
        "Lista completă de întrebări pentru consultul oncologic e în "
        "TODO.md secțiunea pentru viitorul oncolog digestiv (reprioritizată azi cu IHC ca PRIORITATE 1).",
    )

    # ============ MORAL ============
    add_h1(doc, "6. Pe scurt, pentru moral")

    add_callout(
        doc,
        "Lucruri obiective FAVORABILE",
        "• CT-ul a arătat M0 PROBABIL (fără metastaze la distanță vizibile) — neschimbat.\n"
        "• Funcția renală e bună (creatinină 0.83), suportă chemoterapie și contrast.\n"
        "• Nu are scădere ponderală — semn că starea generală nutrițională e ok.\n"
        "• Nu are disfagie progresivă clasică — încă se alimentează rezonabil.\n"
        "• Cardiologic e stabil (post-stent 2012, controlat farmacologic).\n"
        "• Are o echipă medicală formată și un consult la o clinică oncologică reputabilă programat.",
        kind="ok",
    )

    add_callout(
        doc,
        "Lucruri obiective DE URMĂRIT",
        "• Suspiciunea oncologică nu s-a infirmat — biopsia a lăsat-o deschisă.\n"
        "• Ascita observată la CT necesită clarificare (Anater zice probabil cardiacă, dar de exclus cu certitudine).\n"
        "• Diagnosticul de certitudine va veni la 1-3 săptămâni după consult — nu mai devreme.\n"
        "• Plan-ul terapeutic real (FLOT vs altă schemă vs paliativ) se decide după IHC sau rebiopsie.",
        kind="warn",
    )

    add_callout(
        doc,
        "Mesaj final",
        "Nu tragem concluzii pe biopsia singură. Așteptăm consultul de luni cu Dr. Anater "
        "(care a fost foarte echilibrat în mailuri — separând clar oncologic de cardiac) "
        "și comisia tumor board care va decide pasul următor cu echipa completă.",
        kind="info",
    )

    # ============ FOOTER ============
    doc.add_page_break()
    add_h2(doc, "Surse și referințe")

    add_para(
        doc,
        "Acest document a fost generat de Claude Opus 4.7 (Claude Code) pe 28 aprilie 2026, pe baza:",
    )
    add_bullet(doc, "Buletinul Bioclinica nr. 26417A0362 / cod T26H06044 (semnat Dr. Glăja Romanița 27.04)")
    add_bullet(doc, "JSON canonic: Dosar_Medical/2026-04-17_biopsie_esofagiana_histopatologic.json")
    add_bullet(doc, "Interpretare clinică completă: CONTEXT_MEDICAL.md §7.4")
    add_bullet(doc, "Ghidul ESMO pentru cancer eso-gastric (referință generală pentru protocoalele FLOT/CROSS)")
    add_bullet(doc, "Standardele NCCN privind imunohistochimia în cancere eso-gastrice")

    add_h2(doc, "Atenționare obligatorie")
    add_callout(
        doc,
        "Acest document NU înlocuiește consultul medical",
        "Conform Regulii 17 din regulamentul dosarului, această explicație are scop educativ pentru familie. "
        "Diagnosticul, tratamentul și deciziile clinice aparțin EXCLUSIV medicilor curanți "
        "(Dr. Anater - oncolog OncoHelp, Dr. Glăja - anatomopatolog Bioclinica, comisia tumor board OncoHelp).",
        kind="warn",
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("Document generat: 28 aprilie 2026 · Claude Opus 4.7 (1M context) · Pentru Roland Petrila")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = GRAY_DARK

    doc.save(str(OUTPUT))
    size_kb = OUTPUT.stat().st_size / 1024
    print(f"OK saved: {OUTPUT} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    main()
