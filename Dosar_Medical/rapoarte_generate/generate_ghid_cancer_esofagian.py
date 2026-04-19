# -*- coding: utf-8 -*-
"""
Generator DOCX: Ghid complet cancer esofagian pentru familia Petrila.
Compilat din 4 agenti de cercetare paralela (tratament, trial-uri, centre, suport practic).
Sursa datelor: studii pivot NCCN/ESMO + clinicaltrials.gov + SEER + CNAS + surse oficiale.
Ruleaza: python generate_ghid_cancer_esofagian.py
Output: 2026-04-19_ghid_cancer_esofagian_complet.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN = RGBColor(0x0B, 0x6E, 0x4F)
BLUE = RGBColor(0x1F, 0x4E, 0x79)
GRAY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x00, 0x00, 0x00)
LIGHT_BLUE = RGBColor(0x3B, 0x82, 0xF6)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc, text="", style=None, bold=False, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold)
    return p


def add_heading_custom(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_marker(paragraph, marker_type):
    colors = {
        "CERT": GREEN,
        "PROBABIL": ORANGE,
        "INCERT": ORANGE,
        "NEGASIT": GRAY,
        "RECOMANDAT": BLUE,
        "URGENT": RED,
    }
    color = colors.get(marker_type, BLACK)
    run = paragraph.add_run(f"[{marker_type}] ")
    run.bold = True
    run.font.color.rgb = color


def add_bullet(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            marker, text = item
            add_marker(p, marker)
            p.add_run(text)
        else:
            p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def add_table(doc, headers, rows, header_color="1F4E79", first_col_bold=True, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if c_idx == 0 and first_col_bold:
                run.bold = True

    if col_widths:
        for row in tbl.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = col_widths[i]
    return tbl


def add_callout(doc, title, body, color=ORANGE):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    bg = "FFF4E6" if color == ORANGE else ("FEECEC" if color == RED else ("E8F5E9" if color == GREEN else "E3F2FD"))
    set_cell_shading(cell, bg)
    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.add_run(body).font.size = Pt(10)


def add_source(doc, text):
    """Adauga sursa cu font mic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Setare margini
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ============================================================
    # COPERTA
    # ============================================================
    title = doc.add_heading("GHID COMPLET", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Cancer esofagian — de la suspiciune la tratament")
    run.bold = True
    run.font.size = Pt(16)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Document informativ pentru familie")
    run.italic = True
    run.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Pacient: Petrilă Viorel-Mihai (66 ani, Nădlac)\n"
        "Data raport: 19 aprilie 2026\n"
        "Context: suspiciune proces proliferativ esofagian (endoscopie 17.04.2026)\n"
        "Biopsie în așteptare · CT programat 20.04.2026 ora 17:00 (Genesis Micălaca)\n"
        "Autor: Claude Opus 4.7 (1M context) — compilat din 4 sesiuni de cercetare paralele"
    ).font.size = Pt(10)

    doc.add_paragraph()

    add_callout(
        doc,
        "⚠ Atenționare esențială",
        "Acest ghid NU înlocuiește consultul medical. Este un document de orientare pentru "
        "familie, destinat să clarifice vocabular medical și să enumere opțiunile disponibile. "
        "TOATE deciziile de tratament aparțin echipei medicale curante (oncolog, chirurg, "
        "radioterapeut). La momentul finalizării acestui document, biopsia și CT-ul tatălui sunt "
        "încă în desfășurare — tot conținutul este condițional pe rezultate.",
        color=RED,
    )
    doc.add_page_break()

    # ============================================================
    # CUPRINS MANUAL
    # ============================================================
    add_heading_custom(doc, "Cuprins", level=1)
    toc_items = [
        "I. Cum se citește acest document",
        "II. Rezumat în 10 puncte esențiale",
        "III. Ce este „procesul proliferativ esofagian”",
        "IV. Statistici — câți au cancer și câți nu",
        "V. Sistemul TNM — cele 3 direcții de extindere",
        "VI. Ce se întâmplă la fiecare nivel de extindere",
        "VII. Când vine rezultatul CT",
        "VIII. Tratamente pentru leziune BENIGNĂ",
        "IX. Tratamente pentru leziune PRECANCEROASĂ",
        "X. Tratamente pentru CANCER — pe stadii 1, 2, 3, 4",
        "XI. Operația esofagectomie — cât de grea este",
        "XII. Imunoterapia — revoluția 2020-2025",
        "XIII. Trial-uri clinice active cu centre în România",
        "XIV. Centre oncologice — România",
        "XV. Centre oncologice — UE (accesibile prin S2/E112)",
        "XVI. Second opinion internațional (fără deplasare)",
        "XVII. Programul Național de Oncologie CNAS",
        "XVIII. Drepturi pacient oncologic — indemnizație, gradul de handicap",
        "XIX. Nutriție, suport psihologic, paliație",
        "XX. Logistică tratament în București — cazare gratuită",
        "XXI. Întrebări concrete de pus fiecărui specialist",
        "XXII. Plan de acțiune — săptămâna 1, 2, 3 și următoarele",
        "XXIII. Ce NU am găsit (transparență)",
        "XXIV. Surse citate",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.add_run("• " + item).font.size = Pt(10)

    doc.add_page_break()

    # ============================================================
    # I. CUM SE CITEȘTE
    # ============================================================
    add_heading_custom(doc, "I. Cum se citește acest document", level=1, color=BLUE)
    doc.add_paragraph(
        "Documentul e scris pentru cititor fără pregătire medicală. Fiecare afirmație factuală "
        "are un marcaj de certitudine înaintea ei (conform Regulii 17 din regulamentul proiectului):"
    )
    add_bullet(doc, [
        ("CERT", " — confirmat din sursă primară (ghid ESMO/NCCN, studiu peer-reviewed, RCP oficial, document CNAS). Cu URL citat."),
        ("PROBABIL", " — susținut de literatura medicală, dar fără dovada specifică pentru cazul pacientului."),
        ("INCERT", " — date conflictuale sau parțiale; necesită validare de la medic."),
        ("NEGASIT", " — căutat în surse oficiale, nu a fost găsit; indică unde s-a căutat."),
        ("RECOMANDAT", " — opinie de organizare practică, nu afirmație medicală."),
        ("URGENT", " — acțiune imediată sau alertă critică."),
    ])
    doc.add_paragraph()
    doc.add_paragraph(
        "Culori semantice: ROȘU = alertă/urgență/rezultat grav. VERDE = confirmat, siguranță, veste bună. "
        "ORANJ = incert sau atenționare. ALBASTRU = titluri/organizare."
    )
    doc.add_page_break()

    # ============================================================
    # II. REZUMAT
    # ============================================================
    add_heading_custom(doc, "II. Rezumat în 10 puncte esențiale", level=1, color=BLUE)
    rezumat = [
        "Termenul „proces proliferativ esofagian” = descriere a unei leziuni vizibile la endoscopie. NU înseamnă automat cancer. Biopsia va clarifica.",
        "Din leziunile esofagiene vizibile, aproximativ 40-50% sunt benigne, 15-20% precanceroase, 35-45% maligne (valoare orientativă).",
        "Rezultatul CT vine în 2-3 zile lucrătoare (Genesis/Affidea/SANADOR). Biopsia histopatologică: 7-14 zile lucrătoare (monitor automat este activ).",
        "Cancerul esofagian se „extinde” pe 3 direcții simultane (T = adâncime, N = ganglioni, M = metastaze). Fiecare are consecințe diferite pentru tratament.",
        "Pentru stadii timpurii (T1a) există tratament ENDOSCOPIC fără operație — vindecare 90-95%. Pentru stadii 2-3 se face chimio-radio ÎNAINTE + operație + imunoterapie DUPĂ.",
        "Protocolul FLOT (chimio perioperatorie) este mai bun decât CROSS (chimio-radio) pentru adenocarcinom (studiul ESOPEC 2024, confirmat ASCO 2024).",
        "Imunoterapia (pembrolizumab, nivolumab) a schimbat prognosticul în ultimii 4 ani — adjuvant după operație (CheckMate-577) și în stadiul 4 (KEYNOTE-590, CheckMate-648). Decontată în România prin Programul Național de Oncologie.",
        "Pentru pacient cu stent coronarian, tehnica chirurgicală MINIMAL-INVAZIVĂ ROBOTICĂ (RAMIE) reduce complicațiile cardiace cu ~53% și pulmonare cu ~46% față de operația clasică deschisă.",
        "Centre de excelență accesibile: SCJU Arad (local, Prof. Totolici, enumeră explicit chirurgie esofag), Ponderas București (robot Da Vinci, acreditare SRC), ICF Chirurgie IV (Prof. Bacalbașa, București), UE: Charité Berlin, Heidelberg, Gustave Roussy Paris.",
        "Procedura S2 CNAS: aprobare în 5 zile lucrătoare (legal) pentru tratament în UE când nu se poate face în termen rezonabil în România. Imunoterapie + chirurgie oncologică = DECONTATE prin CAS dacă dosarul e complet.",
    ]
    add_numbered(doc, rezumat)
    doc.add_page_break()

    # ============================================================
    # III. CE ESTE PROCESUL PROLIFERATIV
    # ============================================================
    add_heading_custom(doc, "III. Ce este „procesul proliferativ esofagian”", level=1, color=BLUE)

    doc.add_paragraph(
        "Esofagul este tubul muscular lung de ~25 cm care transportă mâncarea din gură în stomac. "
        "Peretele esofagului are 4 straturi concentrice (ca o ceapă cu 4 foi):"
    )
    rows = [
        ["Mucoasa", "Stratul cel mai interior, atinge mâncarea. Aici pornesc majoritatea cancerelor."],
        ["Submucoasa", "Strat cu vase de sânge și vase limfatice. Aici sunt stațiile de plecare pentru metastaze."],
        ["Muscularis propria", "Stratul muscular care împinge mâncarea prin undă peristaltică."],
        ["Adventice", "Foița subțire de țesut conjunctiv care acoperă exteriorul esofagului."],
    ]
    add_table(doc, ["Strat", "Descriere"], rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "„Proliferativ” înseamnă că celulele s-au multiplicat anormal, formând o leziune vizibilă "
        "endoscopic. Termenul este DESCRIPTIV — medicul a văzut ceva anormal, a luat biopsie pentru "
        "a afla ce anume este. Biopsia (examenul histopatologic) este singura metodă care decide "
        "cu certitudine dacă celulele sunt benigne, precanceroase sau maligne."
    )

    doc.add_paragraph()
    add_heading_custom(doc, "Tipuri de leziuni esofagiene posibile (rezultate biopsie)", level=2)

    rows = [
        ["BENIGN", "Polip, papilom, leiomiom (tumoră musculară), chist, esofagită, strictură peptică.", "Tratament minor sau urmărire."],
        ["PRECANCEROS", "Displazie de grad scăzut/înalt (celule modificate dar nu canceroase); esofag Barrett.", "Rezecție endoscopică preventivă."],
        ["MALIGN — carcinom scuamocelular (ESCC)", "Pornește din stratul superficial. Mai frecvent în treimea superioară/medie. Asociat cu fumat + alcool.", "Radio-sensibil, răspunde bine la imunoterapie."],
        ["MALIGN — adenocarcinom (EAC)", "Pornește din glandele esofagului. Mai frecvent în treimea inferioară/joncțiune gastro-esofagiană (GEJ). Asociat cu reflux cronic.", "Tratament diferit; se evaluează HER2."],
    ]
    add_table(doc, ["Categorie", "Ce înseamnă", "Implicație"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Pentru context — în România",
        "Carcinomul scuamocelular (ESCC) este mai frecvent decât adenocarcinomul în populația din "
        "Europa Centrală și de Est (inclusiv România), comparativ cu SUA/Europa de Vest unde "
        "adenocarcinomul este mai frecvent. Ambele tipuri sunt tratabile; diferența contează pentru "
        "protocolul specific (CROSS preferat ESCC, FLOT preferat adenocarcinom GEJ).",
        color=BLUE,
    )
    doc.add_page_break()

    # ============================================================
    # IV. STATISTICI
    # ============================================================
    add_heading_custom(doc, "IV. Statistici — câți au cancer și câți nu", level=1, color=BLUE)

    p = doc.add_paragraph()
    add_marker(p, "PROBABIL")
    p.add_run(
        "La o populație de pacienți cu LEZIUNE esofagiană vizibilă la endoscopie care se face biopsie, "
        "distribuția aproximativă este:"
    )

    rows = [
        ["Benign", "40-50%", "Polipi, papiloame, leiomioame, esofagite, stricturi benigne."],
        ["Precanceros", "15-20%", "Displazie (celule modificate dar nu canceroase încă)."],
        ["Malign", "35-45%", "Carcinom scuamocelular sau adenocarcinom."],
    ]
    add_table(doc, ["Categorie", "Probabilitate", "Exemple"], rows)

    doc.add_paragraph()
    add_heading_custom(doc, "Supraviețuire cancer esofagian pe stadii (dacă biopsia confirmă malign)", level=2)

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: SEER Cancer Stat Facts — Esophageal Cancer (date agregate 2015-2021), accesat 19.04.2026.")

    rows = [
        ["Localizat (limitat la esofag)", "Stadiu 0-I-IIA", "48%"],
        ["Regional (ganglioni sau invazie locală)", "Stadiu IIB-III", "28%"],
        ["Distant (metastatic)", "Stadiu IV", "6%"],
        ["OVERALL (toate stadiile combinate)", "—", "22%"],
    ]
    add_table(doc, ["Stadiu SEER", "Echivalent TNM", "Supraviețuire 5 ani"], rows)

    doc.add_paragraph()
    add_heading_custom(doc, "Detaliere pe sub-stadii (aproximativă)", level=2)

    p = doc.add_paragraph()
    add_marker(p, "PROBABIL")
    p.add_run("Cifre din NCCN + UpToDate + literatura OAE — estimative pre-imunoterapie.")

    rows = [
        ["Stadiu 0 (Tis N0 M0)", "Carcinom in situ", ">90%"],
        ["Stadiu IA (T1a N0 M0)", "Mucosal", "70-85%"],
        ["Stadiu IB (T1b N0 M0)", "Submucosal", "55-70%"],
        ["Stadiu IIA (T2 N0 M0)", "Muscularis propria", "40-55%"],
        ["Stadiu IIB (T3 N0 M0 / T1-2 N1 M0)", "Adventice sau N1 minim", "25-40%"],
        ["Stadiu IIIA (T1-2 N2 / T3 N1)", "Ganglionar extins", "20-30%"],
        ["Stadiu IIIB (T3 N2 / T4a N0-1)", "Local avansat", "15-20%"],
        ["Stadiu IVA (T4b / N3)", "Nerezecabil local", "5-10%"],
        ["Stadiu IVB (Any T Any N M1)", "Metastatic", "<5% (14-18 luni median cu imunoterapie)"],
    ]
    add_table(doc, ["Stadiu AJCC 8", "Descriere", "Supraviețuire 5 ani"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Important — cifrele se îmbunătățesc",
        "Supraviețuirea a crescut în 2020-2024 cu ~3-5 puncte procentuale la stadiile III-IV "
        "datorită: (1) CheckMate-577 (nivolumab adjuvant post-CROSS, dublare DFS), "
        "(2) KEYNOTE-590 (pembrolizumab + chimio prima linie), (3) ESOPEC/FLOT "
        "(protocol chimioterapeutic superior). Pacientul de azi are șanse mai bune decât statisticile "
        "vechi.",
        color=GREEN,
    )
    doc.add_page_break()

    # ============================================================
    # V. TNM
    # ============================================================
    add_heading_custom(doc, "V. Sistemul TNM — cele 3 direcții de extindere", level=1, color=BLUE)

    doc.add_paragraph(
        "Cancerul esofagian se extinde SIMULTAN pe 3 direcții, evaluate separat:"
    )

    add_heading_custom(doc, "T = cât de adânc a pătruns în peretele esofagului", level=2)
    rows = [
        ["Tis", "Carcinom in situ — celule doar la suprafață", "Tratament endoscopic, vindecare >95%"],
        ["T1a", "Invazie lamina propria / muscularis mucosae (mucoasă)", "ESD endoscopic, vindecare 90-95%"],
        ["T1b", "Invazie submucoasă", "ESD sau operație; risc ganglioni 15-25%"],
        ["T2", "Invazie muscularis propria", "OPERAȚIE OBLIGATORIE"],
        ["T3", "Invazie adventice", "Chimio-radio ÎNAINTE + operație"],
        ["T4a", "Invazie pleură / pericard / diafragm / peritoneu", "Operație dificilă dar posibilă"],
        ["T4b", "Invazie aortă / trahee / coloană", "NEOPERABIL; chimio-radio paliativ"],
    ]
    add_table(doc, ["Stadiu T", "Ce înseamnă", "Consecință"], rows)

    add_heading_custom(doc, "N = câți ganglioni limfatici regionali sunt afectați", level=2)
    doc.add_paragraph(
        "Ganglionii limfatici sunt „stațiile de filtrare” ale limfei. Cancerul se propagă prin ei "
        "ca un tren prin stații."
    )
    rows = [
        ["N0", "0 ganglioni afectați", "Cel mai bun prognostic"],
        ["N1", "1-2 ganglioni afectați", "Reduce supraviețuirea cu ~15%"],
        ["N2", "3-6 ganglioni afectați", "Reduce cu ~30%"],
        ["N3", "≥7 ganglioni afectați", "Reduce cu ~50%"],
    ]
    add_table(doc, ["Stadiu N", "Ganglioni pozitivi", "Impact"], rows)

    add_heading_custom(doc, "M = există metastaze la distanță?", level=2)
    rows = [
        ["M0", "Fără metastaze la distanță", "Tratament curativ posibil"],
        ["M1", "Metastaze la distanță", "Tratament paliativ (prelungire viață + calitate)"],
    ]
    add_table(doc, ["Stadiu M", "Ce înseamnă", "Consecință"], rows)

    doc.add_paragraph()
    add_heading_custom(doc, "Unde merge cancerul esofagian când metastazează", level=2)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Journal of Thoracic Disease 2024 (Ai et al.), date SEER.")
    rows = [
        ["Ficat", "15.6%", "Cel mai frecvent — ficatul filtrează sânge din zona digestivă"],
        ["Plămâni", "9.7%", "Al doilea — prin vase de sânge"],
        ["Oase", "7.7%", "Al treilea — coloana, bazinul, coaste"],
        ["Creier", "2-3%", "Mai rar, dar grav"],
        ["Suprarenale", "2-3%", "Descoperite de obicei la CT"],
    ]
    add_table(doc, ["Loc metastază", "Frecvență", "Notă"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Simptome care pot semnala metastaze",
        "• Ficat → icter (piele galbenă), durere hipocondru drept, oboseală, scădere apetit\n"
        "• Plămâni → tuse persistentă, dispnee, durere toracică\n"
        "• Oase → dureri osoase persistente localizate, fracturi spontane\n"
        "• Creier → cefalee matinală, vărsături fără greață, confuzie, convulsii\n\n"
        "CT-ul de pe 20.04 acoperă torace + abdomen + pelvis — verifică plămâni, ficat, ganglioni, "
        "suprarenale, rinichi. Dacă e nevoie de verificare suplimentară: scintigrafie osoasă, RMN "
        "cerebral, PET-CT.",
        color=ORANGE,
    )
    doc.add_page_break()

    return doc


def build_document_part2(doc):
    # ============================================================
    # VI. CE SE INTAMPLA LA FIECARE NIVEL DE EXTINDERE
    # ============================================================
    add_heading_custom(doc, "VI. Ce se întâmplă la fiecare nivel de extindere", level=1, color=BLUE)

    doc.add_paragraph(
        "Această secțiune explică, pentru fiecare nivel de extindere, ce schimbări aduce în "
        "tratament și în prognostic."
    )

    add_heading_custom(doc, "1. Extindere în adâncime (T)", level=2)
    add_bullet(doc, [
        "T1a (doar mucoasa): tratament endoscopic, fără operație. Spitalizare 1-2 zile. Vindecare 90-95%.",
        "T1b (submucoasa): ESD endoscopic sau operație minoră. Risc de 15-25% să existe deja 1-2 ganglioni afectați.",
        "T2 (muscularis): operație obligatorie (esofagectomie). Vindecare 40-55% la 5 ani.",
        "T3 (adventice): Chimio-radio obligatoriu ÎNAINTE de operație (neoadjuvant). Tumora se micșorează, apoi se operează.",
        "T4a (pleură, pericard, diafragm): operabil dar mai dificil. Chimio-radio agresiv înainte. Posibilă rezecție „în bloc” cu țesutul invadat.",
        "T4b (aortă, trahee, coloană): NU se operează. Tratament paliativ cu chimio-radio + imunoterapie + stent esofagian dacă obstrucție.",
    ])

    add_heading_custom(doc, "2. Extindere în ganglioni (N)", level=2)
    p = doc.add_paragraph(
        "Ganglionii limfatici sunt o rețea de stații mici răspândite în jurul esofagului, în torace "
        "și abdomenul superior. Există zeci de grupuri. Cancerul se propagă prin limfă dintr-o "
        "stație în următoarea, deci extinderea ganglionară apare înainte de metastazele la distanță."
    )
    p = doc.add_paragraph()
    p.add_run("Ce înseamnă practic:").bold = True
    add_bullet(doc, [
        "N0 (fără ganglioni afectați): operația poate fi curativă singură, fără chimio obligatoriu.",
        "N1-N3: chimio-radio ÎNAINTE devine obligatoriu, indiferent de stadiul T. Imunoterapie adjuvantă utilă dacă rămâne boală reziduală după operație.",
        "Extinderea ganglionară ≥3 (N2-N3) dublează riscul de recidivă după operație, de aceea imunoterapia post-operator devine esențială.",
    ])

    add_heading_custom(doc, "3. Metastaze la distanță (M)", level=2)
    p = doc.add_paragraph(
        "Metastazele apar când celulele canceroase ajung în sânge și se stabilesc în organe "
        "îndepărtate. Este schimbarea cea mai importantă în logica tratamentului:"
    )
    rows = [
        ["M0 (fără metastaze)", "Tratament CURATIV — țintă = vindecare. Operație + chimio + radio + imunoterapie combinate."],
        ["M1 (cu metastaze)", "Tratament PALIATIV — țintă = prelungire viață + calitate. Chimio + imunoterapie continuu. Operația NU se face (excepții rare)."],
    ]
    add_table(doc, ["Scenariu", "Implicație"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Veste bună: imunoterapia a schimbat stadiul 4",
        "Înainte de 2020, stadiul 4 însemna supraviețuire medie 9-10 luni cu chimio simplă. "
        "Astăzi, cu pembrolizumab + chimio (KEYNOTE-590) sau nivolumab + chimio (CheckMate-648), "
        "supraviețuirea mediană a crescut la 13-15 luni, cu pacienți care trăiesc 3-5+ ani. "
        "PD-L1 CPS ≥10 identifică pacienții cu beneficiu maxim (răspuns complet posibil).",
        color=GREEN,
    )
    doc.add_page_break()

    # ============================================================
    # VII. CAT TIMP CT
    # ============================================================
    add_heading_custom(doc, "VII. Când vine rezultatul CT", level=1, color=BLUE)

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Surse: Affidea.ro, Memorial.ro, Reginamaria.ro, Medicales.ro (accesate 19.04.2026).")

    rows = [
        ["Genesis Medical Arad", "2-3 zile lucrătoare", "24h cu mențiune urgentă"],
        ["Affidea Fundeni", "2-3 zile lucrătoare", "2-3h în urgență"],
        ["SANADOR București", "24-48h", "2h verbal"],
        ["Sistem privat general", "2-3 zile", "2h verbal, scris în 24h"],
        ["Memorial.ro (rapid)", "Verbal 2h în urgență", "-"],
        ["Clinica Eminescu 100", "5-7 zile lucrătoare", "-"],
    ]
    add_table(doc, ["Clinică", "Timp standard", "Timp în urgență"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Concret pentru tatăl tău",
        "CT programat: luni 20.04.2026 ora 17:00 la Genesis Medical Clinic Micălaca.\n"
        "Rezultat scris așteptat: miercuri-joi (22-23.04.2026).\n\n"
        "ACȚIUNE: cere-i medicului Dr. Noufal (gastroenterolog, a emis biletul de trimitere) să scrie "
        "pe bilet „URGENT ONCOLOGIC” sau să contacteze direct radiologul Genesis → rezultat "
        "în 24h scris + comunicare verbală imediată după scanare.\n\n"
        "Cere raportul CT în format DICOM (CD) — îl poți trimite pentru a 2-a opinie la Charité, "
        "Heidelberg sau Cleveland MyConsult.",
        color=BLUE,
    )

    add_heading_custom(doc, "Biopsie histopatologică — timing", level=2)
    p = doc.add_paragraph()
    add_marker(p, "PROBABIL")
    p.add_run(
        "Rezultat histopatologic biopsie prelevată 17.04.2026: estimat 7-14 zile lucrătoare, "
        "deci între 24.04-01.05.2026. Monitor automat activ pe portalul Bioclinica (notificare "
        "instantanee când apare rezultatul)."
    )
    doc.add_page_break()

    # ============================================================
    # VIII. BENIGN
    # ============================================================
    add_heading_custom(doc, "VIII. Tratamente pentru leziune BENIGNĂ", level=1, color=GREEN)

    p = doc.add_paragraph(
        "Dacă biopsia arată leziune benignă, drumul este mult mai scurt. Tipuri posibile și "
        "tratamente:"
    )

    rows = [
        ["Polip / papilom scuamos", "Excrescență din mucoasă", "Rezecție endoscopică simplă, aceeași sesiune. Fără anestezie generală."],
        ["Leiomiom", "Tumoră benignă din musculatură (cea mai frecventă benignă)", "Urmărire periodică dacă <3cm asimptomatic; enucleere chirurgicală dacă >3cm sau simptomatic."],
        ["Chist esofagian", "Sac cu lichid", "Urmărire sau marsupializare."],
        ["Esofagită eozinofilică", "Inflamație cronică alergică", "Dietă + corticoizi."],
        ["Strictură peptică", "Îngustare din reflux cronic", "Dilatare endoscopică + PPI cronic."],
    ]
    add_table(doc, ["Leziune", "Ce este", "Tratament"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Scenariu tipic — leziune benignă",
        "EGD de control la 6-12 luni. Fără chimio, fără radio, fără operație majoră. "
        "Pacient trăiește normal, fără restricții majore. Medicație: eventual PPI (omeprazol) "
        "cronic dacă a existat reflux. Fără impact asupra speranței de viață.",
        color=GREEN,
    )
    doc.add_page_break()

    # ============================================================
    # IX. PRECANCEROS
    # ============================================================
    add_heading_custom(doc, "IX. Tratamente pentru leziune PRECANCEROASĂ", level=1, color=ORANGE)

    p = doc.add_paragraph(
        "Precanceros înseamnă celule modificate dar care NU sunt încă cancer. Fără tratament, "
        "30-40% dezvoltă cancer în 5-10 ani. Cu tratament endoscopic, vindecare aproape completă."
    )

    add_heading_custom(doc, "Tratamente endoscopice (FĂRĂ operație)", level=2)

    add_heading_custom(doc, "EMR — Endoscopic Mucosal Resection", level=3)
    add_bullet(doc, [
        "Se injectează lichid sub leziune → se ridică → se taie cu laț electric.",
        "Durată: 30-60 minute.",
        "Internare: ambulatoriu sau 1 noapte.",
        "Rată de succes: 80-90% pentru leziuni <2cm.",
        "Complicații: sângerare 1-5%, perforație <1%.",
    ])

    add_heading_custom(doc, "ESD — Endoscopic Submucosal Dissection", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Pimentel-Nunes et al., Endoscopy 2022 ESGE Guidelines + meta-analize ScienceDirect 2025.")
    add_bullet(doc, [
        "Tehnică mai avansată, permite rezecție „en bloc” (o singură bucată, nu fragmentată).",
        "Rată rezecție en-bloc: 90-95%.",
        "Rată rezecție curativă R0: 85-90% (superior EMR cu factor 2.29×).",
        "Recidivă locală: <5%.",
        "Folosit pentru leziuni >2cm sau cu suspiciune invazie submucosală.",
        "Centre cu experiență ESD în România: IOB Fundeni, Institutul Clinic Fundeni, Colțea, Regina Maria Băneasa — cer recomandare specifică la programare.",
    ])

    add_heading_custom(doc, "RFA — Radiofrequency Ablation", level=3)
    add_bullet(doc, [
        "Dispozitiv care „arde” controlat stratul superficial al esofagului.",
        "Eficacitate: 80-90% eradicare displazie.",
        "Indicat pentru esofag Barrett cu displazie (leziune difuză, nu focală).",
        "Mai puțin folosit în România pentru cazuri individuale (necesită sistem specializat).",
    ])

    doc.add_paragraph()
    add_callout(
        doc,
        "Scenariu tipic — leziune precanceroasă",
        "Pacient 66 ani cu Barrett + displazie înaltă → ESD într-o singură sesiune (60 minute, "
        "internare 1-2 zile) → urmărire endoscopică la 3, 6, 12 luni. Dacă rezecția a fost R0 "
        "(marginile curate) și urmărirea e normală → tratat definitiv.\n\n"
        "Fără chimio, fără radio, fără operație majoră.",
        color=GREEN,
    )
    doc.add_page_break()

    # ============================================================
    # X. MALIGN - STADII
    # ============================================================
    add_heading_custom(doc, "X. Tratamente pentru CANCER — pe stadii 1, 2, 3, 4", level=1, color=RED)

    # Stadiu 1
    add_heading_custom(doc, "Stadiul 1 (T1-T2 N0 M0) — „cancer foarte timpuriu”", level=2, color=BLUE)

    add_heading_custom(doc, "Stadiu 1 / T1a — NU necesită operație", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: NCCN Esophageal V1.2025 + ESMO 2022.")
    add_bullet(doc, [
        "Tratament: ESD endoscopic.",
        "Internare: 1-2 zile.",
        "Vindecare: 90-95%.",
        "Supraviețuire la 5 ani: ~85-90%.",
        "ZERO mortalitate legată de tratament.",
    ])

    add_heading_custom(doc, "Stadiu 1 / T1b — 2 opțiuni", level=3)
    rows = [
        ["ESD endoscopic", "Leziune <2cm, bine diferențiată, fără invazie limfovasculară", "Recuperare 1 săptămână", "Mortalitate <1%", "Supraviețuire 5 ani ~75%"],
        ["Esofagectomie", "Leziune mare, prost diferențiată, invazie limfovasculară", "Recuperare 6-12 săptămâni", "Mortalitate 3-5%", "Supraviețuire 5 ani 75-80%"],
    ]
    add_table(doc, ["Opțiune", "Când se alege", "Recuperare", "Mortalitate", "Supraviețuire"], rows)

    add_heading_custom(doc, "Stadiu 1 / T2 N0", level=3)
    add_bullet(doc, [
        "Standard: esofagectomie directă (pentru pacient fit).",
        "Alternativă 2024-2025 în discuție: chimio-radio neoadjuvant + operație.",
        "NCCN V1.2025 listează ambele opțiuni.",
        "Decizia depinde de starea generală, vârsta, comorbidități.",
    ])
    doc.add_paragraph()

    # Stadiu 2-3
    add_heading_custom(doc, "Stadiile 2-3 (local avansat) — standardul modern", level=2, color=BLUE)

    add_heading_custom(doc, "CROSS — chimio-radio neoadjuvant (5 săptămâni)", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: van Hagen et al., NEJM 2012; Eyck et al., JCO 2021 (10-year follow-up).")
    add_bullet(doc, [
        "Carboplatin + paclitaxel săptămânal × 5 săptămâni + radioterapie zilnică (41.4 Gy în 23 fracții).",
        "Pauză 6-8 săptămâni (pacientul se reface).",
        "Apoi esofagectomie.",
        "Supraviețuire mediană: 49 luni CROSS+chirurgie vs 24 luni chirurgie singură.",
        "Rezecție R0: 92% vs 69%.",
        "Răspuns patologic complet (pCR): 29% overall, 49% în ESCC.",
        "PREFERAT pentru ESCC + pacient cu risc cardiac — toxicitate mai mică decât FLOT.",
    ])

    add_heading_custom(doc, "FLOT — chimio perioperatorie (8 cicluri, fără radio)", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Al-Batran et al., Lancet 2019 (FLOT4); ESOPEC 2024 (NEJM).")
    add_bullet(doc, [
        "4 cicluri ÎNAINTE de operație + 4 cicluri DUPĂ (total 8 cicluri, 16 săptămâni).",
        "Medicamente: docetaxel + oxaliplatin + 5-FU + leucovorin (la 2 săptămâni).",
        "Indicat pentru adenocarcinom GEJ / gastric / distal esofagian — NU pentru ESCC.",
        "Studiu ESOPEC 2024: FLOT superior CROSS pentru adenocarcinom — supraviețuire mediană 66 vs 37 luni.",
        "ATENȚIE CARDIAC: 5-FU are risc cardiac (spasm coronarian). La pacient cu stent vechi (>10 ani, asimptomatic), FLOT posibil cu monitorizare ECG + troponină.",
    ])

    add_heading_custom(doc, "CheckMate-577 — nivolumab adjuvant DUPĂ operație", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Kelly et al., NEJM 2021.")
    add_bullet(doc, [
        "Nivolumab 240mg la 2 săptămâni × 1 an, la pacienți cu boală reziduală post-CROSS+operație.",
        "Supraviețuire fără recidivă: 22.4 luni (nivolumab) vs 11.0 luni (placebo) — dublare.",
        "Aprobat EMA + DECONTAT în România prin Programul Național de Oncologie.",
        "STANDARD DE ÎNGRIJIRE actual pentru non-pCR după neoadjuvant.",
    ])

    add_heading_custom(doc, "CROSS vs FLOT — când ce se alege", level=3)
    rows = [
        ["ESCC (orice stadiu local)", "CROSS"],
        ["EAC esofagian distal / cardia", "CROSS sau FLOT (discuție tumor board)"],
        ["Adenocarcinom GEJ Siewert II-III", "FLOT"],
        ["Pacient vârstnic (70+) / cardiac / renal", "CROSS (toxicitate mai mică)"],
        ["Pacient tânăr fit, GEJ/gastric", "FLOT"],
    ]
    add_table(doc, ["Situație", "Preferat"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Pentru tatăl tău — cardiac, 66 ani",
        "Dacă rezultatul este adenocarcinom GEJ (cel mai probabil, dat fiind vârsta și factorii de "
        "risc): discuție cu oncologul despre FLOT vs CROSS. Dacă FLOT e ales, cere MONITORIZARE "
        "CARDIACĂ — ECG pre/post fiecare ciclu + troponină. Dacă apar semne de ischemie cardiacă "
        "sub FLOT → SWITCH la CROSS.\n\n"
        "Dacă e ESCC → CROSS de prima alegere (profil cardiac mai bun).",
        color=BLUE,
    )
    doc.add_paragraph()

    # Stadiu 4
    add_heading_custom(doc, "Stadiul 4 (metastatic) — tratament paliativ modern", level=2, color=RED)

    doc.add_paragraph(
        "Scopul se schimbă fundamental: NU mai e vindecare, ci PRELUNGIRE + CALITATE. "
        "Imunoterapia modernă a schimbat dramatic acest stadiu în 2020-2024."
    )

    add_heading_custom(doc, "KEYNOTE-590 — pembrolizumab + chimio prima linie", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Sun et al., Lancet 2021; update ESMO 2023, 5-year follow-up 2025.")
    add_bullet(doc, [
        "Pembrolizumab 200mg la 3 săpt + cisplatin/5-FU, prima linie.",
        "Supraviețuire mediană: 12.4 luni (pembro+chimio) vs 9.8 luni (chimio singur).",
        "ESCC PD-L1 CPS ≥10: supraviețuire 13.9 vs 8.8 luni (beneficiu maxim).",
        "Aprobat EMA + DECONTAT prin CNAS pentru PD-L1 CPS ≥10.",
    ])

    add_heading_custom(doc, "CheckMate-648 — nivolumab combinații în ESCC", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Doki et al., NEJM 2022.")
    add_bullet(doc, [
        "Nivolumab + chimio: supraviețuire 15.4 luni vs 9.1 luni (chimio singur) — pentru ESCC PD-L1 ≥1%.",
        "Nivolumab + ipilimumab (dublă imunoterapie, FĂRĂ chimio): 13.7 luni — opțiune pentru cei care nu tolerează chimio.",
    ])

    add_heading_custom(doc, "Trastuzumab + imunoterapie (pentru HER2+ adenocarcinom)", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Janjigian et al. (KEYNOTE-811), Nature 2024.")
    add_bullet(doc, [
        "Doar pentru adenocarcinom HER2+ (IHC 3+ sau IHC 2+/FISH amplificat) — ~15-25% din EAC.",
        "Combinație: trastuzumab + chimio + pembrolizumab.",
        "Supraviețuire: 20.0 luni vs 16.8 luni (fără pembro).",
        "Aprobat EMA 2024 pentru HER2+ PD-L1 CPS ≥1.",
    ])

    add_heading_custom(doc, "Zolbetuximab (CLDN18.2+) — medicament nou aprobat 2024", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: SPOTLIGHT (Lancet 2023), GLOW (Nat Med 2023), aprobat EMA septembrie 2024.")
    add_bullet(doc, [
        "Pentru adenocarcinom GEJ/gastric CLDN18.2+ (≥75% tumoră), HER2-negativ.",
        "Zolbetuximab + chimio (FOLFOX sau CAPOX).",
        "Supraviețuire: 18.2 luni vs 15.5 luni.",
        "INCERT: disponibilitate în România prin PNO la aprilie 2026 — verificat pe CNAS, date publice consolidate lipsesc. Aprobat EMA din 09.2024, posibil încă nu în PNO.",
        "Pentru acces: cerere specială la CNAS + contract cost-volum.",
    ])

    add_heading_custom(doc, "Tabel sumar — stadiul 4 prima linie", level=3)
    rows = [
        ["ESCC PD-L1 CPS ≥10", "Pembrolizumab + cisplatin/5-FU (KEYNOTE-590)"],
        ["ESCC PD-L1 ≥1%", "Nivolumab + chimio sau Nivo+Ipi (CheckMate-648)"],
        ["EAC/GEJ PD-L1 CPS ≥5", "Nivolumab + FOLFOX/CapeOX (CheckMate-649)"],
        ["EAC/GEJ HER2+", "Trastuzumab + chimio ± pembrolizumab (KEYNOTE-811)"],
        ["EAC/GEJ CLDN18.2+", "Zolbetuximab + chimio (SPOTLIGHT/GLOW)"],
    ]
    add_table(doc, ["Biomarker / situație", "Tratament recomandat"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "IMPORTANT — testare biomarker",
        "Toate aceste tratamente depind de biomarkerii de pe biopsie:\n"
        "• PD-L1 CPS (Combined Positive Score) — obligatoriu pentru imunoterapie\n"
        "• HER2 (IHC/FISH) — pentru adenocarcinom, decide trastuzumab\n"
        "• MSI/MMR — predictor puternic pentru imunoterapie\n"
        "• CLDN18.2 — pentru zolbetuximab\n\n"
        "CERE explicit oncologului să comande TOATE aceste teste pe biopsia deja prelevată.",
        color=ORANGE,
    )
    doc.add_page_break()

    return doc


def build_document_part3(doc):
    # ============================================================
    # XI. OPERATIE
    # ============================================================
    add_heading_custom(doc, "XI. Operația esofagectomie — cât de grea este", level=1, color=BLUE)

    doc.add_paragraph(
        "Esofagectomia este una din cele mai mari operații din chirurgia digestivă. NU este "
        "imposibilă la 66 ani cu stent — dar cere alegerea atentă a tehnicii și a centrului."
    )

    add_heading_custom(doc, "Tipuri de operație", level=2)
    rows = [
        ["Ivor Lewis", "Laparotomie + toracotomie dreaptă", "Intratoracică", "5-7h", "Tumori distale/GEJ"],
        ["McKeown", "Laparotomie + toracotomie + cervicotomie", "Cervicală", "7-9h", "Tumori medii/proximale"],
        ["Transhiatal", "Laparotomie + cervicotomie", "Cervicală", "4-6h", "Probleme pulmonare"],
        ["MIE (Minim Invaziv)", "Laparoscopic + toracoscopic", "Variabilă", "6-8h", "STANDARD modern"],
        ["RAMIE (Robotic)", "Asistat robotic Da Vinci", "Variabilă", "7-9h", "Complicații cardiace ↓53%"],
    ]
    add_table(doc, ["Tehnică", "Abordare", "Anastomoză", "Durată", "Indicație"], rows)

    add_heading_custom(doc, "Pentru tatăl tău (66 ani, stent 2012) — recomandare", level=2)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: Annals of Esophagus 2024 (Yip et al.) + MIRO trial NEJM 2019 + ROBOT trial Ann Surg 2019.")
    add_bullet(doc, [
        "RAMIE (Robotic MIE) PREFERAT — reduce complicațiile cardiopulmonare cu 53% vs open.",
        "MIE clasic (laparoscopic + toracoscopic) — complicații pulmonare 15.6% vs 30.1% open.",
        "EVITĂ operația deschisă clasică — risc cardiac dublu la pacient cu stent.",
        "Internare: 10 zile RAMIE vs 14 zile open.",
        "Mortalitate 90 zile: 3-4% RAMIE vs 8-10% open.",
    ])

    add_heading_custom(doc, "Centre cu RAMIE / MIE în România", level=2)
    add_bullet(doc, [
        "Ponderas Academic Hospital București — robot Da Vinci, singurul centru RO acreditat SRC.",
        "SANADOR București — robot Da Vinci Xi.",
        "Spitalul Universitar de Urgență București — are robot.",
        "Institutul Clinic Fundeni (Chirurgie IV — Prof. Bacalbașa) — enumeră explicit chirurgie esofagiană.",
        "Institutul Oncologic București (IOB) — Chirurgie Toracică (Dr. Natalia Motaș).",
    ])

    add_heading_custom(doc, "Volume — relația operații/an și rezultate", level=2)
    rows = [
        ["Centre expert (>50/an)", "1.5-3%", "3-5%"],
        ["High-volume (>30/an)", "2-5%", "4-7%"],
        ["Mid-volume (10-30/an)", "5-8%", "7-12%"],
        ["Low-volume (<10/an)", "8-12%", "12-18%"],
    ]
    add_table(doc, ["Tip centru", "Mortalitate 30 zile", "Mortalitate 90 zile"], rows)

    add_heading_custom(doc, "Complicații tipice + frecvență", level=2)
    rows = [
        ["Pneumonie", "25-30%", "15-20%"],
        ["Anastomotic leak (fistulă)", "10-20%", "8-12%"],
        ["Chilotorax", "2-5%", "1-3%"],
        ["Fibrilație atrială", "15-20%", "12-15%"],
        ["Infecție plagă", "5-10%", "3-5%"],
        ["Embolism pulmonar", "2-4%", "1-2%"],
        ["Leziune nerv recurent", "5-10%", "3-5%"],
        ["Reintervenție", "10-15%", "5-10%"],
        ["IMA perioperator (cardiac)", "1-3% pop generală", "3-6% la stent"],
    ]
    add_table(doc, ["Complicație", "Frecvență overall", "Centre expert"], rows)

    add_heading_custom(doc, "Viața după esofagectomie", level=2)
    add_bullet(doc, [
        "Dumping syndrome (greață după mese cu carbohidrați) — 40-60% primele 6 luni.",
        "Reflux gastro-esofagian cronic — 60-80% în primii 2 ani; PPI cronic.",
        "Pierdere greutate: 10-15% din inițial, primele 6 luni.",
        "Disfagie / strictură anastomotică: 15-30% necesită dilatare.",
        "Dieta: 5-6 mese mici/zi, porții 150-250ml, NU lichide în timpul mesei.",
        "Calitatea vieții la 12-18 luni: 70-80% din preoperator.",
    ])

    doc.add_paragraph()
    add_callout(
        doc,
        "Evaluare preoperatorie OBLIGATORIE pentru pacient cu stent",
        "1. Consult cardiologic complet (cu raportul vechi stent 2012).\n"
        "2. Ecocardiografie (FEVS, funcție diastolică).\n"
        "3. Stress ecocardiografie / scintigrafie miocardică.\n"
        "4. Eventual coronarografie de control (stent >10 ani).\n"
        "5. Spirometrie.\n"
        "6. Scor ASA (ideal 2-3).\n"
        "7. Scor RCRI (Revised Cardiac Risk Index).\n\n"
        "Fără aceste evaluări, operația majoră NU este sigură.",
        color=RED,
    )
    doc.add_page_break()

    # ============================================================
    # XII. IMUNOTERAPIE
    # ============================================================
    add_heading_custom(doc, "XII. Imunoterapia — revoluția 2020-2025", level=1, color=GREEN)

    doc.add_paragraph(
        "Imunoterapia a schimbat peisajul tratamentului pentru cancer esofagian începând cu 2020. "
        "Medicamentele (anticorpi monoclonali) „dau voie” sistemului imun să recunoască și să "
        "atace celulele canceroase."
    )

    add_heading_custom(doc, "Cum funcționează — în termeni simpli", level=2)
    doc.add_paragraph(
        "Celulele canceroase au la suprafață o proteină PD-L1 care funcționează ca un „scut” — "
        "„nu mă ataca, sunt celulă normală”. Pembrolizumab și nivolumab blochează acest scut → "
        "sistemul imun recunoaște celulele canceroase și le distruge."
    )

    add_heading_custom(doc, "Medicamente aprobate în UE pentru cancer esofagian", level=2)
    rows = [
        ["Pembrolizumab (Keytruda, MSD)", "ESCC + EAC/GEJ PD-L1 CPS ≥10 + chimio", "Prima linie"],
        ["Nivolumab (Opdivo, BMS)", "ESCC PD-L1 ≥1%; EAC/GEJ CPS ≥5 + chimio", "Prima + adjuvant"],
        ["Nivolumab adjuvant", "Post-CROSS+operație cu boală reziduală", "1 an adjuvant"],
        ["Ipilimumab (Yervoy)", "+ nivolumab ESCC PD-L1 ≥1%", "Prima linie"],
        ["Trastuzumab (Herceptin)", "EAC/GEJ HER2+ + chimio", "Prima linie"],
        ["Trastuzumab + Pembrolizumab", "HER2+ PD-L1 CPS ≥1 (KEYNOTE-811)", "Prima linie"],
        ["Zolbetuximab (Vyloy)", "CLDN18.2+ (≥75%) + chimio (EMA 09.2024)", "Prima linie"],
        ["T-DXd (Enhertu)", "HER2+ refractar post-trastuzumab", "Linia 2+"],
        ["Ramucirumab (Cyramza)", "EAC/GEJ linia 2 (RAINBOW)", "Linia 2"],
    ]
    add_table(doc, ["Medicament", "Indicație EMA", "Linia"], rows)

    add_heading_custom(doc, "Testare biomarker OBLIGATORIE pe biopsie", level=2)
    p = doc.add_paragraph(
        "Aceste teste NU se fac automat — trebuie cerute explicit oncologului. Esențiale pentru "
        "imunoterapie:"
    )
    rows = [
        ["PD-L1 CPS", "% celule PD-L1+ / total × 100", "CPS ≥1 nivolumab ESCC; CPS ≥10 pembro"],
        ["HER2 (IHC/FISH)", "Amplificare HER2", "HER2+ → trastuzumab"],
        ["MSI/MMR", "Instabilitate microsateliți", "MSI-H → răspuns EXCELENT la imunoterapie"],
        ["CLDN18.2", "Claudina 18.2 (≥75%)", "Pozitiv → zolbetuximab"],
        ["NTRK fusion", "Fuziuni gene NTRK (rare)", "Pozitiv → larotrectinib"],
    ]
    add_table(doc, ["Biomarker", "Ce este", "Implicație"], rows)

    add_heading_custom(doc, "Accesare în România prin PNO", level=2)
    add_numbered(doc, [
        "Confirmare histopatologică (biopsie).",
        "Testare biomarker (PD-L1, HER2, MSI, CLDN18.2).",
        "Stadializare completă (CT, PET-CT dacă indicat).",
        "Consult multidisciplinar (tumor board).",
        "Dosar la Comisia de Oncologie CNAS prin oncolog.",
        "Aprobare: 2-8 săptămâni (variabil).",
        "Administrare GRATUITĂ prin contract cost-volum producător.",
    ])

    doc.add_paragraph()
    add_callout(
        doc,
        "Efecte adverse majore ale imunoterapiei",
        "Pneumonita interstițială (2-5%) — URGENȚĂ, oprire medicament + corticoizi.\n"
        "Colita imună (10-20%) — diaree severă, oprire + corticoizi.\n"
        "Hepatita imună — transaminaze crescute.\n"
        "Endocrinopatii — hipotiroidism, insuficiență suprarenală (1-5%).\n\n"
        "Orice simptom nou sub imunoterapie → contact imediat oncolog.",
        color=ORANGE,
    )
    doc.add_page_break()

    # ============================================================
    # XIII. TRIAL-URI
    # ============================================================
    add_heading_custom(doc, "XIII. Trial-uri clinice active cu centre în România", level=1, color=BLUE)

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: clinicaltrials.gov API v2, consultat 19.04.2026.")

    doc.add_paragraph(
        "Trial-urile clinice oferă acces la medicamente încă neaprobate GRATUIT. Medicamentul "
        "experimental e furnizat de sponsor. Pacientul suportă doar transport/cazare la vizite."
    )

    add_heading_custom(doc, "Trial-uri RECRUTING ACTIV în România (aprilie 2026)", level=2)

    add_heading_custom(doc, "NCT06644781 — IDeate-Esophageal-01 (Ifinatamab Deruxtecan)", level=3)
    add_bullet(doc, [
        "Faza 3, sponsor Daiichi Sankyo.",
        "Pentru: ESCC avansat/metastatic, linia 2.",
        "Medicament: I-DXd 12 mg/kg IV la 3 săpt (ADC anti-B7-H3).",
        "Centre RO: Memorial București, SC Radiotherapy Cluj-Floresti, Sf. Nectarie Craiova, Sigmedical Suceava.",
        "Contact: +1 908 992 6400, CTRinfo_us@daiichisankyo.com.",
    ])

    add_heading_custom(doc, "NCT06532006 — HLX22 + Trastuzumab + Pembrolizumab (HER2+ GEJ)", level=3)
    add_bullet(doc, [
        "Faza 3, sponsor Shanghai Henlius Biotech.",
        "Pentru: adenocarcinom gastric/GEJ HER2+ prima linie.",
        "Criteriu IMPORTANT: LVEF ≥55%.",
        "Centre RO: Inst. Regional Gastro Cluj, Amethyst Craiova, Sf. Nectarie Floresti, Pelican Oradea.",
        "Contact: Ying Li +86 18810366427, ying_li1@henlius.com.",
    ])

    add_heading_custom(doc, "Trial-uri NOT_YET_RECRUITING (deschidere 2026)", level=2)
    add_bullet(doc, [
        "NCT07221149 — Pumitamig + FOLFOX (adenocarcinom GEJ prima linie) — BMS, Cluj/Craiova/Iași.",
        "NCT06901531 — LUCERNA (Zolbetuximab + Pembrolizumab + chimio) — Astellas, Cluj/Floresti/Craiova/TM/Iași/București.",
        "NCT07431281 — CLARITY-Gastric 02 (Sonesitatug Vedotin) — AstraZeneca.",
    ])

    add_heading_custom(doc, "Aplicare la un trial", level=2)
    add_numbered(doc, [
        "Oncolog curant oferă trimitere după stadializare + biomarkeri.",
        "Contact direct PI local (telefon/email sponsor).",
        "Screening la centrul trial-ului — reconfirmare biomarkeri centralizată.",
        "Consimțământ informat.",
        "Randomizare → tratament.",
        "Total: 3-6 săptămâni screening→tratament.",
    ])

    doc.add_paragraph()
    add_callout(
        doc,
        "Recomandare practică",
        "La primul consult oncologic după rezultate complete, cere explicit: „Există trial-uri "
        "clinice la care aș putea fi eligibil?” Pentru Arad, cel mai aproape cu trial-uri ACTIVE "
        "este Cluj-Napoca + Timișoara (OncoHelp) + București (IOB, Memorial).",
        color=BLUE,
    )
    doc.add_page_break()

    return doc


def build_document_part4(doc):
    # ============================================================
    # XIV. CENTRE ROMANIA
    # ============================================================
    add_heading_custom(doc, "XIV. Centre oncologice — România", level=1, color=BLUE)

    add_heading_custom(doc, "Centre locale (Arad + Timișoara — prioritar geografic)", level=2)

    add_heading_custom(doc, "SCJU Arad — Chirurgie Generală I", level=3)
    add_bullet(doc, [
        "Adresă: Str. Andrényi Károly 2-4, 310037 Arad.",
        "Telefon: 0357 407 200 (int. 283 pentru Chirurgie I).",
        "Email: secretariat@scjarad.ro.",
        "Șef: Prof. Univ. Dr. Bogdan Totolici.",
        "Echipă: Conf. Dr. Carmen Neamțu, Dr. Adrian Pavel, Dr. Petre Ciucuriță.",
        "Specialitate declarată explicit: cancer esofag, stomac, hepatic, biliar, pancreatic.",
        "PRIMA LINIE pentru consult inițial — aproape de casă.",
    ])

    add_heading_custom(doc, "SCJU Pius Brînzeu Timișoara (~60 km)", level=3)
    add_bullet(doc, [
        "Adresă: Bd. Liviu Rebreanu 156, Timișoara.",
        "Programări: 0752 297 252 / 0748 331 296.",
        "Email: chirurgie1@hosptm.ro.",
        "Șef Chirurgie I: Prof. Dr. Hab. Sorin Olariu.",
    ])

    add_heading_custom(doc, "OncoHelp Timișoara (ONG cu contract CAS)", level=3)
    add_bullet(doc, [
        "Str. C. Porumbescu 57-59, Timișoara 300239.",
        "Telefon: 0256 495 403 / 0356 460 995 / 0753 595 959.",
        "Radioterapie: 0752 266 170.",
        "Email: programari@oncohelp.ro.",
        "Președinte: Șerban Negru.",
        "Servicii GRATUITE pentru asigurați CJAS Timiș.",
    ])

    add_heading_custom(doc, "Centre naționale de referință (București)", level=2)

    add_heading_custom(doc, "Institutul Oncologic București (IOB) Fundeni", level=3)
    add_bullet(doc, [
        "Șos. Fundeni 252, Sector 2, București.",
        "Centrală: 021 227 1000 / 1001.",
        "Cameră gardă: 021 227 1094.",
        "Programări: 021 227 1595.",
        "Chirurgie Toracică: 021 227 1347 (Dr. Natalia Motaș).",
        "Email: secretariat@iob.ro.",
        "Cel mai mare centru oncologic din România.",
    ])

    add_heading_custom(doc, "Institutul Clinic Fundeni — Chirurgie IV", level=3)
    add_bullet(doc, [
        "Șos. Fundeni 258, Sector 2 (instituție diferită de IOB).",
        "Telefon: 021 275 0700 (L-V 09:00-15:00).",
        "Șef: Prof. Dr. Nicolae Bacalbașa.",
        "Specialități declarate: chirurgie toracică + ESOFAGIANĂ, ficat, pancreas.",
        "Singura secție ICF care declară explicit chirurgie esofagiană.",
    ])

    add_heading_custom(doc, "Ponderas Academic Hospital (Regina Maria)", level=3)
    add_bullet(doc, [
        "Str. Nicolae Caramfil 85A, Sector 1.",
        "Telefon: 021 9886.",
        "Email: office.ponderas@reginamaria.ro.",
        "SINGURUL Centru de Excelență Chirurgie Robotică acreditat SRC.",
        "Dr. Felix Dobrițoiu — șef Chirurgie Toracică-Digestivă.",
        "Angajament: diagnostic + inițiere tratament în 7 zile.",
    ])

    add_heading_custom(doc, "SANADOR — Centrul Oncologic", level=3)
    add_bullet(doc, [
        "Str. Sevastopol 5, Sector 1.",
        "Telefon: 021 9699.",
        "Radioterapie DECONTATĂ CAS; chirurgie = privat.",
        "Da Vinci Xi + Tumor Board.",
    ])

    add_heading_custom(doc, "Institutul Oncologic Cluj (IOCN)", level=3)
    add_bullet(doc, [
        "Str. Republicii 34-36, Cluj-Napoca.",
        "Telefoane: 0264 598 362/3/4.",
        "Programări online: programari.iocn.ro.",
        "Al doilea centru de referință după IOB.",
    ])
    doc.add_page_break()

    # ============================================================
    # XV. CENTRE UE
    # ============================================================
    add_heading_custom(doc, "XV. Centre oncologice UE — accesibile prin S2/E112", level=1, color=BLUE)

    add_callout(
        doc,
        "ALERTE importante",
        "• AKH Viena NU mai primește pacienți internaționali din 2019. Alternativa: Wiener Privatklinik.\n"
        "• Rinecker Proton Therapy München ÎNCHIS din 2019. Alternativa: HIT Heidelberg.",
        color=RED,
    )

    add_heading_custom(doc, "Wiener Privatklinik (WPK) — Viena", level=2)
    add_bullet(doc, [
        "Pelikangasse 15, A-1090 Vienna.",
        "Telefon: +43 1 40180-7051 / International: +43 1 40180-8700.",
        "Email: office@wpk.at, info@wpk.at.",
        "~100 medici de la AKH colaborează cu WPK (90% profesori).",
        "Birou de reprezentanță București + Timișoara — contact prin office@wpk.at.",
        "Acceptă S2 UE. 550+ pacienți români tratați H1 2019.",
    ])

    add_heading_custom(doc, "Charité Berlin", level=2)
    add_bullet(doc, [
        "Int. Patient Office: +49 30 450 578 244 (L-V 12:00-14:00).",
        "Email: international-patients@charite.de.",
        "Stomach/Esophageal Center: +49 30 450 564 222.",
        "Tumor coordinator esofag: PD Dr. Annika Kurreck.",
        "Certificări DKG cancer gastric + esofagian.",
        "Cel mai mare centru oncologic Europa certificat DKG (>14000 pacienți/an).",
        "Acceptă S2 UE.",
    ])

    add_heading_custom(doc, "Heidelberg University Hospital + NCT + HIT", level=2)
    add_bullet(doc, [
        "Int. Office: Im Neuenheimer Feld 400, 69120 Heidelberg.",
        "Telefon: +49 6221 56-6243.",
        "Email: international.office@med.uni-heidelberg.de.",
        "Visceral Surgery: +49 6221 56-6110.",
        "10 zile lucrătoare pentru ofertă scrisă.",
        "HIT Heidelberg = UNUL DIN PUȚINELE centre lume cu terapie proton + carbon-ion + heliu.",
        "Timp așteptare pre-chirurgie: 1-2 săptămâni.",
    ])

    add_heading_custom(doc, "Gustave Roussy Paris (rank #6 mondial)", level=2)
    add_bullet(doc, [
        "114 rue Édouard-Vaillant, 94805 Villejuif Cedex.",
        "Telefon: +33 1 42 11 42 11.",
        "Email: internationalpatients@gustaveroussy.fr.",
        "Head: Pr Axel Le Cesne.",
        "PRIMUL centru luptă împotriva cancerului în Europa.",
        "ESMO Designated Centre.",
        "12000+ pacienți noi/an. Tratează cancer esofag explicit.",
        "Acceptă formular S2 UE.",
    ])

    add_heading_custom(doc, "Semmelweis University Budapest", level=2)
    add_bullet(doc, [
        "Budapesta: ~280 km de Arad (3h cu mașina).",
        "Principala universitate medicală din Ungaria.",
        "https://semmelweis.hu/english/patient-services/",
    ])

    add_heading_custom(doc, "Centre Proton Therapy", level=2)
    rows = [
        ["PTC Praga (Cehia)", "Budínova 2437/1a, Praha 8", "+420 222 999 000", "~890 km"],
        ["MedAustron Wr. Neustadt (Austria)", "Marie Curie-Str. 5", "+43 2622 26100 300", "~780 km"],
        ["HIT Heidelberg (Germania)", "Im Neuenheimer Feld", "via int. office", "~1600 km"],
    ]
    add_table(doc, ["Centru", "Adresă", "Telefon", "Distanță"], rows)

    doc.add_paragraph()
    add_callout(
        doc,
        "Proton therapy — când e utilă",
        "Pentru cancer esofagian, proton therapy e evaluată individual (nu standard). Indicată "
        "SPECIAL la pacientul cardiac — reduce doza pe cord cu 30-50% vs IMRT. S2 la CNAS "
        "justificată dacă radioterapia clasică ar crește riscul cardiac.",
        color=BLUE,
    )
    doc.add_page_break()

    # ============================================================
    # XVI. SECOND OPINION
    # ============================================================
    add_heading_custom(doc, "XVI. Second opinion internațional (fără deplasare)", level=1, color=BLUE)

    add_heading_custom(doc, "Dana-Farber — Online Second Opinion", level=2)
    add_bullet(doc, [
        "Preț: $3000 USD. Timp: 3-4 săptămâni.",
        "Telefon: +1 877 442 3324.",
        "Email: DanaFarberExpertSecondOpinion@myaccesshope.org.",
        "Doar raport scris, fără video.",
        "INCERT pentru România — verifică email înainte de plată.",
        "Esophageal Center tratează >200 cazuri/an.",
    ])

    add_heading_custom(doc, "Cleveland Clinic MyConsult", level=2)
    add_bullet(doc, [
        "Preț pacient internațional: $4500 USD.",
        "Timp: câteva zile.",
        "Telefon: +1 216 444 3223.",
        "Oncology inclusă. România ACCEPTATĂ.",
        "Cleveland colectează documente cu consimțământul pacientului.",
    ])

    add_heading_custom(doc, "Memorial Sloan Kettering (MSK)", level=2)
    add_bullet(doc, [
        "Telefon: +1 212 639 4900.",
        "Bobst International Center, 160 East 53rd St, NYC.",
        "Interpreți gratuiți.",
        "Fără recomandare medicală necesară.",
    ])

    add_heading_custom(doc, "Alternative europene (mai ieftine)", level=2)
    add_bullet(doc, [
        "Charité Second Opinion (Berlin) — preț mult mai mic + S2 posibil.",
        "WPK Bucharest Office — a 2-a opinie prin Wiener Privatklinik.",
    ])
    doc.add_page_break()

    # ============================================================
    # XVII. PNO CNAS
    # ============================================================
    add_heading_custom(doc, "XVII. Programul Național de Oncologie CNAS", level=1, color=BLUE)

    doc.add_paragraph(
        "Programul Național de Oncologie (PNO) acoperă tratamentele oncologice complexe pentru "
        "toți asigurații. Sub-programe:"
    )
    rows = [
        ["P6.7.1", "Medicamentos: citostatice, imunoterapie, terapii țintite"],
        ["P6.7.2", "Radioterapie (IMRT, stereotaxie, brahiterapie)"],
        ["P6.7.3", "PET-CT (cu aprobare comisie)"],
    ]
    add_table(doc, ["Sub-program", "Acoperă"], rows)

    add_heading_custom(doc, "Decontat prin PNO pentru cancer esofagian", level=2)
    add_bullet(doc, [
        ("CERT", "Citostatice: 5-FU, cisplatin, oxaliplatin, irinotecan, taxani, capecitabina."),
        ("CERT", "Imunoterapie: pembrolizumab, nivolumab, durvalumab."),
        ("CERT", "Terapii țintite: trastuzumab (HER2+), ramucirumab."),
        ("CERT", "Radioterapie IMRT pentru cancer esofagian."),
        ("PROBABIL", "Zolbetuximab (aprobat EMA 09.2024) — posibil încă NU în PNO; verificare directă."),
    ])

    add_heading_custom(doc, "Procedura aprobare dosar", level=2)
    add_numbered(doc, [
        "Oncolog pregătește dosar: referat + consimțământ + diagnostic + biomarkeri + stadializare.",
        "Depunere la CAS Arad (CASARAD).",
        "Evaluare Comisia de Oncologie CNAS.",
        "Aprobare: 2-8 săptămâni.",
        "Administrare GRATUIT prin spital cu contract cost-volum.",
    ])

    add_heading_custom(doc, "Formular S2 pentru tratament UE", level=2)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursa: CNAS Ordinul 762/2021 + Directiva 2011/24/EU.")

    add_bullet(doc, [
        "Cele 2 condiții CUMULATIVE: (1) serviciu în pachetul RO; (2) nu poate fi furnizat în RO în timp rezonabil medical.",
        "Termen legal soluționare: 5 zile lucrătoare.",
        "Documente: cerere tip, CI, dosar medical + argumentare, ACCEPT scris de la spital UE.",
        "CNAS NU poate refuza emiterea S2 dacă condițiile sunt îndeplinite.",
    ])

    add_heading_custom(doc, "Contact CASARAD (Casa Arad)", level=3)
    add_bullet(doc, [
        "Adresă: Bd. Revoluției 45, 310181 Arad.",
        "Telefon: 0257 270 202 / 0357 401 180 / 0357 401 175.",
        "Tel verde: 0800 800 968.",
        "Email: casar@casar.ro.",
        "Card european: casar_card@casar.ro.",
    ])
    doc.add_page_break()

    return doc


def build_document_part5(doc):
    # ============================================================
    # XVIII. DREPTURI PACIENT
    # ============================================================
    add_heading_custom(doc, "XVIII. Drepturi pacient oncologic — indemnizație și handicap", level=1, color=BLUE)

    add_heading_custom(doc, "Grad de handicap — Legea 448/2006", level=2)
    rows = [
        ["GRAV", "Dependență totală, prognostic sever", "Indemnizație însoțitor 2574 RON/lună; scutire impozit venit; 12 CFR/an + 24 urban/an gratis"],
        ["ACCENTUAT", "Limitări funcționale importante", "Scutire impozit venit; 6 CFR/an; 12 urban/an"],
        ["MEDIU", "Stări mai puțin severe", "Reduceri locale"],
    ]
    add_table(doc, ["Grad", "Descriere", "Beneficii"], rows)

    add_heading_custom(doc, "Dosar certificat handicap (DGASPC Arad)", level=2)
    add_numbered(doc, [
        "Cerere tip.",
        "Copie CI + original.",
        "Scrisoare medic familie (max 30-60 zile).",
        "Referat ONCOLOG (diagnostic, stadiu TNM, plan terapeutic).",
        "Bilete ieșire + analize + imagistică.",
        "Ancheta socială primăria Nădlac.",
        "Adeverință Casa Județeană Pensii Arad.",
    ])

    p = doc.add_paragraph()
    p.add_run("Timp procesare: 60 zile evaluare + 15 zile decizie.").bold = True

    add_heading_custom(doc, "Pensie pacient pensionar", level=2)
    add_bullet(doc, [
        "Pensia limită vârstă CONTINUĂ neschimbată.",
        "NU se schimbă pe invaliditate.",
        "SE ADAUGĂ indemnizația de însoțitor (~2574 RON/lună grad GRAV).",
        "Cumul pensie + indemnizație.",
    ])

    add_heading_custom(doc, "Transport la tratament CNAS", level=2)
    add_bullet(doc, [
        "Ambulanța pentru pacient intransportabil — decontată.",
        "Transport personal Arad-București NU direct decontat.",
        "Cea mai bună opțiune: gradul GRAV → 12 călătorii CFR/an gratuit.",
    ])
    doc.add_page_break()

    # ============================================================
    # XIX. NUTRITIE PSIHO PALIATIE
    # ============================================================
    add_heading_custom(doc, "XIX. Nutriție, suport psihologic, paliație", level=1, color=BLUE)

    add_heading_custom(doc, "Nutriție perioperator și în chimio/radio", level=2)
    add_bullet(doc, [
        "Aport caloric: 25-30 kcal/kg/zi; 30-35 kcal/kg/zi la cașexie.",
        "Proteină: 1.2-1.5 g/kg/zi; până la 2 g/kg/zi post-operator.",
        "Hidratare: 30-35 mL/kg/zi (ajustat cardiac/renal).",
        "Screening obligatoriu: MUST / NRS-2002 / PG-SGA la fiecare consult.",
    ])

    add_heading_custom(doc, "Imunonutriție preoperator (ESPEN 2023)", level=3)
    add_bullet(doc, [
        "Arginina + omega-3 + nucleotide, 3×250ml/zi × 5-7 zile pre-op.",
        "Reducere complicații infecțioase 20-30% (Cochrane 2022).",
        "Produse: Impact Advanced Recovery (Nestle) — INCERT disponibilitate RO.",
        "Alternativă: Cubitan (Nutricia).",
    ])

    add_heading_custom(doc, "Produse ONS disponibile România", level=3)
    rows = [
        ["Fortimel Extra", "Nutricia", "eMAG, Farmacia Tei", "~80 RON/4×200ml"],
        ["Fortimel Compact Protein", "Nutricia", "eMAG", "~140-150 RON/4 pack"],
        ["Supportan Drink", "Fresenius Kabi", "Farmacia Tei", "~149 RON/4×200ml"],
        ["Nutridrink", "Nutricia", "Lanțuri mari", "70-90 RON/4×200ml"],
        ["Diasip (diabetici)", "Nutricia", "Farmacia Tei, Dr.Max", "Similar Nutridrink"],
    ]
    add_table(doc, ["Produs", "Producător", "Disponibilitate", "Preț"], rows)

    add_heading_custom(doc, "Dieta post-esofagectomie", level=3)
    add_bullet(doc, [
        "5-6 mese mici/zi, porții 150-200 ml.",
        "Consistență progresivă: lichid → pasat → semi-solid → solid moale.",
        "Mestecă încet. NU lichide în timpul mesei.",
        "Șezut drept 30-45 min după masă.",
        "Evită: acid, picant, grăsimi saturate, cofeină, alcool, ciocolată, carbogazoase.",
        "Suplimentare: B12, fier, calciu + vit D.",
        "Evită dumping: carbohidrați simpli mari pe stomac gol.",
    ])

    add_heading_custom(doc, "Suport psihologic", level=2)
    rows = [
        ["Lidia Stoica Psiho-Oncologie", "Str. Esarfei 55, Sector 3 București", "0720 655 618"],
        ["Comunitatea Oncologică RO", "București + online", "comunitateaoncologica.ro"],
        ["Sus Inima Centru București", "Str. Mircea Vulcănescu 79, Sector 1", "+40 734 157 442"],
        ["Oncohelp Timișoara", "Str. Porumbescu 57-59", "0256 495 403"],
        ["P.A.V.E.L. TelVerde", "Bucuresti", "0800 800 421 (gratuit)"],
    ]
    add_table(doc, ["Organizație", "Locație", "Contact"], rows)

    add_heading_custom(doc, "Paliație — HOSPICE Casa Speranței", level=2)
    add_bullet(doc, [
        "HQ Brașov: Str. Ghioceilor 1, Făgăraș — +40 372 577 900.",
        "Sediu București (Centrul de Resurse): Calea Călărașilor 109 — 021 326 3771.",
        "Servicii GRATUITE: domiciliu, centre zi, internare.",
        "Local Arad: Str. Păcurarilor 8 (18 locuri); Dr. Crainic Daniel, Str. Poetului bl. Z7.",
        "CNAS deconteaza 90 zile/an îngrijire la domiciliu pentru adulți.",
    ])
    doc.add_page_break()

    # ============================================================
    # XX. LOGISTICA BUCURESTI
    # ============================================================
    add_heading_custom(doc, "XX. Logistică tratament în București — cazare gratuită", level=1, color=BLUE)

    add_heading_custom(doc, "Transport Arad-București", level=2)
    rows = [
        ["Mașina personală A1", "8-9h cu pauze", "150-200 RON motorină", "Cel mai confortabil"],
        ["Tren CFR", "9-13h", "150-300 RON cls II", "Prelungit"],
        ["Avion Timișoara-Otopeni", "1h 10min zbor + transferuri", "150-400 RON", "Rapid dar cu drum TM + transfer OTP"],
        ["Autocar", "11-13h", "120-180 RON", "Confort limitat"],
    ]
    add_table(doc, ["Mod", "Durată", "Preț", "Comentariu"], rows)

    add_heading_custom(doc, "Cazare GRATUITĂ — Casa Sus Inima București", level=2)
    add_bullet(doc, [
        "Str. Luigi Cazzavillan 44, Sector 1.",
        "Telefon: +40 734 157 442 (WhatsApp).",
        "Email: contact@susinima.eu.",
        "16 paturi, 3 băi, 2 bucătării, curte.",
        "Pacienți oncologici (focus radioterapie; chimio — verifică direct).",
        "Acompaniator inclus. GRATUIT.",
    ])

    add_heading_custom(doc, "Cazare privată aproape de IOB Fundeni", level=2)
    rows = [
        ["Casa Speranța Fundeni", "Șos. Fundeni 147A", "100 RON/noapte", "Vizavi IOB"],
        ["Casa Ghețu", "Șos. Fundeni 141", "-", "100m IOB"],
        ["Hoteluri 3*-4*", "Zona Fundeni", "300-500 RON/noapte", "Booking.com"],
    ]
    add_table(doc, ["Unitate", "Adresă", "Preț", "Distanță"], rows)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Cost 6 luni tratament:").bold = True
    add_bullet(doc, [
        "La Sus Inima GRATUIT: 0-3000 RON (mâncare, farmacie, transport local).",
        "Cazare privată: 10.000-30.000 RON.",
    ])
    doc.add_page_break()

    # ============================================================
    # XXI. INTREBARI
    # ============================================================
    add_heading_custom(doc, "XXI. Întrebări concrete de pus fiecărui specialist", level=1, color=BLUE)

    add_heading_custom(doc, "Pentru gastroenterolog (Dr. Noufal)", level=2)
    add_bullet(doc, [
        "Poate fi emis biletul CT cu „URGENT ONCOLOGIC”?",
        "Pot obține rezultatul biopsiei mai rapid prin contact direct cu laboratorul?",
        "Ce marcaj de ambiguitate există în aspectul endoscopic?",
        "Dacă T1a/precanceros, faci ESD local sau trimiți la Cluj/București?",
    ])

    add_heading_custom(doc, "Pentru radiolog (după CT)", level=2)
    add_bullet(doc, [
        "Ce dimensiune are leziunea (mm)?",
        "Ce straturi esofagiene sunt invadate (T clinic)?",
        "Câți ganglioni locoregionali sunt suspecți?",
        "Leziuni suspecte în ficat / plămâni / oase / suprarenale?",
        "Necesar PET-CT pentru clarificare?",
        "Raport în format DICOM (CD)?",
    ])

    add_heading_custom(doc, "Pentru patolog (după biopsie)", level=2)
    add_bullet(doc, [
        "Tip celular (scuamos / adenocarcinom)?",
        "Grad diferențiere (G1/G2/G3)?",
        "Invazie vasculară sau perineurală?",
        "PD-L1 CPS — obligatoriu pentru imunoterapie.",
        "HER2 (IHC + FISH dacă 2+).",
        "MSI / MMR.",
        "Dacă adenocarcinom: CLDN18.2 IHC.",
        "A 2-a opinie patologică la IOB Fundeni?",
    ])

    add_heading_custom(doc, "Pentru oncolog (primul consult)", level=2)
    add_bullet(doc, [
        "Stadiu TNM exact?",
        "Protocol neoadjuvant — CROSS sau FLOT? De ce?",
        "Adenocarcinom: de ce nu FLOT (ESOPEC 2024)?",
        "Imunoterapie adjuvantă post-operator (nivolumab CheckMate-577)?",
        "Biomarkeri gata pe biopsie? Ce lipsește?",
        "Trial-uri clinice disponibile pentru stadiul meu?",
        "Indicația e pe lista CNAS mai 2025 pentru pembro/nivo?",
        "Al 2-lea consult la IOB Fundeni înainte de decizia finală?",
    ])

    add_heading_custom(doc, "Pentru chirurg", level=2)
    add_bullet(doc, [
        "Tehnică recomandată — Ivor Lewis / McKeown / MIE / RAMIE?",
        "Câte esofagectomii faci pe an în acest centru?",
        "Rata mortalitate 30 zile și 90 zile?",
        "Rata anastomotic leak?",
        "Jejunostomă intraoperator profilactic?",
        "Durată internare?",
        "Când se poate reîncepe dieta normală?",
    ])

    add_heading_custom(doc, "Pentru cardiolog (OBLIGATORIU la stent)", level=2)
    add_bullet(doc, [
        "Funcția ventricul stâng (FEVS)?",
        "Stentul 2012 încă funcțional (angiografie)?",
        "Pot tolera operație majoră + anestezie generală?",
        "Pot tolera FLOT (5-FU risc spasm coronarian)?",
        "Pauză perichirurgicală Triplixam — când și cum?",
        "Aspenter perioperator — DA/NU?",
        "Troponină + ECG seriale sub chimio?",
    ])

    add_heading_custom(doc, "Pentru nutriționist", level=2)
    add_bullet(doc, [
        "Scor NRS-2002 / MUST?",
        "Aport caloric + proteic țintă zilnic?",
        "Imunonutriție Impact/Cubitan — de unde procur?",
        "ONS rambursate prin CAS?",
        "Plan dietetic post-esofagectomie?",
    ])

    add_heading_custom(doc, "Pentru DGASPC Arad", level=2)
    add_bullet(doc, [
        "Grad handicap eligibil în stadiul meu?",
        "Documente exacte pentru dosar?",
        "Timp procesare 2026?",
        "Indemnizație însoțitor — cumul cu pensie limită vârstă?",
    ])
    doc.add_page_break()

    # ============================================================
    # XXII. PLAN ACTIUNE
    # ============================================================
    add_heading_custom(doc, "XXII. Plan de acțiune — săptămâna 1, 2, 3 și următoarele", level=1, color=BLUE)

    add_heading_custom(doc, "Săptămâna 1 (19.04 → 26.04.2026)", level=2)
    add_bullet(doc, [
        "19.04 DUMINICĂ: hidratare activă pre-CT (2-3L apă + electroliți).",
        "20.04 LUNI 17:00: CT torace+abdomen+pelvis cu contrast la Genesis Micălaca.",
        "20.04: cere raportul în DICOM + mențiune urgent oncologic.",
        "21.04: consult cardiologic preventiv (funcție VS pre-posibil-tratament).",
        "22.04: reluare Jamesi (H+48 post-CT) DOAR după creatinină normală.",
        "22-23.04: așteaptă rezultatul CT scris.",
        "Monitor automat biopsie activ — notificare pe telefon la rezultat (estimat 24.04-01.05).",
    ])

    add_heading_custom(doc, "Săptămâna 2 (27.04 → 03.05.2026)", level=2)
    add_bullet(doc, [
        "Citește CT cu Dr. Noufal sau oncolog local.",
        "Scanează și salvează rezultatul CT în Dosar_Medical/.",
        "Dacă biopsia a venit: citire cu oncolog, cere retestare biomarkeri.",
        "Programare consult oncolog digestiv — SCJU Arad / Oncohelp Timișoara.",
        "Paralel: contact email Ponderas (office.ponderas@reginamaria.ro) pentru 2-lea consult.",
    ])

    add_heading_custom(doc, "Săptămâna 3 (04.05 → 10.05.2026)", level=2)
    add_bullet(doc, [
        "Consult oncolog digestiv primar.",
        "Discuție tumor board multidisciplinar.",
        "Al 2-lea consult la IOB Fundeni SAU Ponderas.",
        "Dacă scenariu rezonabil: considera S2 la Charité/Heidelberg/Gustave Roussy pentru al 2-lea consult extern.",
    ])

    add_heading_custom(doc, "Săptămânile 4+ (11.05+)", level=2)
    add_bullet(doc, [
        "Începere tratament (neoadjuvant CROSS/FLOT dacă operabil; chimio+imunoterapie dacă metastatic).",
        "Dosar DGASPC Arad pentru grad handicap (paralel cu tratamentul).",
        "Aplicare formular S2 la CASARAD DACĂ tratament UE decis.",
        "Contact Sus Inima pentru cazare București dacă necesar.",
    ])

    doc.add_paragraph()
    add_callout(
        doc,
        "Principiu-cheie pentru familie",
        "Nu tot trebuie făcut săptămâna aceasta. Tratamentul oncologic modern are cadran propriu: "
        "2-4 săptămâni între diagnostic și început tratament, timp în care se face stadializarea "
        "completă, testare biomarkeri, consult multidisciplinar. Urgența e clinică, nu psihologică.\n\n"
        "Prioritate: CT + biopsie → consult oncolog → al 2-lea consult → decizie = ~3 săptămâni.\n"
        "Tratament = 6-18 luni (depinde de stadiu).",
        color=GREEN,
    )
    doc.add_page_break()

    # ============================================================
    # XXIII. CE NU AM GASIT
    # ============================================================
    add_heading_custom(doc, "XXIII. Ce NU am găsit (transparență)", level=1, color=GRAY)

    doc.add_paragraph(
        "Conform Regulii 17, listez aici informațiile căutate dar negăsite sau incerte. Sunt "
        "întrebări legitime pentru echipa medicală."
    )

    add_bullet(doc, [
        ("NEGASIT", "Volume anuale esofagectomii per centru în România — nepublicate."),
        ("NEGASIT", "Lista exactă 2026 imunoterapii compensate PNO pentru esofag — site CNAS nu publică lista structurată."),
        ("NEGASIT", "Timp mediu aprobare PNO 2026 pentru imunoterapie esofag."),
        ("NEGASIT", "Disponibilitate zolbetuximab în PNO RO 2026."),
        ("NEGASIT", "Preț exact proton therapy Praga/Viena cu S2 — verificare directă."),
        ("INCERT", "Dana-Farber Online Second Opinion pentru pacienți români — email obligatoriu înainte de plată."),
        ("INCERT", "Aplicații mobile RO dedicate pacienți oncologici."),
        ("INCERT", "Grupuri suport specifice cancer esofagian în RO — cele existente sunt generaliste."),
        ("INCERT", "Acoperire HOSPICE în Nădlac/Arad."),
        ("INCERT", "Toți biomarkerii pot fi testați pe biopsia deja prelevată (volumul material)."),
    ])
    doc.add_page_break()

    # ============================================================
    # XXIV. SURSE
    # ============================================================
    add_heading_custom(doc, "XXIV. Surse citate", level=1, color=GRAY)
    p = doc.add_paragraph("Toate URL-urile accesate pe 19.04.2026.")
    p.italic = True

    add_heading_custom(doc, "Ghiduri clinice", level=2)
    add_bullet(doc, [
        "NCCN Esophageal V1.2025 — nccn.org/guidelines",
        "ESMO Clinical Practice Guidelines 2022 — annalsofoncology.org",
        "AJCC Cancer Staging Manual 8th Ed (2017)",
        "ESPEN Clinical Nutrition in Cancer 2021 — PMID 33946039",
    ])

    add_heading_custom(doc, "Studii pivot", level=2)
    add_bullet(doc, [
        "CROSS — van Hagen NEJM 2012",
        "FLOT4 — Al-Batran Lancet 2019",
        "ESOPEC — NEJM 2024",
        "CheckMate-577 — Kelly NEJM 2021",
        "KEYNOTE-590 — Sun Lancet 2021",
        "CheckMate-648 — Doki NEJM 2022",
        "KEYNOTE-811 — Janjigian Nature 2024",
        "SPOTLIGHT — zolbetuximab Lancet 2023",
        "DESTINY-Gastric01 — Shitara NEJM 2020",
        "MATTERHORN — NEJM 2025",
        "MIRO — Mariette NEJM 2019",
        "ROBOT — van der Sluis Ann Surg 2019",
    ])

    add_heading_custom(doc, "Centre + CNAS + asociații", level=2)
    add_bullet(doc, [
        "IOB — iob.ro",
        "ICF — icfundeni.ro",
        "IOCN — iocn.ro",
        "Colțea — coltea.ro",
        "Ponderas — reginamaria.ro/ponderas",
        "SANADOR — sanador.ro",
        "OncoHelp — oncohelp.ro",
        "WPK Viena — wiener-privatklinik.com",
        "Charité — charite.de",
        "Heidelberg — heidelberg-university-hospital.com",
        "Gustave Roussy — gustaveroussy.fr",
        "PNO 2024 — cnas.ro",
        "Formular S2 — cnas.ro/state-membre-ue",
        "FABC — fabc.ro",
        "Sus Inima — susinima.eu",
        "HOSPICE — hospice.ro",
    ])

    doc.add_paragraph()

    # ============================================================
    # FINAL
    # ============================================================
    add_heading_custom(doc, "Control versiune", level=1, color=GRAY)
    add_bullet(doc, [
        "Autor: Claude Opus 4.7 (1M context)",
        "Data generare: 2026-04-19",
        "Versiune raport: 1.0",
        "Sursa: 4 agenți de cercetare paralele (tratament, trial-uri, centre, suport)",
        "Marcaje aplicate: [CERT], [PROBABIL], [INCERT], [NEGASIT], [RECOMANDAT]",
        "Conformitate: Regula 17 CLAUDE.md",
    ])

    add_callout(
        doc,
        "ATENȚIONARE FINALĂ",
        "Acest document este compilație de informații medicale din surse primare (NCCN, ESMO, "
        "EMA, studii peer-reviewed, documente CNAS, clinicaltrials.gov). NU constituie consult "
        "medical și NU înlocuiește decizia oncologului curant.\n\n"
        "Stadializarea definitivă depinde de biopsia (rezultat în lucru) + CT 20.04.2026 + "
        "eventuale investigații suplimentare + consult multidisciplinar. Orice decizie "
        "terapeutică se ia de echipa oncologică cu acces la dosarul complet.\n\n"
        "În caz de urgență medicală: 112.",
        color=RED,
    )

    return doc


def main():
    doc = build_document()
    doc = build_document_part2(doc)
    doc = build_document_part3(doc)
    doc = build_document_part4(doc)
    doc = build_document_part5(doc)

    output_path = Path(__file__).parent / "2026-04-19_ghid_cancer_esofagian_complet.docx"
    doc.save(str(output_path))
    print(f"Document salvat: {output_path}")
    print(f"Dimensiune: {output_path.stat().st_size:,} bytes")
    return output_path


if __name__ == "__main__":
    main()
