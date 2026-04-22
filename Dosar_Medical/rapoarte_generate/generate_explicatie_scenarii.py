"""
Generator DOCX pentru Explicatia Consult Oncolog + Scenarii Ascita/Biopsie
Pacient: Petrila Viorel-Mihai
Data generare: 22 aprilie 2026

Script-as-source-of-truth: genereaza DOCX-ul din acest fisier Python.
Documentul Markdown corespunzator:
  Documente_Informative/EXPLICATIE_CONSULT_ONCOLOG_SCENARII.md

Output:
  2026-04-22_explicatie_consult_oncolog_scenarii.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# =========================================================================
# PALETA CULORI
# =========================================================================

COLOR_PRIMARY = RGBColor(0x1E, 0x40, 0xAF)      # Albastru medical
COLOR_SECONDARY = RGBColor(0x64, 0x74, 0x8B)    # Gri-albastru
COLOR_TEXT = RGBColor(0x1E, 0x29, 0x3B)         # Negru-albastru
COLOR_MUTED = RGBColor(0x94, 0xA3, 0xB8)        # Gri
COLOR_OK = RGBColor(0x16, 0xA3, 0x4A)           # Verde
COLOR_WARN = RGBColor(0xD9, 0x77, 0x06)         # Portocaliu
COLOR_CRIT = RGBColor(0xDC, 0x26, 0x26)         # Rosu
COLOR_INFO = RGBColor(0x08, 0x91, 0xB2)         # Cyan
COLOR_QUOTE_BG = "E0F2FE"                        # Cyan deschis (hex string pt XML)
COLOR_CRIT_BG = "FEE2E2"
COLOR_WARN_BG = "FEF3C7"
COLOR_OK_BG = "DCFCE7"
COLOR_INFO_BG = "CFFAFE"
COLOR_GRAY_BG = "F1F5F9"
COLOR_HEADER_BG = "1E40AF"                       # Albastru inchis pt header celula

# =========================================================================
# HELPERS STILIZARE
# =========================================================================


def set_cell_bg(cell, color_hex):
    """Seteaza background color pentru o celula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def set_paragraph_bg(paragraph, color_hex):
    """Seteaza background color pentru un paragraf (callout)."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    p_pr.append(shd)


def add_border_left(paragraph, color_hex, width_pt=6):
    """Adauga un border stanga unui paragraf (callout style)."""
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_pt * 4))  # size in 1/8 pt
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), color_hex)
    pbdr.append(left)
    p_pr.append(pbdr)


def add_heading_bar(doc, text, level=1, color=COLOR_PRIMARY, bg_hex=None):
    """Heading custom cu posibil background."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(22)
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        run.font.size = Pt(16)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        run.font.size = Pt(13)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    run.font.color.rgb = color
    if bg_hex:
        set_paragraph_bg(p, bg_hex)
    return p


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=COLOR_TEXT, alignment=None):
    """Paragraf normal cu formatare."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def add_mixed_paragraph(doc, segments, alignment=None, space_after=6):
    """
    Paragraf cu segmente multiple (fiecare poate avea formatare proprie).
    segments: list of dicts {text, bold, italic, color, size}
    """
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    for seg in segments:
        run = p.add_run(seg["text"])
        run.bold = seg.get("bold", False)
        run.italic = seg.get("italic", False)
        run.font.size = Pt(seg.get("size", 11))
        run.font.color.rgb = seg.get("color", COLOR_TEXT)
    return p


def add_bullet(doc, text, level=0, bold_prefix=None):
    """Bullet list cu posibil prefix bold."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(" " + text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def add_numbered(doc, text, bold_prefix=None):
    """Numbered list."""
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(" " + text)
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p


def add_quote(doc, text, color_left=COLOR_INFO, bg_hex=COLOR_INFO_BG):
    """Citat cu border stanga + fundal (pentru analogii)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    add_border_left(p, "{:02X}{:02X}{:02X}".format(*color_left), width_pt=3)
    set_paragraph_bg(p, bg_hex)
    return p


def add_callout(doc, label, text, color=COLOR_INFO, bg_hex=COLOR_INFO_BG):
    """Callout colorat cu label bold."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = COLOR_TEXT
    add_border_left(p, "{:02X}{:02X}{:02X}".format(*color), width_pt=3)
    set_paragraph_bg(p, bg_hex)
    return p


def add_table(doc, headers, rows, header_bg=COLOR_HEADER_BG, col_widths=None):
    """
    Tabel formatat cu antet colorat.
    headers: list of strings
    rows: list of lists of strings
    col_widths: list of Inches
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False

    # Header
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            # Zebra
            if r_idx % 2 == 1:
                set_cell_bg(cell, COLOR_GRAY_BG)

    # Col widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # Space after
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    return table


def add_separator(doc):
    """Linie de separare subtila."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("━" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break()


def configure_document(doc):
    """Seteaza margini, font default, etc."""
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)

    # Font default
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_TEXT


# =========================================================================
# CONSTRUIRE DOCUMENT
# =========================================================================

def build_document():
    doc = Document()
    configure_document(doc)

    # ============================================================
    # COVER PAGE
    # ============================================================

    cover_title = doc.add_paragraph()
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_title.paragraph_format.space_before = Pt(72)
    r = cover_title.add_run("EXPLICAȚIE CONSULT ONCOLOG URGENT")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = COLOR_PRIMARY

    cover_subtitle = doc.add_paragraph()
    cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover_subtitle.add_run("Cele 4 Scenarii Ascită vs. Biopsie — ghid complet pentru familie")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_SECONDARY
    cover_subtitle.paragraph_format.space_after = Pt(72)

    # Box cu datele pacientului
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, "F1F5F9")
    cell.text = ""
    for line in [
        ("Pacient:", "Petrilă Viorel-Mihai (66 ani)"),
        ("CNP:", "1590518024486"),
        ("Diagnostic suspectat:", "Proces proliferativ esofagian circumferențial nedepășibil endoscopic"),
        ("Clasificare imagistică:", "Siewert II probabil (T3-T4, N0-N1, M0 probabil — cu ascită de evaluat)"),
        ("Investigații de bază:", "Endoscopie 17.04.2026 + CT TAP 20.04.2026"),
        ("Document creat:", "22 aprilie 2026"),
        ("Autor:", "Roland Petrilă (fiul pacientului) + Claude Code"),
    ]:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(line[0] + " ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(line[1])
        r2.font.size = Pt(11)
    # sterge primul paragraf gol al celulei
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    disclaimer_p = doc.add_paragraph()
    disclaimer_p.paragraph_format.space_before = Pt(36)
    disclaimer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = disclaimer_p.add_run(
        "NU înlocuiește consultul medical. Material explicativ pentru înțelegere. "
        "Deciziile terapeutice aparțin exclusiv oncologului digestiv."
    )
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = COLOR_MUTED

    add_page_break(doc)

    # ============================================================
    # CUPRINS
    # ============================================================

    add_heading_bar(doc, "CUPRINS", level=1, color=COLOR_PRIMARY)
    toc_items = [
        "1. Povestea — casa veche cu probleme",
        "2. Ascita — ce e de fapt apa din subsol",
        "3. Întrebările tale — răspunsuri directe (4 întrebări)",
        "4. Investigațiile posibile — explicate simplu",
        "5. Markerii moleculari — cheile care deschid tratamente",
        "6. Cele 4 scenarii combinatorii cu plan de tratament",
        "7. Protocolul FLOT — explicat în detaliu",
        "8. Imunoterapia — o categorie nouă de tratament",
        "9. Nutriția — cum trebuie să mănânce tata acum",
        "10. Semnale de alarmă — când suni 112, când aștepți",
        "11. Întrebări frecvente ale familiei",
        "12. Tabel sumarizat — 4 scenarii + plan",
        "13. Ce trebuie să faci TU (Roland) — concret",
        "14. Timeline vizual — cronologia realistă",
        "15. Mesajul principal — în 30 secunde",
        "16. Glossar — termeni medicali explicați simplu",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item.split(".", 1)[1].strip())
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(4)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 1 — POVESTEA
    # ============================================================

    add_heading_bar(doc, "1. Povestea — casa veche cu probleme", level=1, color=COLOR_PRIMARY)

    add_paragraph(doc,
        "Tatăl tău e ca o casă bine construită, care a rezistat 66 de ani. "
        "În ultima lună au apărut semne că ceva nu e în regulă pe interior.")

    add_mixed_paragraph(doc, [
        {"text": "Endoscopia", "bold": True},
        {"text": " a fost ca și cum ai trimis un meșter cu o lanternă pe hol — a găsit o zidărie dăunată în coridorul principal (esofagul) care blochează pasajul. Meșterul a luat o bucățică din zid ("},
        {"text": "biopsia", "bold": True},
        {"text": ") și a trimis-o la laborator ca să analizeze exact ce material e afectat — mucegai, igrasie, carii, cărămidă deteriorată? Rezultatul vine în 1–2 săptămâni."},
    ])

    add_mixed_paragraph(doc, [
        {"text": "Între timp, pentru a vedea dacă problema e doar în acel coridor sau dacă s-a răspândit, ai chemat un "},
        {"text": "expert cu dronă", "bold": True},
        {"text": " (CT-ul) care a zburat prin toată casa și a filmat totul — de la acoperiș la subsol, inclusiv pereți, instalații, fundație."},
    ])

    add_heading_bar(doc, "Ce a văzut drona — raportul CT (20.04.2026)", level=2)

    add_paragraph(doc, "Drona a raportat 3 lucruri importante:", bold=True)

    add_callout(doc, "Vești bune", "acoperișul, fundația, pereții altor camere sunt OK (M0 probabil). Ganglionii limfatici sub pragul de alarmă (max 7,5 mm vs. 10 mm). Sistemele vitale (inimă, rinichi) funcționează normal.", color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_callout(doc, "Vești de urmărit", "zidăria dăunată e mai extinsă decât părea — cuprinde tot coridorul circumferențial și intră puțin în sufragerie (joncțiunea stomacului). Clasificare Siewert II.", color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    add_callout(doc, "Semnalul care schimbă totul", "pungi cu apă în subsol (ascită) — 15 mm în jurul ficatului + 28 mm în pelvis. Acest semnal trebuie elucidat URGENT.", color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 2 — ASCITA
    # ============================================================

    add_heading_bar(doc, "2. Ascita — ce e de fapt apa din subsol", level=1, color=COLOR_PRIMARY)

    add_heading_bar(doc, "2.1 Cum arată un corp sănătos", level=2)
    add_paragraph(doc, "Imaginează-ți burta ca pe un borcan mare care conține:")
    add_bullet(doc, "Ficatul (ca un burete de curățat), stomacul, intestinele, rinichii, pancreasul")
    add_bullet(doc, "Toate organele sunt împachetate într-o folie transparentă numită peritoneu — ca celofanul de la brânză")
    add_bullet(doc, "Între folie și peretele burții e un strat foarte subțire de lichid lubrifiant (~50–100 ml, nu-l simți)")

    add_heading_bar(doc, "2.2 Ce e ascita", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Ascita = ", "bold": True},
        {"text": "acumulare anormală de lichid în burtă, peste cantitatea normală de lubrifiere. E ca și cum ai avea un acvariu uscat în mod normal, și deodată apare apă pe fundul lui. Nu mult, dar suficient cât să o vezi."},
    ])
    add_paragraph(doc,
        "În cazul tatei: 15 mm în jurul ficatului + 28 mm în pelvis = cantitate ușoară–moderată. "
        "Nu e masivă (ascita masivă poate însemna 10+ litri), dar e semnificativă pentru diagnostic.")

    add_heading_bar(doc, "2.3 Cele 4 cauze posibile", level=2)

    # Cauza 1
    add_heading_bar(doc, "Cauza 1 — „Conductă care picură în subsol\" (inflamație locală)", level=3, color=COLOR_WARN)
    add_quote(doc, "Analogie: ai o conductă veche din subsol care începe să picure din cauza vibrațiilor de la zidăria prăbușită din coridorul principal. Apa e curată, doar apă.", color_left=COLOR_WARN, bg_hex=COLOR_WARN_BG)
    add_mixed_paragraph(doc, [
        {"text": "Medical: ", "bold": True},
        {"text": "tumora esofagiană avansată irită peritoneul din apropiere, provoacă inflamație locală, și peritoneul secretă mai mult lichid ca reacție. "},
        {"text": "Nu e cancer răspândit, doar reacție.", "bold": True},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Cum recunoaște laboratorul: ", "bold": True},
        {"text": "puncție + analiză = "},
        {"text": "fără celule tumorale", "bold": True},
        {"text": ", eventual leucocite inflamatorii."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Impact clinic: ", "bold": True},
        {"text": "bun. Tratamentul principal (chimio + operație) rezolvă cauza — tumora —, iar apa dispare."},
    ])

    # Cauza 2
    add_heading_bar(doc, "Cauza 2 — „Filtre înfundate\" (hipoalbuminemie)", level=3, color=COLOR_WARN)
    add_quote(doc, "Analogie: sistemul de filtrare al casei (ficatul) sau pompa care menține presiunea (rinichii) nu mai funcționează bine. Proteinele care „țin apa\" în vase de sânge scad, și apa se scurge în burtă.", color_left=COLOR_WARN, bg_hex=COLOR_WARN_BG)
    add_mixed_paragraph(doc, [
        {"text": "Medical: ", "bold": True},
        {"text": "la pacienții oncologici cu apetit scăzut și inflamație cronică, nivelul albuminei poate scădea. Sub 3 g/dL, apa începe să migreze din vasele de sânge în țesuturi și cavități."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Impact clinic: ", "bold": True},
        {"text": "se tratează cu nutriție intensivă + eventual perfuzie cu albumină. Nu schimbă planul oncologic."},
    ])

    # Cauza 3
    add_heading_bar(doc, "Cauza 3 — „Drenajul blocat\" (obstrucție venoasă / limfatică)", level=3, color=COLOR_WARN)
    add_quote(doc, "Analogie: scurgerea de la subsol e blocată pentru că tumora apasă pe conducta principală de evacuare. Apa vine normal, dar nu mai iese.", color_left=COLOR_WARN, bg_hex=COLOR_WARN_BG)
    add_mixed_paragraph(doc, [
        {"text": "Medical: ", "bold": True},
        {"text": "tumora sau ganglionii măriți comprimă vena portă sau vasele limfatice din jur. Lichidul se acumulează în amonte de blocaj."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Impact clinic: ", "bold": True},
        {"text": "intermediar. Dispare dacă tratamentul reduce tumora, poate reveni dacă tumora nu răspunde."},
    ])

    # Cauza 4
    add_heading_bar(doc, "Cauza 4 — Mucegai invizibil pe toți pereții subsolului (carcinomatoză peritoneală)", level=3, color=COLOR_CRIT)
    add_quote(doc, "Analogie: mucegaiul din coridorul principal s-a răspândit prin spori în aer, și acum e pe toți pereții subsolului, invizibil cu ochiul liber, dar prezent. Apa conține aceste „spori\".", color_left=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)
    add_mixed_paragraph(doc, [
        {"text": "Medical: ", "bold": True},
        {"text": "celulele tumorale au migrat de la tumora primară prin cavitatea peritoneală și s-au implantat pe folia-peritoneu. Fiecare mic focar tumoral secretă lichid și face peritoneul să „plângă\" în plus."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Cum recunoaște laboratorul: ", "bold": True},
        {"text": "puncție + analiză citologică = celule tumorale vizibile la microscop."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Impact clinic: ", "bold": True, "color": COLOR_CRIT},
        {"text": "major. Transformă stadiul din III (curabil) în IV (paliativ). Protocolul terapeutic se schimbă radical.", "color": COLOR_CRIT},
    ])

    add_heading_bar(doc, "2.4 De ce 2 locații (perihepatic + intrapelvin)?", level=2)
    add_paragraph(doc, "Apa, prin gravitație, se adună în cele 2 „găleți\" naturale ale burții:")
    add_numbered(doc, "Spațiul perihepatic (în jurul ficatului) — găleata de sus, unde apa se colectează când stai în picioare")
    add_numbered(doc, "Fundul pelvisului (în bazin) — găleata de jos, unde apa se colectează când stai întins")
    add_paragraph(doc,
        "În cazul tatei, ambele locuri au apă simultan = cantitate suficientă încât ambele „găleți\" se umplu. "
        "NU indică direcția sursă (apa se distribuie uniform prin gravitație), doar confirmă că există ascită reală, nu e artefact CT.")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 3 — INTREBARI
    # ============================================================

    add_heading_bar(doc, "3. Întrebările tale — răspunsuri directe", level=1, color=COLOR_PRIMARY)

    # Q1
    add_heading_bar(doc, "3.1 Trebuie să mergem la oncolog ACUM, înainte de rezultatul biopsiei?", level=2)
    add_mixed_paragraph(doc, [
        {"text": "DA.", "bold": True, "size": 13, "color": COLOR_CRIT},
        {"text": " Și sunt 2 motive separate care amândouă o fac urgentă:"},
    ])
    add_callout(doc, "Motiv 1 — timing", "programarea la un oncolog bun ia 2–3 săptămâni oricum. Dacă aștepți biopsia (1 mai) + programarea, ajungi la tratament pe ~1 iunie. Prea târziu — tumora nu așteaptă.", color=COLOR_WARN, bg_hex=COLOR_WARN_BG)
    add_callout(doc, "Motiv 2 — coordonarea investigațiilor", "nimeni altcineva nu poate decide ce să facă cu apa din subsol. Doar oncologul coordonează paracenteza, laparoscopia, PET-CT.", color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    # Q2
    add_heading_bar(doc, "3.2 Ascita poate fi cancerigenă indiferent de unde a apărut?", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Răspuns simplu: ", "bold": True},
        {"text": "DA, teoretic. Dar în contextul tatei, sursa cea mai probabilă tot tumora eso-gastrică rămâne."},
    ])
    add_callout(doc, "[CERT]", "Carcinomatoza peritoneală poate veni din orice tumoră malignă internă care ajunge la peritoneu.", color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_paragraph(doc, "Surse posibile cu probabilitate la tatăl tău:", bold=True)
    add_table(doc,
        headers=["Sursă posibilă", "Probabilitate"],
        rows=[
            ["Cancer gastric", "HIGH — e practic aceeași zonă (Siewert II include stomacul)"],
            ["Cancer esofagian avansat", "HIGH — e diagnosticul suspectat principal"],
            ["Cancer pancreatic", "MEDIU — posibil fără simptome clare încă"],
            ["Cancer colorectal", "SCĂZUT — colonoscopia din 17.04 a arătat doar polip benign"],
            ["Cancer ovarian", "NU SE APLICĂ — tata e bărbat"],
            ["Mezoteliom peritoneal primar", "FOARTE RAR — dar posibil cu istoric de fumat"],
            ["Cancer biliar / hepatic", "SCĂZUT — ficatul apare normal la CT"],
            ["Cancer prostatic", "SCĂZUT — prostata apare normală la CT"],
        ],
        col_widths=[Inches(1.8), Inches(4.5)])

    add_paragraph(doc, "Concluzia practică:", bold=True)
    add_paragraph(doc,
        "Dacă ascita e malignă, în 85–95% din cazuri în contextul lui e legată de tumora eso-gastrică. "
        "Dar oncologul trebuie să CONFIRME asta — nu poate presupune.")

    # Q3
    add_heading_bar(doc, "3.3 Ce poate constata oncologul FĂRĂ analize suplimentare?", level=2)
    add_mixed_paragraph(doc, [
        {"text": "[PROBABIL] Ce POATE spune la primul consult:", "bold": True, "color": COLOR_OK},
    ])
    for item in [
        ("Clasificarea Siewert", "confirmă dacă e I, II sau III. Afectează tipul de chirurgie."),
        ("Stadializare CLINICĂ preliminară", "T + N estimat, cu marcaj „necesită confirmare histologică și M\"."),
        ("Risc operator aproximativ", "pe baza comorbidităților (stent cardiac, diabet, hernie operată)."),
        ("Protocol terapeutic preliminar", "FLOT vs. CROSS vs. alt protocol."),
        ("Ce investigații sunt necesare", "paracenteză? PET-CT? laparoscopie?"),
        ("Urgența tratamentului", "săptămâni (urgent) vs. o lună."),
        ("Intervenții paliative", "stent? jejunostomă? prehabilitare nutrițională?"),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0] + " —")

    add_mixed_paragraph(doc, [
        {"text": "[NEGASIT] Ce NU POATE spune la primul consult:", "bold": True, "color": COLOR_MUTED},
    ])
    for item in [
        ("Tipul histologic exact", "adenocarcinom vs. scuamos → obligatoriu biopsia"),
        ("Gradul de diferențiere", "bine / moderat / slab → obligatoriu biopsia"),
        ("Markerii moleculari", "HER2, PD-L1, MSI, claudin-18.2 → biopsia + IHC"),
        ("Natura ascitei", "reactivă vs. carcinomatoză → obligatoriu paracenteză"),
        ("Stadializarea M definitivă", "fără PET-CT, CT e doar „M0 probabil\""),
        ("Decizia finală pe protocol", "se cimentează doar cu toate datele"),
    ]:
        add_bullet(doc, item[1], bold_prefix=item[0] + " —")

    # Q4
    add_heading_bar(doc, "3.4 E obligatoriu oncologului să facă o altă serie de analize?", level=2)
    add_mixed_paragraph(doc, [
        {"text": "DA, în practică — inevitabil.", "bold": True, "size": 13, "color": COLOR_CRIT},
        {"text": " Oncologul NU poate începe tratament fără:"},
    ])

    add_paragraph(doc, "Analize minime obligatorii pre-tratament:", bold=True)
    add_table(doc,
        headers=["Investigație", "De ce", "Timp aproximativ"],
        rows=[
            ["Hemoleucogramă completă", "Baseline pre-chimio — leucocite, trombocite, hemoglobină", "1 zi"],
            ["Biochimie completă", "Funcție hepatică, renală, glicemie", "1 zi"],
            ["Markeri tumorali serici", "CEA, CA 19-9, CA 72-4 — urmărire răspuns", "1–3 zile"],
            ["ECG + ECHO cardiac", "Evaluare funcție cardiacă (pacient post-stent)", "3–7 zile"],
            ["Paracenteză + citologie", "OBLIGATORIE pentru ascită", "1–7 zile"],
            ["PET-CT (posibil)", "Sensibilitate superioară pentru M distant", "2–3 săpt"],
            ["IHC pe biopsie (HER2, PD-L1, MSI)", "Decide imunoterapie / terapii țintite", "2–4 săpt"],
        ],
        col_widths=[Inches(2.0), Inches(3.3), Inches(1.2)])

    add_paragraph(doc, "Plus, posibil: laparoscopie diagnostică, EUS, consult nutriționist, consult cardiolog.")
    add_callout(doc, "Timp estimativ de la consult la tratament", "3–6 săptămâni, variind în funcție de cât de rapid se fac investigațiile suplimentare.", color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 4 — INVESTIGATII
    # ============================================================

    add_heading_bar(doc, "4. Investigațiile posibile — explicate simplu", level=1, color=COLOR_PRIMARY)

    # 4.1 Paracenteza
    add_heading_bar(doc, "4.1 Paracenteza diagnostică — „pipeta care testează apa\"", level=2)
    add_quote(doc, "Analogie: imaginează-ți că ai un acvariu și vrei să știi dacă apa e curată sau contaminată. Iei o pipetă, extragi un eșantion mic, îl pui sub microscop. Exact asta e paracenteza.")

    add_paragraph(doc, "Cum se face (pas cu pas):", bold=True)
    for i, step in enumerate([
        "Pregătire — tata stă întins pe pat, cu burta dezvelită",
        "Anestezie locală — xilină (ca la dentist) într-un punct al burții. Nu doare, doar o înțepătură",
        "Ghidaj ecografic — ecografie portabilă pentru a vedea unde e apa",
        "Introducere ac subțire (ca cele pentru puncție venoasă, dar mai lung)",
        "Extragere lichid — 10–30 ml aspirate în seringă",
        "Trimitere la laborator — împărțit în 4–5 eprubete pentru analize diferite",
        "Scoatere ac + pansament — totul durează 10–15 minute",
    ], start=1):
        add_numbered(doc, step)

    add_paragraph(doc, "Ce analize se fac pe lichid:", bold=True)
    add_table(doc,
        headers=["Analiză", "Ce caută", "De ce"],
        rows=[
            ["Citologie", "Celule tumorale", "E malignă sau nu?"],
            ["Biochimie", "Proteine, albumină, LDH", "Tipul de ascită"],
            ["Număr celule", "Leucocite, hematii", "Infecție sau hemoragie"],
            ["Culturi microbiene", "Bacterii", "Exclude infecție"],
            ["NGS (opțional)", "Mutații ADN tumoral", "Identifică originea celulelor"],
        ],
        col_widths=[Inches(1.4), Inches(2.0), Inches(3.0)])

    add_callout(doc, "[CERT] Riscuri", "complicații majore sub 1%. Efecte comune: disconfort local 1–2 zile.", color=COLOR_OK, bg_hex=COLOR_OK_BG)
    add_callout(doc, "Durată rezultate", "citologie 3–7 zile; biochimie aceeași zi; NGS 2–3 săpt.", color=COLOR_INFO, bg_hex=COLOR_INFO_BG)
    add_paragraph(doc,
        "Locație efectuare: ambulator, fie la oncolog, fie la spital de zi, fie la radiologie intervențională. "
        "Nu e nevoie de internare.")

    # 4.2 Laparoscopia
    add_heading_bar(doc, "4.2 Laparoscopia diagnostică — „camera care vede direct subsolul\"", level=2)
    add_quote(doc, "Analogie: dacă pipeta din acvariu e neclară (citologia negativă dar suspiciune clinică rămâne), bagi o cameră video mică în acvariu să vezi direct cu ochii dacă sunt alge pe pereți.")

    add_paragraph(doc, "Cum se face:", bold=True)
    for step in [
        "Anestezie generală, tata doarme pe durata procedurii",
        "Incizie mică (5–10 mm) lângă ombilic",
        "Introducere cameră video miniaturală pe tub flexibil",
        "Insuflare CO2 — gazul umflă burta pentru spațiu",
        "Examinare directă a peritoneului, organelor, tumorii",
        "Biopsii țintite — dacă se văd noduli pe peritoneu",
        "Aspirare ascită pentru citologie suplimentară",
        "Scoatere instrumente + închidere cu 1–2 fire",
    ]:
        add_numbered(doc, step)

    add_callout(doc, "Când e indicată",
        "paracenteza negativă dar suspiciune clinică rămâne; cantitate prea mică pentru puncție sigură; "
        "se dorește vizualizare directă; biopsie direct din peritoneu.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_paragraph(doc, "Durată și recuperare:", bold=True)
    add_bullet(doc, "Procedură: 30–60 minute")
    add_bullet(doc, "Spitalizare: 1 noapte (uneori ambulator)")
    add_bullet(doc, "Recuperare: 3–7 zile")
    add_bullet(doc, "Durere: ușoară, controlată cu paracetamol")

    add_callout(doc, "[CERT] Risc", "complicații majore 1–3% (mai mari decât paracenteza, dar rare). Anestezie generală — risc cardiovascular la pacient post-stent. DE DISCUTAT CU CARDIOLOG ÎNAINTE.", color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    # 4.3 PET-CT
    add_heading_bar(doc, "4.3 PET-CT — „drona termică\"", level=2)
    add_quote(doc, "Analogie: CT-ul obișnuit e ca o fotografie a casei din afară — vezi structura. PET-CT e ca o cameră termică care suprapune peste fotografia normală o hartă cu „unde e metabolism activ\" — unde celulele „ard\" mult glucoză.")

    add_paragraph(doc, "Cum funcționează:", bold=True)
    add_numbered(doc, "Injectare cu FDG (glucoză radioactivă). Nu doare, doar o înțepătură.")
    add_numbered(doc, "Așteptare 60 min pentru absorbție FDG. Celulele tumorale mănâncă mai multă glucoză.")
    add_numbered(doc, "Scanare PET — detectează radioactivitatea. Zonele „aprinse\" apar colorate.")
    add_numbered(doc, "Scanare CT simultan pentru anatomie.")
    add_numbered(doc, "Suprapunere — imagini 3D color cu zone suspecte evidențiate.")

    add_paragraph(doc, "Ce vede PET-CT ce nu vede CT normal:", bold=True)
    add_bullet(doc, "Metastaze mici (< 1 cm) invizibile morfologic")
    add_bullet(doc, "Ganglioni „aprinși\" normali dimensional dar activi metabolic")
    add_bullet(doc, "Activitate tumorală reziduală după tratament (răspuns?)")
    add_bullet(doc, "Recidive precoce înainte să apară ca masă vizibilă")

    add_callout(doc, "Când e indicat pentru tata",
        "Înainte de tratament (confirmare M0) | După chimio preoperatorie (evaluare răspuns) | "
        "După chirurgie (baseline urmărire) | La suspiciune recidivă.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_callout(doc, "Limitări",
        "Fals pozitive: inflamații, infecții, țesut post-chirurgical. Fals negative: tumori lente, mucinoase, foarte mici. "
        "Cost privat 3000–4500 lei; decontare CAS parțială.",
        color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    # 4.4 EUS
    add_heading_bar(doc, "4.4 EUS (Eco-endoscopia) — „ultrasunetele din interior\"", level=2)
    add_quote(doc, "Analogie: ecografia abdominală obișnuită (cea din 14.04) e ca și cum ai pune urechea la peretele exterior al casei. EUS e ca și cum ai pune urechea direct pe peretele interior, cu camera în coridor.")

    add_paragraph(doc, "Cum funcționează:")
    add_bullet(doc, "Endoscop (ca cel de la gastroscopie) cu probă de ultrasunete la vârf")
    add_bullet(doc, "Se introduce pe gură până la esofag/stomac")
    add_bullet(doc, "Scanează structura peretelui esofagian în detaliu")
    add_bullet(doc, "Poate face biopsii profunde din ganglioni sau tumori")

    add_paragraph(doc, "Utilitate:", bold=True)
    add_bullet(doc, "Stadializare T mai precisă (cât adânc a penetrat tumora)")
    add_bullet(doc, "Biopsii ganglioni periesofagieni")
    add_bullet(doc, "Evaluare răspuns la chimioterapie")

    add_paragraph(doc, "Procedură: sedare ușoară (midazolam/propofol), durata 20–40 minute, recuperare 1–2 ore.")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 5 — MARKERI MOLECULARI
    # ============================================================

    add_heading_bar(doc, "5. Markerii moleculari — cheile care deschid tratamente", level=1, color=COLOR_PRIMARY)

    add_quote(doc, "Gândește-te la celulele tumorale ca la uși încuiate. Medicii au chei specifice (medicamente țintite) care funcționează DOAR pe anumite uși. Marker molecular = verificarea ce lacăt e pe ușă, ca să știm ce cheie să folosim.")

    # HER2
    add_heading_bar(doc, "5.1 HER2 — „lacătul care reacționează la Herceptin\"", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Analogie: ", "bold": True},
        {"text": "HER2 e ca un receptor de creștere pe suprafața celulei — dacă celula tumorală are MULȚI receptori HER2, crește foarte repede. E o „super-antenă\" care amplifică semnalele de multiplicare."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Cum se testează: ", "bold": True},
        {"text": "imunohistochimie pe biopsie. Rezultat: IHC 0 / 1+ / 2+ / 3+."},
    ])
    add_bullet(doc, "IHC 0/1+ (negativ): chimio normală FLOT")
    add_bullet(doc, "IHC 3+ (sau 2+ confirmat ISH) pozitiv: se adaugă trastuzumab (Herceptin)")
    add_mixed_paragraph(doc, [
        {"text": "Ce face trastuzumab: ", "bold": True},
        {"text": "o moleculă-cheie care se prinde de receptorul HER2 și îl blochează — celula nu mai primește semnale de multiplicare și moare."},
    ])
    add_callout(doc, "[CERT] Frecvență și impact",
        "15–20% din adenocarcinoamele eso-gastrice sunt HER2+. Studiul Keynote-811 (2024) a arătat că pacienții HER2+ "
        "tratați cu trastuzumab + pembrolizumab + chimio au supraviețuire mediană ~20 luni vs ~14 luni fără imunoterapie.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # PD-L1
    add_heading_bar(doc, "5.2 PD-L1 — „scutul invizibilității tumorii\"", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Analogie: ", "bold": True},
        {"text": "sistemul imunitar e o armată care patrulează prin corp și omoară intruși. PD-L1 e ca un steag alb pe celula tumorală care îi spune soldatului „sunt un cetățean normal, nu mă ataca\". Astfel tumora se ascunde de apărarea corpului."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Cum se testează: ", "bold": True},
        {"text": "imunohistochimie pe biopsie cu măsurarea CPS (Combined Positive Score)."},
    ])
    add_bullet(doc, "CPS < 1: imunoterapie mai puțin eficientă")
    add_bullet(doc, "CPS ≥ 1: imunoterapie eligibilă")
    add_bullet(doc, "CPS ≥ 10: imunoterapie foarte eficientă")
    add_mixed_paragraph(doc, [
        {"text": "Ce face pembrolizumab: ", "bold": True},
        {"text": "„scoate steagul alb\" de pe celula tumorală. Sistemul imunitar îl vede acum ca intrus și îl atacă."},
    ])
    add_callout(doc, "[CERT] Impact",
        "50–60% din adenocarcinoame au CPS ≥ 1, ~30% au CPS ≥ 10. Keynote-590 (2021): pembrolizumab + chemo → "
        "supraviețuire mediană 13.5 luni vs 9.4 luni la CPS ≥ 10 — un plus de 50% la durată de viață.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # MSI
    add_heading_bar(doc, "5.3 MSI / MMR — „defecte în fotocopiator\"", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Analogie: ", "bold": True},
        {"text": "celula normală are un „fotocopiator ADN\" care se verifică înainte de a face o copie. MSI-High înseamnă că fotocopiatorul e stricat — celula face multe greșeli genetice, devine instabilă, dar e și mai vizibilă pentru sistemul imunitar."},
    ])
    add_paragraph(doc, "De ce e util: tumorile MSI-H au mutații multe = mulți antigeni = răspund foarte bine la imunoterapie.")
    add_callout(doc, "Frecvență",
        "[PROBABIL] 5–15% din adenocarcinoame. Dacă MSI-H, pembrolizumab singur (fără chimio) poate fi suficient. Răspuns durabil la 40–50% din pacienți.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # Claudin
    add_heading_bar(doc, "5.4 Claudin-18.2 — „noua țintă terapeutică\" (2024)", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Analogie: ", "bold": True},
        {"text": "claudin-18.2 e o proteină specifică stomacului — normal e ascunsă între celule, dar în cancerele gastrice și eso-gastrice devine expusă pe suprafață. E ca un stigmat vizibil care nu există la celule normale."},
    ])
    add_mixed_paragraph(doc, [
        {"text": "Medicamentul țintit: ", "bold": True},
        {"text": "zolbetuximab (Vyloy) — moleculă-cheie aprobată 2024 (SPOTLIGHT)."},
    ])
    add_callout(doc, "Frecvență și impact",
        "[CERT] ~30–40% din adenocarcinoamele eso-gastrice. Adăugarea zolbetuximab la chimio prelungește supraviețuirea mediană cu ~2–3 luni.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # Tabel recap markers
    add_heading_bar(doc, "5.5 Tabel recapitulativ markeri — ce speri să găsească biopsia", level=2)
    add_table(doc,
        headers=["Marker", "Frecvență așteptată", "Medicament țintit", "Beneficiu tipic"],
        rows=[
            ["HER2+", "15–20%", "Trastuzumab + pembrolizumab", "+6 luni supraviețuire"],
            ["PD-L1 ≥10", "~30%", "Pembrolizumab", "+4–6 luni supraviețuire"],
            ["MSI-H", "5–15%", "Pembrolizumab (poate singur)", "Răspuns durabil, uneori remisiune"],
            ["Claudin-18.2+", "30–40%", "Zolbetuximab", "+2–3 luni supraviețuire"],
            ["Toate negative", "—", "FLOT standard", "Supraviețuire mediană standard"],
        ],
        col_widths=[Inches(1.2), Inches(1.4), Inches(2.0), Inches(2.0)])

    add_callout(doc, "[PROBABIL]",
        "La pacient oncologic modern, există 50–70% șansă ca biopsia să găsească MINIM UN marker țintit. "
        "Asta deschide tratamente mai eficiente.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 6 — 4 SCENARII
    # ============================================================

    add_heading_bar(doc, "6. Cele 4 scenarii combinatorii — planul de tratament", level=1, color=COLOR_PRIMARY)

    add_paragraph(doc, "Ai 2 probe majore în evaluare:")
    add_bullet(doc, "Biopsia (prelevată 17.04, rezultat așteptat 24.04–01.05) → ce tip de celule are leziunea din esofag")
    add_bullet(doc, "Ascita (observată la CT 20.04, de analizat prin paracenteză) → dacă conține celule tumorale")
    add_paragraph(doc, "Fiecare poate fi MALIGNĂ sau BENIGNĂ. Deci avem 4 combinații:")

    # Scenariul A
    add_heading_bar(doc, "SCENARIU A — Biopsie MALIGNĂ + Ascită BENIGNĂ", level=2, color=COLOR_OK)
    add_quote(doc, "Analogie casă: Zidul din coridor E infectat cu mucegai, DAR apa din subsol e doar infiltrație locală — nu s-a răspândit dincolo.", color_left=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_callout(doc, "[PROBABIL] Cel mai probabil scenariu",
        "~60–70% probabilitate, bazat pe: aspectul CT (masă localizată fără metastaze), ganglioni sub pragul patologic, lipsa simptomelor sistemice dramatice.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_mixed_paragraph(doc, [
        {"text": "Diagnostic clinic: ", "bold": True},
        {"text": "Cancer esofagian Siewert II, stadiu probabil III (T3-T4, N0-N1, M0)."},
    ])

    add_mixed_paragraph(doc, [
        {"text": "Plan de tratament — CURATIV (scopul e vindecarea):", "bold": True, "size": 12, "color": COLOR_OK},
    ])

    # Tabel plan A
    add_table(doc,
        headers=["Fază", "Durată", "Ce se face"],
        rows=[
            ["Faza 1: Chimio preoperatorie", "8 săpt", "FLOT: 4 cicluri (5-FU + Leucovorină + Oxaliplatin + Docetaxel)"],
            ["Faza 2: Pauză + evaluare", "4–6 săpt", "CT/PET-CT repeat pentru răspuns + recuperare"],
            ["Faza 3: Chirurgie", "1 op + 4–6 săpt recuperare", "Ezofago-gastrectomie parțială + limfadenectomie"],
            ["Faza 4: Chimio postoperatorie", "8 săpt", "FLOT 4 cicluri suplimentare"],
        ],
        col_widths=[Inches(2.2), Inches(1.4), Inches(3.0)])

    add_callout(doc, "Total calendar",
        "~6 luni tratament activ intenționat CURATIV. Studiul FLOT4 (2019): supraviețuire mediană la 5 ani ~45%. "
        "Răspuns patologic complet (tumora dispărută la microscop): ~15–20% din pacienți.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # Scenariul B
    add_heading_bar(doc, "SCENARIU B — Biopsie MALIGNĂ + Ascită MALIGNĂ", level=2, color=COLOR_CRIT)
    add_quote(doc, "Analogie casă: Zidul din coridor E infectat cu mucegai, ȘI mucegaiul s-a răspândit invizibil pe toți pereții subsolului.", color_left=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_callout(doc, "[PROBABIL] Scenariu mai îngrijorător",
        "~15–25% probabilitate, bazat pe: cantitatea de ascită (15+28 mm), aspectul infiltrativ la CT, extensia la fundul stomacului.",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_mixed_paragraph(doc, [
        {"text": "Diagnostic clinic: ", "bold": True},
        {"text": "Cancer esofagian Siewert II cu carcinomatoză peritoneală, stadiu IV (M1 confirmat)."},
    ])

    add_mixed_paragraph(doc, [
        {"text": "Plan de tratament — PALIATIV SISTEMIC (scopul e CONTROLUL, nu vindecarea):", "bold": True, "size": 12, "color": COLOR_CRIT},
    ])

    add_bullet(doc, "Chimioterapie sistemică linia 1: FLOT sau FOLFOX", bold_prefix="FAZA 1:")
    add_bullet(doc, "+ Pembrolizumab dacă PD-L1 CPS ≥ 1 (Keynote-590/811)", level=1)
    add_bullet(doc, "+ Trastuzumab dacă HER2+ (Keynote-811)", level=1)
    add_bullet(doc, "+ Zolbetuximab dacă Claudin-18.2+ (SPOTLIGHT)", level=1)
    add_bullet(doc, "Administrare la 2–3 săptămâni, continuu luni/ani", level=1)
    add_bullet(doc, "Evaluare răspuns prin CT la 3 luni", level=1)

    add_mixed_paragraph(doc, [
        {"text": "NU SE FACE CHIRURGIE CURATIVĂ ", "bold": True, "color": COLOR_CRIT},
        {"text": "(dacă e răspândit, nu mai ai o zonă localizată de scos)."},
    ])

    add_paragraph(doc, "Chirurgie DOAR paliativă (la nevoie):", bold=True)
    add_bullet(doc, "Stent esofagian dacă deglutiția e severă")
    add_bullet(doc, "Jejunostomă pentru alimentație")
    add_bullet(doc, "Laparoscopie dacă ocluzie intestinală")

    add_heading_bar(doc, "[CERT] Supraviețuire mediană în scenariul B (date 2024–2025)", level=3, color=COLOR_OK)
    add_bullet(doc, "FLOT/FOLFOX linia 1 fără imunoterapie: 10–14 luni median")
    add_bullet(doc, "FLOT + pembrolizumab (PD-L1+): 17–20 luni median (Keynote-590)")
    add_bullet(doc, "Trastuzumab + chemo + pembro (HER2+): ~20 luni median (Keynote-811)")
    add_bullet(doc, "Mulți pacienți ajung la 2–3 ani cu protocoale moderne")

    add_callout(doc, "Mesaj-cheie",
        "Scenariul B NU ESTE RENUNȚARE. E strategie diferită. Pacienții cu stadiu IV trăiesc ani buni cu chimio + imunoterapie, deseori cu calitate rezonabilă a vieții.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    # Scenariul C
    add_heading_bar(doc, "SCENARIU C — Biopsie BENIGNĂ + Ascită BENIGNĂ", level=2, color=COLOR_INFO)
    add_quote(doc, "Analogie casă: Zidul din coridor NU e mucegai — e doar igrasie severă cu inflamație. Apa din subsol e de la o conductă veche care picură.", color_left=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_callout(doc, "[INCERT] Scenariu mai puțin probabil",
        "~5–10% probabilitate, dar NU imposibil.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_paragraph(doc, "Ce POATE să însemne biopsia „benignă\":", bold=True)
    add_numbered(doc, "Esofagită severă erozivă + stenoză inflamatorie — reflux cronic")
    add_numbered(doc, "Displazie high-grade (precanceroasă) — ATENȚIE, necesită tratament proactiv")
    add_numbered(doc, "Esofag Barrett cu displazie — intervenție endoscopică preventivă")
    add_numbered(doc, "Tumoră benignă rară — GIST benign, leiomioma, lipoma")
    add_numbered(doc, "Sindrom Zenker sau diverticul — dilatare benignă")
    add_numbered(doc, "ARTEFACT DE BIOPSIE — necesită REBIOPSIE")

    add_paragraph(doc, "Plan de tratament — depinde exact de ce arată biopsia:", bold=True)
    add_paragraph(doc, "C.1. Dacă e esofagită / inflamație: IPP în doză mare + dilatare endoscopică + lifestyle anti-reflux + endoscopii control.")
    add_paragraph(doc, "C.2. Dacă e displazie high-grade: REZECȚIE ENDOSCOPICĂ (EMR sau ESD), ablație prin radiofrecvență, urmărire.")
    add_paragraph(doc, "C.3. Dacă e tumoră benignă: Rezecție locală endoscopică sau chirurgie minim-invazivă + urmărire imagistică.")
    add_paragraph(doc, "C.4. Dacă biopsia e non-reprezentativă: REBIOPSIE obligatorie + EUS.")
    add_paragraph(doc, "În toate sub-scenariile C, oncologul predă înapoi la gastroenterolog — nu mai ai nevoie de oncologie activă.")

    # Scenariul D
    add_heading_bar(doc, "SCENARIU D — Biopsie BENIGNĂ + Ascită MALIGNĂ", level=2, color=COLOR_WARN)
    add_quote(doc, "Analogie casă: Zidul din coridor e doar igrasie, DAR apa din subsol conține mucegai răspândit — vine de undeva ALTUNDE.", color_left=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    add_callout(doc, "[INCERT] Scenariu rar dar posibil",
        "~1–3% probabilitate. Situație diagnostică complexă.",
        color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    add_paragraph(doc, "Ce poate însemna:", bold=True)
    add_numbered(doc, "Biopsia a ratat tumora — artefact de prelevare. NECESITĂ REBIOPSIE URGENTĂ.")
    add_numbered(doc, "Există o ALTĂ tumoră primară (gastric, pancreatic, biliar, mezoteliom) care a generat carcinomatoza.")
    add_numbered(doc, "Eroare de laborator — rar, dar se întâmplă.")

    add_paragraph(doc, "Plan de investigație obligatoriu:", bold=True)
    add_bullet(doc, "Rebiopsie endoscopică URGENT cu prelevare multiplă", bold_prefix="PAS 1:")
    add_bullet(doc, "Căutarea tumorii primare: PET-CT, colonoscopie repetat, RMN pancreas, markeri tumorali (CEA, CA 19-9, CA 72-4, CA 125), NGS ascită", bold_prefix="PAS 2:")
    add_bullet(doc, "Decizie tratament după clarificare (scenariu B sau protocol specific alt primar)", bold_prefix="PAS 3:")

    add_callout(doc, "Durată investigație",
        "2–4 săptămâni suplimentare înainte de început tratament. [CERT] Paracenteza cu citologie și NGS poate identifica originea celulelor tumorale din ascită.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 7 — FLOT DETALIAT
    # ============================================================

    add_heading_bar(doc, "7. Protocolul FLOT — explicat în detaliu", level=1, color=COLOR_PRIMARY)

    add_paragraph(doc, "FLOT = acronim pentru 4 medicamente administrate simultan:", bold=True)
    add_bullet(doc, "5-Fluorouracil (5-FU)", bold_prefix="F —")
    add_bullet(doc, "Leucovorină", bold_prefix="L —")
    add_bullet(doc, "Oxaliplatin", bold_prefix="O —")
    add_bullet(doc, "Docetaxel (Taxotere)", bold_prefix="T —")

    add_callout(doc, "[CERT] Standardul european",
        "FLOT este standardul european de îngrijire pentru cancer eso-gastric Siewert II/III din 2019, bazat pe studiul FLOT4 (Al-Batran et al., Lancet 2019) care a demonstrat supraviețuire mai bună decât protocolul anterior (ECF/ECX).",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_heading_bar(doc, "7.1 Ce face fiecare medicament — analogii", level=2)

    # 5-FU
    add_heading_bar(doc, "5-Fluorouracil — „Terminatorul de ADN\"", level=3)
    add_bullet(doc, "E o substanță care imită o cărămidă de ADN (uracilul)")
    add_bullet(doc, "Celulele tumorale îl „mănâncă\" gândind că e material de construcție normal")
    add_bullet(doc, "Odată integrat, blochează sinteza ADN-ului — celula nu se mai poate divide")
    add_bullet(doc, "Eficient pe celulele cu diviziune rapidă = celule tumorale (dar și măduvă, mucoasă, foliculi păr)")

    # Leucovorin
    add_heading_bar(doc, "Leucovorină — „Amplificatorul\"", level=3)
    add_bullet(doc, "Nu e o chimioterapie propriu-zisă")
    add_bullet(doc, "E un cofactor care intensifică acțiunea 5-FU")
    add_bullet(doc, "Face 5-FU să funcționeze 3–5x mai puternic în celulele tumorale")

    # Oxaliplatin
    add_heading_bar(doc, "Oxaliplatin — „Legătorul de ADN\"", level=3)
    add_bullet(doc, "Conține platină — atom care se leagă de ADN-ul celulei")
    add_bullet(doc, "Provoacă „cross-links\" între cele 2 spirale ale ADN-ului")
    add_bullet(doc, "Celula nu mai poate deschide ADN-ul pentru replicare → moare")
    add_bullet(doc, "Efect secundar specific: neuropatie periferică (furnicături degete, senzație rece mâini)")

    # Docetaxel
    add_heading_bar(doc, "Docetaxel — „Paralizatorul de diviziune\"", level=3)
    add_bullet(doc, "Se extrage din scoarța de tisa")
    add_bullet(doc, "Blochează „scheletul\" celulei (microtubulii) în timpul diviziunii")
    add_bullet(doc, "Celula încearcă să se divida, dar nu poate → moare")
    add_bullet(doc, "Efect secundar specific: căderea părului (temporară), reducerea leucocitelor")

    add_heading_bar(doc, "7.2 Cum arată un ciclu FLOT (2 săptămâni)", level=2)

    add_table(doc,
        headers=["Ziua", "Loc", "Acțiune"],
        rows=[
            ["Ziua 1", "LA SPITAL (4–6 ore)", "Analize + premedicație + perfuzii (docetaxel → oxaliplatin → leucovorină → 5-FU bolus)"],
            ["Ziua 1 seara", "ACASĂ", "Pleacă acasă cu pompă portabilă 5-FU (46 ore)"],
            ["Ziua 2–3", "ACASĂ", "Pompa continuă să se perfuzeze lent"],
            ["Ziua 3 seara", "SPITAL (15 min)", "Scoatere pompă"],
            ["Ziua 4–14", "ACASĂ", "Recuperare. Efecte max zile 5–10. Analize la zilele 7+10"],
            ["Ziua 14", "—", "Următorul ciclu începe"],
        ],
        col_widths=[Inches(1.0), Inches(1.6), Inches(3.8)])

    add_callout(doc, "Total",
        "4 cicluri pre-operator + 4 cicluri post-operator = 8 cicluri × 2 săpt = 16 săptămâni chimioterapie activă (~4 luni) + 4–6 săpt pauză chirurgie + chirurgie + 4–6 săpt recuperare = 6 luni total.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_heading_bar(doc, "7.3 Efecte secundare și cum se gestionează", level=2)

    add_table(doc,
        headers=["Efect", "Frecvență", "Când", "Gestionare"],
        rows=[
            ["Greață + vărsături", "50–70%", "Zilele 1–5", "Antiemetice profilactic (ondansetron, aprepitant, dexametazonă)"],
            ["Oboseală", "80–90%", "Continuu, max zile 3–7", "Odihnă, nutriție, activitate moderată"],
            ["Diaree", "30–50%", "Zile 3–10", "Loperamidă, hidratare"],
            ["Mucozită (afte)", "40%", "Zile 5–10", "Igienă orală, clătire bicarbonat, analgezice locale"],
            ["Neutropenie (leucocite ↓)", "60–80%", "Zile 7–14", "Analize regulate; la < 1000 → G-CSF (filgrastim)"],
            ["Căderea părului", "60%", "Din săpt 3", "Temporară, revine după tratament"],
            ["Neuropatie (furnicături)", "40–60%", "Progresiv cumulativ", "La grad 2 → reducere doză oxaliplatin; evitare frig"],
            ["Reacții alergice oxaliplatin", "5–10%", "Uneori ciclul 3+", "Premedicație antihistaminice; rareori oprire"],
            ["Infecții", "15–25%", "În neutropenie", "Evitare aglomerații, igienă, febră > 38°C → URGENȚĂ"],
        ],
        col_widths=[Inches(1.5), Inches(0.9), Inches(1.3), Inches(2.6)])

    add_callout(doc, "[CERT] Ajustare doze vârstnici",
        "La pacienții > 65 ani cu comorbidități, dozele se reduc de rutină la 80–85% din cele standard pentru tolerabilitate mai bună.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_heading_bar(doc, "7.4 Ce aștepți în primele 2 săptămâni de FLOT", level=2)
    add_bullet(doc, "senzație OK, eventual ușoară oboseală sau greață spre seară", bold_prefix="Zi 1 (luni):")
    add_bullet(doc, "de obicei cele mai rele — greață mai pregnantă, oboseală marcată, lipsă apetit", bold_prefix="Ziua 2–3:")
    add_bullet(doc, "se ameliorează treptat. Revine pofta de mâncare", bold_prefix="Ziua 4–7:")
    add_bullet(doc, "ATENȚIE la leucocite scăzute. Orice febră > 38°C → telefon oncolog IMEDIAT", bold_prefix="Ziua 7–10:")
    add_bullet(doc, "revenire progresivă, pregătire pentru ciclul următor", bold_prefix="Ziua 10–14:")

    add_heading_bar(doc, "7.5 Alternative la FLOT dacă nu e tolerat", level=2)
    add_bullet(doc, "FLO (fără docetaxel) — reducere efecte secundare, eficiență ~80% din FLOT")
    add_bullet(doc, "FOLFOX (5-FU + Leucovorină + Oxaliplatin) — fără docetaxel, standard stadiu IV")
    add_bullet(doc, "FOLFIRI (5-FU + Leucovorină + Irinotecan) — linia 2 după eșec primă linie")
    add_bullet(doc, "Capecitabină + Oxaliplatin (CAPOX) — capecitabina orală în loc de 5-FU perfuzabil")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 8 — IMUNOTERAPIA
    # ============================================================

    add_heading_bar(doc, "8. Imunoterapia — o categorie nouă de tratament", level=1, color=COLOR_PRIMARY)

    add_heading_bar(doc, "8.1 Ce e diferit față de chimio", level=2)
    add_paragraph(doc, "Chimioterapia ucide direct celulele cu diviziune rapidă (include și celule normale).")
    add_paragraph(doc, "Imunoterapia EDUCĂ sistemul imunitar să recunoască și să atace celulele tumorale. E ca și cum instruiești armata să vadă intruși pe care nu-i vedea înainte.")

    add_heading_bar(doc, "8.2 Pembrolizumab (Keytruda)", level=2)
    add_mixed_paragraph(doc, [
        {"text": "Ce face: ", "bold": True},
        {"text": "„Scoate steagul alb\" (PD-L1) de pe celula tumorală. Sistemul imunitar o vede acum și o atacă."},
    ])
    add_paragraph(doc, "Administrare: perfuzie la 3 săptămâni (200 mg), 30 min per doză.")
    add_callout(doc, "[CERT] Eficiență",
        "Keynote-590 (2021): în esofagian avansat + chemo + pembrolizumab vs chemo singur → supraviețuire mediană 12.4 vs 9.8 luni. Cu CPS ≥ 10 → 13.5 vs 9.4 luni.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)
    add_paragraph(doc, "Efecte secundare principale:")
    add_bullet(doc, "Generale: oboseală, greață ușoară, erupție cutanată")
    add_bullet(doc, "Autoimune (rare, 5–15%): tiroidită, pneumonită, colită, hepatită, miocardita — necesită intervenție promptă")

    add_heading_bar(doc, "8.3 Trastuzumab (Herceptin) — dacă HER2+", level=2)
    add_paragraph(doc, "Ce face: se prinde de receptorul HER2 și blochează semnalele de creștere.")
    add_paragraph(doc, "Administrare: perfuzie la 3 săptămâni, combinat cu chimio + posibil pembrolizumab.")
    add_callout(doc, "[CERT] Eficiență",
        "Keynote-811 (2024): HER2+ avansat cu chimio + trastuzumab + pembrolizumab → supraviețuire mediană 20 luni vs ~14 luni fără imunoterapie.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)
    add_callout(doc, "Efect secundar specific",
        "Poate afecta inima (cardiotoxicitate) — necesită ECHO cardiac la fiecare 3 luni. La tata post-stent 2012, cardiologul va fi implicat activ.",
        color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    add_heading_bar(doc, "8.4 Zolbetuximab (Vyloy) — dacă Claudin-18.2+", level=2)
    add_paragraph(doc, "Ce face: se prinde de claudin-18.2 pe celula tumorală și marchează celula pentru distrugere imună.")
    add_paragraph(doc, "Aprobat: FDA 2024, EMA 2024, în România accesibil din 2025–2026.")
    add_callout(doc, "[CERT] Eficiență",
        "SPOTLIGHT (2023): la pacienți cu claudin-18.2+ (30–40% din pacienți) → prelungire supraviețuire mediană +2–3 luni.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_heading_bar(doc, "8.5 Accesibilitate în România (2026)", level=2)
    add_paragraph(doc, "[PROBABIL] Majoritatea imunoterapiilor sunt disponibile prin Programul Național Oncologic (PNO) CNAS:")
    add_bullet(doc, "Pembrolizumab: aprobat pentru esofagian avansat CPS ≥ 1 → decontat CNAS")
    add_bullet(doc, "Trastuzumab: aprobat HER2+ → decontat CNAS")
    add_bullet(doc, "Zolbetuximab: aprobare recentă 2024–2025 → verifică eligibilitate la centrul tău")
    add_bullet(doc, "Nivolumab: pentru unele indicații")

    add_paragraph(doc, "Documente necesare pentru accesare PNO:", bold=True)
    add_numbered(doc, "Dosar medical complet")
    add_numbered(doc, "Rezultate IHC (HER2, PD-L1 CPS, MSI, claudin-18.2)")
    add_numbered(doc, "Prescripție oncolog cu justificare")
    add_numbered(doc, "Decizie Comisie de Experți PNO")
    add_callout(doc, "Timp aprobare", "2–4 săptămâni de la cererea oncologului.", color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 9 — NUTRITIE
    # ============================================================

    add_heading_bar(doc, "9. Nutriția — cum trebuie să mănânce tata acum", level=1, color=COLOR_PRIMARY)

    add_heading_bar(doc, "9.1 Provocarea specifică", level=2)
    add_paragraph(doc, "Tata are 2 challenge-uri simultane:", bold=True)
    add_numbered(doc, "Stenoză esofagiană — canalul e îngustat, nu se mai înghite bine mâncarea solidă")
    add_numbered(doc, "Pregătire pentru chimioterapie — trebuie „întărit\" pentru a tolera tratamentul")

    add_callout(doc, "[CERT] Impact nutriție",
        "Nutriția preoperatorie e factor major INDEPENDENT pentru succesul tratamentului. Pacienți cu status nutrițional slab au: risc crescut infecții post-op (2–3x), recuperare mai lentă, toleranță chimio redusă, supraviețuire redusă cu 20–30%.",
        color=COLOR_WARN, bg_hex=COLOR_WARN_BG)

    add_heading_bar(doc, "9.2 Ce să mănânce TATA acum", level=2)

    add_paragraph(doc, "Principii:", bold=True)
    add_numbered(doc, "Mâncare moale / păstoasă / lichidă — pasează, macerează")
    add_numbered(doc, "Porții mici, frecvent — 6–8 mese mici/zi, nu 3 mari")
    add_numbered(doc, "Calorii concentrate — mai multe calorii într-un volum mic")
    add_numbered(doc, "Proteine prioritare — carne tocată, pește, ouă, lactate, leguminoase")
    add_numbered(doc, "Lichide la masă — apă, ceai, supă ajută pasajul")

    add_heading_bar(doc, "Alimente RECOMANDATE (stenoză + tratament)", level=3, color=COLOR_OK)
    add_table(doc,
        headers=["Categorie", "Exemple"],
        rows=[
            ["Supe", "Cremă de legume, supă de pui cu tăiței mici, ciorbă strecurată"],
            ["Carne", "Piept pui fiert + pireu, pește aburit, chiftele moi, pateu de ficat"],
            ["Pește", "Somon/macrou aburit, pescador moale"],
            ["Lactate", "Iaurt gras, brânză proaspătă, mozzarella, chefir, kașcaval"],
            ["Ouă", "Omletă, ouă moi, scrambled eggs cu smântână"],
            ["Legume", "Fierte până la moliciune, pasate — dovlecei, morcovi, cartofi"],
            ["Fructe", "Banane, compoturi, piure de mere, pere gătite"],
            ["Cereale", "Gris cu lapte, fulgi de ovăz fieri, orez puhav, paste bine fierte"],
            ["Grăsimi sănătoase", "Ulei măsline extra-virgin, unt, avocado pasat"],
            ["Dulciuri", "Budinci, creme, iaurt cu miere"],
            ["Suplimente", "Ensure / Nutridrink (shake-uri nutriționale 300 kcal/bucata)"],
        ],
        col_widths=[Inches(1.6), Inches(4.6)])

    add_heading_bar(doc, "Alimente DE EVITAT", level=3, color=COLOR_CRIT)
    add_bullet(doc, "Pâine uscată, biscuiți crocanți, rondele de orez")
    add_bullet(doc, "Carne fibroasă, pui cu os (bucăți neîmbucătățite)")
    add_bullet(doc, "Nuci, semințe, boabe întregi")
    add_bullet(doc, "Fructe tari (măr crud, pară tare, alune)")
    add_bullet(doc, "Legume crude tari (morcov crud, țelină, ardei crud)")
    add_bullet(doc, "Orez uscat, paste foarte al dente")
    add_bullet(doc, "Mâncare foarte picantă sau acidă (iritant pentru esofag)")

    add_heading_bar(doc, "9.3 Suplimente nutriționale formale", level=2)
    add_callout(doc, "[CERT] Recomandări ESPEN (European Society for Clinical Nutrition 2021–2023)",
        "Pentru pacienți oncologici: Nutridrink / Ensure / Fortimel — 1–2/zi (+300–600 kcal + proteine + vitamine). Proteine pulbere (whey) — 1–2 porții/zi. Multivitamine + minerale. Vitamina D 3000–4000 UI/zi. Omega-3 1–2 g/zi.",
        color=COLOR_OK, bg_hex=COLOR_OK_BG)

    add_heading_bar(doc, "9.4 Monitorizare greutate", level=2)
    add_callout(doc, "Obiectiv", "NU permite scădere în greutate peste 5% / lună.", color=COLOR_WARN, bg_hex=COLOR_WARN_BG)
    add_bullet(doc, "Cântărire săptămânală (aceeași balanță, aceeași oră)")
    add_bullet(doc, "Scădere > 2 kg în 2 săptămâni → escaladare imediată: nutriționist oncolog")
    add_bullet(doc, "Scădere > 10% din greutatea inițială → considerare jejunostomie sau stent esofagian")

    add_heading_bar(doc, "9.5 Dacă deglutiția se agravează brusc", level=2)
    add_paragraph(doc, "Semnale de alarmă:", bold=True)
    add_bullet(doc, "Tuse la lichide (aspirație în plămâni)")
    add_bullet(doc, "Imposibilitate înghițire chiar a lichidelor")
    add_bullet(doc, "Vărsături imediat post-deglutiție")
    add_bullet(doc, "Durere severă la înghițire")

    add_paragraph(doc, "Acțiune: telefon oncolog IMEDIAT + consult gastroenterolog. Opțiuni:")
    add_bullet(doc, "Stent esofagian — proteză metalică care dilată zona (endoscopic)")
    add_bullet(doc, "Jejunostomie — tub alimentare chirurgical")
    add_bullet(doc, "Alimentație parenterală — temporară")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 10 — SEMNALE ALARMA
    # ============================================================

    add_heading_bar(doc, "10. Semnale de alarmă — când suni 112, când aștepți", level=1, color=COLOR_PRIMARY)

    add_heading_bar(doc, "10.1 URGENȚĂ 112 IMEDIAT", level=2, color=COLOR_CRIT)

    add_callout(doc, "ANGIOEDEM — umflare bruscă",
        "La pleoape, buze, limbă, gât. Respirație dificilă bruscă. Sursa risc: interacțiune Jamesi + Triplixam (sitagliptin + perindopril).",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_callout(doc, "INFECȚIE SEVERĂ",
        "Febră > 38.5°C cu tremur / frison, mai ales în chimio. Poate fi sepsis.",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_callout(doc, "SEMNE CARDIACE",
        "Durere toracică opresivă, iradiere în braț stâng / mandibulă. Dispnee bruscă extremă. Leșin.",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_callout(doc, "HEMORAGIE",
        "Vărsături cu sânge roșu sau „zaț de cafea\". Scaune negre ca smoala (melenă). Sângerare prelungită.",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_callout(doc, "OCLUZIE INTESTINALĂ",
        "Durere abdominală severă. Vărsături persistente. Oprire tranzit. Burta foarte balonată.",
        color=COLOR_CRIT, bg_hex=COLOR_CRIT_BG)

    add_heading_bar(doc, "10.2 CONTACT ONCOLOG / URGENT în 24h (nu 112)", level=2, color=COLOR_WARN)
    add_bullet(doc, "Febră 37.5–38.5°C fără alte simptome, persistentă")
    add_bullet(doc, "Greață / vărsături care nu se ameliorează cu antiemetice")
    add_bullet(doc, "Diaree mai mult de 6 scaune/zi")
    add_bullet(doc, "Mucozită severă (gură foarte sensibilă)")
    add_bullet(doc, "Erupție cutanată nouă (imunoterapie)")
    add_bullet(doc, "Oboseală extremă brusc apărută")
    add_bullet(doc, "Scădere în greutate > 2 kg în 2 săpt")
    add_bullet(doc, "Disfagie progresivă")

    add_heading_bar(doc, "10.3 Monitorizare DE RUTINĂ (nu urgent)", level=2, color=COLOR_OK)
    add_bullet(doc, "Cântărire săptămânală")
    add_bullet(doc, "Temperatură zilnic dacă chimio activă")
    add_bullet(doc, "Jurnal efecte secundare (data + tip + severitate 1–10)")
    add_bullet(doc, "Glicemie de control (diabet Jamesi)")
    add_bullet(doc, "Tensiunea arterială săptămânal")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 11 — FAQ
    # ============================================================

    add_heading_bar(doc, "11. Întrebări frecvente ale familiei — răspunsuri pregătite", level=1, color=COLOR_PRIMARY)

    faq_list = [
        ("Q1: „Cât timp va trăi tata?\"",
         "Nu pot să-ți dau o cifră. Nici oncologul nu va da una sigură înainte de toate datele. Statisticile generale există, dar ele includ oameni cu situații foarte diferite. Tata e într-o categorie specifică: 66 de ani, diabet controlat, inimă stabilă, fără metastaze vizibile. Cu tratament corect, în scenariul cel mai probabil (stadiu III) mulți oameni trăiesc ani buni de zile. În cel mai rău caz (stadiu IV), tratamentele moderne permit 2–3 ani de calitate bună. Ceea ce contează acum e să facem pașii concreți la oncolog — nu cifrele."),
        ("Q2: „Va suferi mult?\"",
         "Nu neapărat. Medicina modernă gestionează bine durerea și efectele secundare. Chimioterapia e incomodă, nu e chinuitoare — pacienții continuă să trăiască acasă, să mănânce, să comunice. Greața e gestionată cu medicamente eficiente. Oboseala e reală dar normală. La stadiu IV, există echipe de îngrijire paliativă care se asigură că durerea e controlată complet. Suferința fizică neuitată nu mai e standardul — asta era acum 30 de ani."),
        ("Q3: „Trebuie să ne mutăm la Cluj / București?\"",
         "Nu obligatoriu. Chimioterapia se face lunar, 1–2 zile/ciclu, și multe centre din Arad / Timișoara o administrează. Chirurgia se face o singură dată — acolo da, probabil la un centru specializat. Dar majoritatea timpului, tata va fi acasă. Vom călători câteva zile pentru operație + control. Restul e acasă."),
        ("Q4: „Putem să-l îngrijim acasă pe tot parcursul tratamentului?\"",
         "Da, cu o singură excepție: operația + 1–2 săpt după. Restul — chemoterapie administrată în ambulator, apoi recuperare acasă. Familia e îngrijitorul principal. Vom avea consult nutriționist, asistent social, eventual psiholog. Tu (mama) ești esențială — îngrijirea ta de el contează la fel de mult cât medicamentele."),
        ("Q5: „O să-i cadă părul? Va arăta bolnav?\"",
         "La FLOT, părul cade parțial sau total (din săpt 3–4), dar revine complet după tratament. Mulți pacienți aleg să se radă preventiv — dă senzația de control. Se poate purta pălărie / bandană / perucă. Pielea poate deveni mai palidă, poate slăbi — toate sunt temporare. Va arăta ca cineva care urmează un tratament, nu ca cineva care moare."),
        ("Q6: „Ce pot eu să fac ca să ajut?\" (rude)",
         "Lucrurile practice contează mai mult decât vorbele. Concret: mâncare gata preparată (porții mici, nutritive, congelate); transport la consulturi; timp liber pentru Roland/mama (o după-masa/seară când preluați voi îngrijirea); telefon cu tata despre viață normală, nu despre boală; NU aduceți informații de pe Google; dacă trebuie bani pentru investigații private, fondul colectat ajută."),
        ("Q7: „Va putea mânca normal după tot tratamentul?\"",
         "Depinde mult. Dacă se face chirurgie parțială, după 6 luni de recuperare mănâncă aproape normal, cu porții mai mici (jumătate din ce mânca înainte) și mese mai dese (5–6/zi vs 3). Nu mai tolerează tăieței uscați, pâine veche. Va pierde probabil 5–10 kg permanent. Alcool — moderat, evitare înainte de culcare. Majoritatea revin la 80–90% din normal pe termen lung."),
        ("Q8: „Se transmite genetic? Trebuie să ne facem și noi analize?\"",
         "Cancerele esofagiene sunt în general NU ereditare direct. Factorii de risc (fumat, alcool, reflux, obezitate) sunt de mediu. Dar istoricul familial e util — unchi, mătuși, bunici cu cancere digestive trebuie raportați oncologului. Pentru tine, e util: screening endoscopic la 50 ani + evitare fumat + tratament reflux. Nu e panică — doar prudență."),
        ("Q9: „Ce se întâmplă dacă biopsia arată că nu e cancer?\"",
         "Scenariul C (rar dar posibil). Înseamnă: 1) E altceva tratabil — inflamație severă, displazie, tumoră benignă → tratament specific. 2) Biopsia a ratat tumora → rebiopsie urgentă. În oricare caz, nu ieșim din sistem medical până nu știm clar — biopsie negativă cu simptome și CT suspect nu înseamnă „totul e bine\", înseamnă „avem nevoie de mai multe date\"."),
        ("Q10: „Ce nu-i spunem tatei?\"",
         "NIMIC. Pacientul are dreptul legal și moral la informație completă. Discutăm cu el direct, folosim ghidul de prezentare CT, dar îi spunem adevărul. A ascunde: se observă (o simte), erodează încrederea (pierdem aliatul principal), contravine eticii medicale. Ce putem face: prezentăm treptat, cu compasiune, cu plan. Nu detalii crude brutal. Dar ONEST."),
    ]

    for q, a in faq_list:
        add_heading_bar(doc, q, level=3, color=COLOR_PRIMARY)
        add_paragraph(doc, a)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 12 — TABEL SCENARII
    # ============================================================

    add_heading_bar(doc, "12. Tabel sumarizat — 4 scenarii + plan", level=1, color=COLOR_PRIMARY)

    add_table(doc,
        headers=["Scenariu", "Biopsie", "Ascită", "Prob.", "Stadiu", "Strategie", "Durată", "Obiectiv"],
        rows=[
            ["A", "Malignă", "Benignă", "60–70%", "III", "FLOT + chirurgie + FLOT", "~6 luni", "CURATIV"],
            ["B", "Malignă", "Malignă", "15–25%", "IV", "Chimio + imunoterapie", "Continuu", "CONTROL 2–3 ani"],
            ["C", "Benignă", "Benignă", "5–10%", "N/A", "IPP/rezecție/rebiopsie", "Variabil", "REZOLVARE non-oncologică"],
            ["D", "Benignă", "Malignă", "1–3%", "N/A", "Rebiopsie + căutare primar", "2–4 săpt", "CLARIFICARE"],
        ],
        col_widths=[Inches(0.6), Inches(0.8), Inches(0.8), Inches(0.7), Inches(0.6), Inches(1.6), Inches(0.8), Inches(1.5)])

    add_paragraph(doc, "Probabilitățile sunt estimări [PROBABIL] bazate pe:")
    add_bullet(doc, "Aspect CT (proces expansiv infiltrativ, cantitate moderată ascită)")
    add_bullet(doc, "Context clinic (vârstă, istoric fumat, simptome, comorbidități)")
    add_bullet(doc, "Statistici populaționale pentru leziuni similare Siewert II cu ascită la CT")
    add_bullet(doc, "NU pe analiza specifică a cazului — aceea se face de oncolog cu toate datele")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 13 — CE FACI TU
    # ============================================================

    add_heading_bar(doc, "13. Ce trebuie să faci TU (Roland) — concret", level=1, color=COLOR_PRIMARY)

    add_heading_bar(doc, "ACUM (în următoarele 48h)", level=2, color=COLOR_CRIT)
    add_numbered(doc, "Citește acest document + Documente_Informative/GHID_CONSULT_ONCOLOG.md (pași operaționali)")
    add_numbered(doc, "NU prezentezi scenariile B/D familiei acum — sperie inutil. Prezintă doar „facem consult oncolog ca să stabilim planul\"")
    add_numbered(doc, "Sună la Dr. Noufal mâine dimineață (23.04) pentru recomandare oncolog")

    add_heading_bar(doc, "La primul consult oncolog — 7 întrebări prioritare", level=2, color=COLOR_PRIMARY)
    add_numbered(doc, "„Care dintre cele 4 scenarii considerați cel mai probabil pe baza CT + endoscopie?\"")
    add_numbered(doc, "„Ce investigație ne spune cel mai repede scenariul real — paracenteza, PET-CT, rebiopsia?\"")
    add_numbered(doc, "„Programăm paracenteza obligatorie chiar înainte de rezultatul biopsiei?\"")
    add_numbered(doc, "„Dacă scenariul A se confirmă, începem FLOT în cât timp?\"")
    add_numbered(doc, "„Dacă scenariul B, avem protocoale cu imunoterapie disponibile în 2026 la CNAS?\"")
    add_numbered(doc, "„Ce markeri moleculari vreți să cerem la biopsie (HER2, PD-L1, MSI, claudin-18.2)?\"")
    add_numbered(doc, "„Aveți acces la PET-CT? Dacă nu, unde putem face privat?\"")

    add_heading_bar(doc, "După ce ai toate datele", level=2, color=COLOR_OK)
    add_paragraph(doc,
        "Revii la oncolog pentru decizia finală de tratament. Acest document devine „istoric\" — planul real e cel dat de oncolog cu datele complete (biopsie + paracenteză + eventual PET-CT).")

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 14 — TIMELINE
    # ============================================================

    add_heading_bar(doc, "14. Timeline vizual — cronologia realistă", level=1, color=COLOR_PRIMARY)

    timeline_data = [
        ("APRILIE 2026", "Finalizat + programare"),
        ("17.04", "Endoscopie + colonoscopie + biopsie ✅"),
        ("20.04", "CT stadializare ✅"),
        ("22.04 ← EȘTI AICI", "Jamesi reluat post-CT ✅"),
        ("23.04", "Sună la Dr. Noufal pentru recomandare oncolog 🎯"),
        ("24–30.04", "Programare consult oncolog + pregătire dosar fizic"),
        ("24.04 – 01.05", "Rezultat biopsie (estimativ) ⏳"),
        ("", ""),
        ("MAI 2026", "Consulturi + investigații"),
        ("~05–10.05", "Primul consult oncolog (cu rezultat biopsie) 🎯"),
        ("~10–20.05", "Paracenteza + analize suplimentare ⏳"),
        ("~25.05", "Al doilea consult oncolog — decizie protocol final 🎯"),
        ("", ""),
        ("IUNIE 2026", "Început tratament"),
        ("~01.06", "ÎNCEPE TRATAMENTUL ACTIV — Ciclul 1 FLOT 🚀"),
        ("~15.06", "Ciclul 2 FLOT 💊"),
        ("", ""),
        ("IULIE–SEPTEMBRIE 2026", "Chimio preoperatorie"),
        ("", "4 cicluri, fiecare la 2 săpt → evaluare răspuns"),
        ("", ""),
        ("OCTOMBRIE 2026", "Chirurgie"),
        ("", "Pauză 4–6 săpt + CT/PET-CT repeat"),
        ("~15–20.10", "🏥 Chirurgie (scenariul curativ)"),
        ("", ""),
        ("NOIEMBRIE–DECEMBRIE 2026", "Recuperare + început chimio adjuvantă"),
        ("IANUARIE–APRILIE 2027", "Chimio adjuvantă (4 cicluri post-op)"),
        ("MAI 2027", "FINAL TRATAMENT ACTIV → urmărire periodică la 3 luni"),
    ]

    add_table(doc,
        headers=["Data", "Eveniment"],
        rows=timeline_data,
        col_widths=[Inches(2.2), Inches(4.3)])

    add_callout(doc, "[PROBABIL]",
        "Timeline-ul de mai sus e pentru scenariul A (cel mai probabil). Scenariul B schimbă: fără chirurgie, chimio + imunoterapie pe durată nedeterminată cu evaluări CT la 3 luni.",
        color=COLOR_INFO, bg_hex=COLOR_INFO_BG)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 15 — MESAJ 30 SEC
    # ============================================================

    add_heading_bar(doc, "15. Mesajul principal — în 30 secunde", level=1, color=COLOR_PRIMARY)

    add_quote(doc,
        "Cele 4 scenarii sunt ipoteze statistice bazate pe ce vezi la CT + endoscopie. Realitatea caz "
        "se află doar după biopsia (până pe 1 mai) + paracenteza ascitei (cerută de oncolog). "
        "Consultul oncolog e URGENT pentru că: (1) coordonează paracenteza, (2) economisește 2–3 săptămâni în programare, "
        "(3) accesul la investigații auxiliare (PET-CT, EUS, markeri) e mai rapid prin oncolog. "
        "Cel mai probabil scenariu (A) = stadiu III tratabil curativ cu FLOT. Cel mai agresiv (B) = stadiu IV dar cu "
        "opțiuni moderne de control pentru 2–3 ani cu calitate. Scenariile C și D sunt posibile dar rare.",
        color_left=COLOR_PRIMARY, bg_hex="DBEAFE")

    add_mixed_paragraph(doc, [
        {"text": "Esențialul: ", "bold": True, "size": 14, "color": COLOR_CRIT},
        {"text": "nu aștepți biopsia ca să programezi oncologul. ", "size": 13},
        {"text": "Suni mâine.", "bold": True, "size": 14, "color": COLOR_CRIT},
    ], alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)

    # ============================================================
    # SECTIUNEA 16 — GLOSSAR
    # ============================================================

    add_heading_bar(doc, "16. Glossar — termeni medicali explicați simplu", level=1, color=COLOR_PRIMARY)

    glossar_data = [
        ("Adenocarcinom", "Cancer din celulele glandulare", "„Fabrica de secreție\" dăunată"),
        ("Ascită", "Apă acumulată anormal în burtă", "Acvariu care începe să se umple"),
        ("Biopsie", "Bucățică mică de țesut pentru analiză", "Probă din cărămida zidului"),
        ("Carcinom scuamocelular", "Cancer din celulele pavimentoase", "„Acoperiș\" celular dăunat"),
        ("Carcinomatoză peritoneală", "Celule tumorale pe învelișul burții", "Mucegai peste toți pereții subsolului"),
        ("Chimioterapie", "Medicamente care ucid celule cu diviziune rapidă", "Ploaie de substanțe care afectează ce crește repede"),
        ("CPS (PD-L1)", "Scor care măsoară cât PD-L1 e pe tumoră", "Mărimea „steagului alb\" al tumorii"),
        ("CT", "Computer tomograf — imagini 3D cu raze X", "Drona care zboară prin casă"),
        ("Disfagie", "Dificultate la înghițire", "Scurgere casei înfundată"),
        ("Eco-endoscopie (EUS)", "Endoscop cu ultrasunete la vârf", "Urechea lipită de peretele interior"),
        ("Endoscopie", "Cameră pe tub flexibil introdusă pe gură", "Meșter cu lanternă pe hol"),
        ("Ezofago-gastrectomie", "Scoatere chirurgicală esofag + parte stomac", "Demolare coridor + reconstrucție"),
        ("FLOT", "Protocol chimio (5-FU + Leucov + Oxalipl + Docetax)", "Pachet de artilerie anti-tumorală"),
        ("Ganglioni limfatici", "Noduli care filtrează limfatul", "Stații de control în autostrăzi"),
        ("HER2", "Receptor de creștere pe celule", "„Antenă de multiplicare\""),
        ("Histopatologie", "Analiza microscopică a țesuturilor", "Laborator specializat în tipul cărămizii"),
        ("Hipoalbuminemie", "Nivel scăzut albumină în sânge", "Pompa de presiune slăbită"),
        ("Imunoterapie", "Tratament care activează sistemul imunitar", "Instruire armată să recunoască intruși noi"),
        ("Infiltrativ", "Care pătrunde prin țesuturi fără margini clare", "Cerneală care se răspândește pe hârtie umedă"),
        ("Laparoscopie", "Chirurgie minim-invazivă cu cameră prin burtă", "Cameră video mică în acvariu"),
        ("Malign", "Cancer, cu potențial de răspândire", "Mucegai care „prinde\" și se înmulțește"),
        ("Metastază", "Celulă tumorală migrată la alt organ", "Spore de mucegai în altă cameră"),
        ("MSI-High", "Instabilitate genetică — multe mutații", "Fotocopiator stricat care face greșeli"),
        ("Neoadjuvant", "Tratament PREOPERATOR", "Pregătirea cartierului înainte de demolare"),
        ("Neutropenie", "Leucocite (neutrofile) scăzute", "Armată cu lipsuri în rânduri"),
        ("Paliativ", "Tratament pentru control, nu vindecare", "Întreținere clădire vs. renovare completă"),
        ("Paracenteza", "Puncție cu ac în burtă pentru lichid", "Pipetă care testează apa"),
        ("Pembrolizumab", "Imunoterapie PD-L1 (Keytruda)", "„Scoate steagul alb\" de pe tumoră"),
        ("Peritoneu", "Folia care îmbracă organele din burtă", "Celofanul care împachetează brânzeturile"),
        ("PET-CT", "Scanner care detectează metabolism + anatomie", "Drona cu cameră termică"),
        ("Polip", "Excrescență mică din mucoasă", "„Ciupercă\" mică pe perete"),
        ("Rezecție", "Scoatere chirurgicală a unui țesut", "Demolare zonă afectată"),
        ("Siewert I/II/III", "Clasificare cancere joncțiune eso-gastrică", "Etichetare locație exactă"),
        ("Stadializare (TNM)", "Evaluarea cât de extins e cancerul", "Inventar complet al pagubei"),
        ("Stenoză", "Îngustare / strâmtare a unui canal", "Coridor blocat cu moloz"),
        ("Stent esofagian", "Proteză metalică care dilată o zonă îngustată", "Bară de susținere într-un coridor surpat"),
        ("Trastuzumab", "Medicament țintit HER2+ (Herceptin)", "„Cheie\" care blochează antena creștere"),
        ("Tumor board", "Reuniune multidisciplinară decizie", "Consiliu de arhitecți pentru plan"),
        ("Zolbetuximab", "Medicament țintit claudin-18.2+ (Vyloy)", "Cheie nouă (2024) pentru lacăt specific"),
    ]

    add_table(doc,
        headers=["Termen medical", "Explicație simplă", "Analogia"],
        rows=glossar_data,
        col_widths=[Inches(1.8), Inches(2.4), Inches(2.2)])

    add_page_break(doc)

    # ============================================================
    # SURSE + TRANSPARENTA
    # ============================================================

    add_heading_bar(doc, "Surse citate", level=2, color=COLOR_PRIMARY)
    sources = [
        "Raport CT 20.04.2026 — Dosar_Medical/2026-04-20_ct_torace_abdomen_pelvis.json",
        "Ghid complet cancer esofagian 19.04.2026 — Dosar_Medical/rapoarte_generate/2026-04-19_ghid_cancer_esofagian_complet.docx",
        "Studiu FLOT4 — Al-Batran et al., Lancet 2019",
        "Keynote-590 — Sun et al., Lancet 2021 (pembrolizumab + chemo în cancer esofagian avansat)",
        "Keynote-811 — Janjigian et al., Lancet 2023 (trastuzumab + pembrolizumab în HER2+)",
        "SPOTLIGHT — Shitara et al., Lancet 2023 (zolbetuximab pentru claudin-18.2+)",
        "ESPEN Guidelines — European Society for Clinical Nutrition and Metabolism, 2021–2023",
        "AJCC Cancer Staging Manual, 8th Edition",
        "NCCN Clinical Practice Guidelines Esophageal and Esophagogastric Junction Cancers V1.2025",
        "ESMO Clinical Practice Guidelines Esophageal Cancer 2022",
        "CONTEXT_MEDICAL.md v1.2.1",
        "Ghid operațional GHID_CONSULT_ONCOLOG.md (același folder)",
    ]
    for src in sources:
        add_bullet(doc, src)

    add_heading_bar(doc, "Ce NU am inclus (transparență)", level=2, color=COLOR_MUTED)
    transparency = [
        "[NEGASIT] Probabilități scenarii specifice pentru cazul tatei — doar oncologul cu toate datele poate calcula",
        "[INCERT] Durata exactă de supraviețuire în fiecare scenariu — variază enorm individual",
        "[INCERT] Răspunsul specific al tatei la FLOT — depinde de mutații tumorale, status general, comorbidități",
        "[NEGASIT] Disponibilitatea exactă a trialurilor clinice la centrul unde alegeți — se verifică la primul consult",
        "[INCERT] Costuri exacte pentru PET-CT / zolbetuximab în 2026 — verifică la centrul ales",
        "[NEGASIT] Detalii ale accesului PNO pentru fiecare imunoterapie — oncologul are datele curente",
    ]
    for t in transparency:
        add_bullet(doc, t)

    # ============================================================
    # DISCLAIMER FINAL
    # ============================================================

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("NU înlocuiește consultul medical.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = COLOR_CRIT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(
        "Acest document e material explicativ pentru înțelegere de către familie a situației. "
        "Deciziile terapeutice aparțin EXCLUSIV oncologului digestiv cu toate datele clinice."
    )
    r2.italic = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = COLOR_MUTED

    return doc


# =========================================================================
# MAIN
# =========================================================================

if __name__ == "__main__":
    output_path = Path(__file__).parent / "2026-04-22_explicatie_consult_oncolog_scenarii.docx"
    print(f"Generare DOCX la: {output_path}")
    doc = build_document()
    doc.save(str(output_path))
    print(f"DOCX generat cu succes. Dimensiune: {output_path.stat().st_size / 1024:.1f} KB")
