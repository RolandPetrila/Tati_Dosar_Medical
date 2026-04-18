# -*- coding: utf-8 -*-
"""
Generator raport DOCX: Reactii adverse Jamesi + Triplixam.
Sursa de adevar: structurile Python de mai jos (fiecare claim cu marcaj certitudine).
Ruleaza: python generate_reactii_adverse_docx.py
Output: raport_reactii_adverse_jamesi_triplixam_2026-04-18.docx in acelasi folder
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path


RED = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREEN = RGBColor(0x0B, 0x6E, 0x4F)
GRAY = RGBColor(0x60, 0x60, 0x60)
BLACK = RGBColor(0x00, 0x00, 0x00)


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_paragraph(doc, text="", style=None, bold=False, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        add_run(p, text, bold=bold)
    return p


def add_heading_custom(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    if color:
        for run in h.runs:
            run.font.color.rgb = color
    return h


def add_marker(paragraph, marker_type):
    """Adauga marcaj certitudine colorat la inceputul unui paragraf."""
    colors = {
        "CERT": GREEN,
        "PROBABIL": ORANGE,
        "INCERT": ORANGE,
        "NEGASIT": GRAY,
    }
    color = colors.get(marker_type, BLACK)
    run = paragraph.add_run(f"[{marker_type}] ")
    run.bold = True
    run.font.color.rgb = color


def add_bullet(doc, items, style="List Bullet"):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(item, tuple):
            marker, text = item
            add_marker(p, marker)
            p.add_run(text)
        else:
            p.add_run(item)


def add_table(doc, headers, rows, header_color="4F6D7A", first_col_bold=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"

    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_color)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows):
        cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            if c_idx == 0 and first_col_bold:
                run.bold = True
    return tbl


def add_callout(doc, title, body, color=ORANGE):
    """Cutie evidentiata pentru avertismente / exemple."""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, "FFF4E6" if color == ORANGE else ("FEECEC" if color == RED else "E8F5E9"))
    p1 = cell.paragraphs[0]
    run = p1.add_run(title)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(11)
    p2 = cell.add_paragraph()
    p2.add_run(body).font.size = Pt(10)


def build_document():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # === COPERTA ===
    title = doc.add_heading("RAPORT MEDICAL", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Reacții adverse și precauții pentru Jamesi și Triplixam")
    run.bold = True
    run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        "Pacient: Petrilă Viorel-Mihai (66 ani) | Data raport: 18 aprilie 2026\n"
        "Autor: Claude Code (Opus 4.7) | Versiune: 1.0"
    ).font.size = Pt(10)

    doc.add_paragraph()

    # === CUM SE CITESTE ===
    add_heading_custom(doc, "Cum se citește acest document", level=1)
    p = doc.add_paragraph(
        "Raportul e scris pentru un cititor fără pregătire medicală. "
        "Fiecare afirmație are un marcaj de certitudine în fața ei:"
    )
    add_bullet(doc, [
        ("CERT", " — confirmat din prospectul oficial al medicamentului (Rezumatul Caracteristicilor Produsului, RCP/SmPC), ghid clinic major sau studiu peer-reviewed."),
        ("PROBABIL", " — susținut de literatura medicală generală, dar nu explicit în RCP-ul acestui medicament."),
        ("INCERT", " — date conflictuale sau insuficiente, mecanism ipotezat."),
        ("NEGASIT", " — am căutat în surse oficiale și nu am găsit informația."),
    ])
    doc.add_paragraph()

    add_callout(
        doc,
        "⚠ Atenționare importantă",
        "Acest raport NU înlocuiește consultul medical. Orice suspiciune de reacție adversă se "
        "raportează la medicul curant. În cazurile marcate „URGENȚE”, sunați 112 imediat.",
        color=RED,
    )
    doc.add_paragraph()
    doc.add_page_break()

    # === REZUMAT EXECUTIV ===
    add_heading_custom(doc, "Rezumat în 5 puncte", level=1)
    add_bullet(doc, [
        "Pacientul ia două medicamente combinate: Jamesi (sitagliptin + metformin, pentru diabet) și Triplixam (perindopril + indapamidă + amlodipină, pentru tensiune).",
        "Ambele combinații au profil de siguranță bine cunoscut; cele mai multe reacții adverse sunt minore și se rezolvă singure sau prin ajustarea medicamentului.",
        "Există trei riscuri grave de cunoscut: (1) acidoză lactică la metformin + contrast iodat (de aceea se oprește pre-CT), (2) angioedem (umflare la față) la IECA (perindopril), cu risc crescut în combinație cu sitagliptin, (3) scăderea potasiului în sânge de la diuretic (indapamidă).",
        "Pentru CT luni 20.04: STOP Jamesi sâmbătă 17:00, hidratare duminică, întrebare pentru radiolog despre Triplixam.",
        "În caz de umflare bruscă la față/buze/limbă, durere abdominală severă, respirație anormală, slăbiciune severă → 112 imediat.",
    ])
    doc.add_paragraph()
    doc.add_page_break()

    # === PARTEA I: JAMESI ===
    add_heading_custom(doc, "PARTEA I — Jamesi 50 mg / 1000 mg", level=1, color=RGBColor(0x1F, 0x4E, 0x79))
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "Jamesi este denumirea comercială românească a combinației "
    )
    p.add_run("sitagliptin 50 mg + clorhidrat de metformin 1000 mg").bold = True
    p.add_run(
        ". Produsul de referință european este Janumet® (MSD). Cele două substanțe active "
        "sunt aceleași; informațiile despre reacții adverse se bazează pe RCP Janumet "
        "(Electronic Medicines Compendium UK, actualizat 2024)."
    )

    p = doc.add_paragraph()
    p.add_run("Scop: ").bold = True
    p.add_run("controlul glicemiei la pacienții cu diabet zaharat tip 2.")

    # --- METFORMIN ---
    add_heading_custom(doc, "1. Componenta METFORMIN (1000 mg)", level=2)
    p = doc.add_paragraph(
        "Metforminul e cel mai vechi și cel mai folosit antidiabetic oral. "
        "Funcționează prin două mecanisme: reduce producția de glucoză în ficat și "
        "crește sensibilitatea mușchilor la insulină."
    )

    add_heading_custom(doc, "Reacții adverse FOARTE FRECVENTE (mai mult de 1 din 10 pacienți)", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Janumet, secțiunea 4.8.")

    add_bullet(doc, [
        "Greață",
        "Vărsături",
        "Diaree",
        "Dureri de stomac (abdomen)",
        "Scăderea apetitului",
        "Gust metalic în gură",
    ])

    add_callout(
        doc,
        "Exemplu concret — primele săptămâni",
        "În primele 1-2 săptămâni de la începerea Jamesi, mulți pacienți au scaune moi sau "
        "dureri de stomac după masă. La majoritatea, simptomele dispar în 2-4 săptămâni, pe măsură "
        "ce corpul se obișnuiește. Trucul: luați Jamesi cu masa, nu pe stomacul gol. "
        "Dacă persistă peste o lună sau devin severe, sunați medicul — poate reduce doza temporar.",
    )

    add_heading_custom(doc, "Reacții adverse FRECVENTE (1-10 din 100 pacienți)", level=3)
    rows = [
        ["Scăderea vitaminei B12", "Oboseală, amețeli, furnicături la mâini/picioare",
         "Apare după luni/ani de tratament. Medicul controlează B12 la 2-3 ani."],
        ["Hipoglicemie", "Transpirații, tremurături, foame bruscă, amețeli",
         "DOAR când Jamesi se combină cu sulfoniluree sau insulină. La Jamesi singur = risc foarte mic."],
    ]
    add_table(doc, ["Reacție", "Simptome", "Context"], rows)

    add_heading_custom(doc, "Reacții adverse RARE dar GRAVE", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    run = p.add_run("Acidoza lactică — URGENȚĂ MEDICALĂ.")
    run.bold = True
    run.font.color.rgb = RED

    p = doc.add_paragraph(
        "Se întâmplă foarte rar (aprox. 3-10 cazuri la 100.000 pacienți pe an). "
        "Dar când apare, e o urgență: corpul acumulează acid lactic în sânge, sângele devine "
        "prea acid, organele încep să cedeze."
    )

    p = doc.add_paragraph()
    p.add_run("Când crește riscul:").bold = True
    add_bullet(doc, [
        "Funcție renală proastă (rinichii nu mai elimină metforminul)",
        "Deshidratare severă (diaree prelungită, vărsături, febră mare, transpirație în exces)",
        "Contrast iodat pentru CT (fără pauză de 48h) — de aceea se oprește Jamesi pre-CT",
        "Abuz de alcool",
        "Infecție severă / sepsis",
        "Insuficiență cardiacă acută",
    ])

    p = doc.add_paragraph()
    p.add_run("Simptome (SUNAȚI 112):").bold = True
    add_bullet(doc, [
        "Respirație rapidă și adâncă, ca după efort",
        "Greață severă cu vărsături",
        "Durere musculară / abdominală foarte puternică",
        "Somnolență severă, dificultate de a gândi clar",
        "Corp rece, temperatură sub normal",
    ])

    add_callout(
        doc,
        "De ce se oprește Jamesi 48h înainte de CT",
        "Contrastul iodat poate să supraîncarce temporar rinichii. Dacă metforminul e în "
        "circulație în acel moment și rinichii nu-l elimină eficient, metforminul se acumulează "
        "→ risc de acidoză lactică. Oprirea cu 48h înainte = organismul are timp să elimine "
        "metforminul. Reluarea după CT se face doar cu creatinină (funcție renală) confirmată normală.",
        color=RED,
    )

    # --- SITAGLIPTIN ---
    doc.add_paragraph()
    add_heading_custom(doc, "2. Componenta SITAGLIPTIN (50 mg)", level=2)
    p = doc.add_paragraph(
        "Sitagliptinul aparține clasei „inhibitori DPP-4” (incretin-mimetice). Ajută pancreasul "
        "să elibereze mai multă insulină DOAR când glicemia crește (după mese). "
        "Prin el însuși, NU cauzează hipoglicemie — singura excepție e combinația cu sulfoniluree/insulină."
    )

    add_heading_custom(doc, "Reacții adverse FRECVENTE (1-10 din 100)", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Janumet, secțiunea 4.8.")
    add_bullet(doc, [
        "Dureri de cap",
        "Infecții respiratorii ușoare (răceală, nas înfundat, durere în gât)",
    ])

    add_heading_custom(doc, "Reacții adverse RARE dar SERIOASE", level=3)

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    run = p.add_run("Pancreatita acută.")
    run.bold = True
    run.font.color.rgb = RED
    p.add_run(
        " Inflamația severă a pancreasului. Avertizare FDA. Raportări postmarketing inclusiv cazuri "
        "fatale (hemoragice sau necrotice)."
    )
    p = doc.add_paragraph()
    p.add_run("Simptome (URGENȚE):").bold = True
    add_bullet(doc, [
        "Durere foarte puternică în partea de SUS a abdomenului (sub coaste, centru)",
        "Durerea iradiază în spate",
        "Greață și vărsături persistente",
        "Uneori febră",
    ])
    p = doc.add_paragraph(
        "Dacă apar aceste simptome, OPRIȚI Jamesi și mergeți la urgențe. "
        "Pancreatita netratată poate fi fatală."
    )

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    run = p.add_run("Reacții alergice severe.")
    run.bold = True
    p.add_run(
        " Anafilaxie (reacție alergică generală), angioedem (umflare la față/buze/limbă), "
        "sindrom Stevens-Johnson (piele cu vezicule și exfoliere). De obicei în primele 3 luni, "
        "dar posibil oricând. URGENȚE imediat."
    )

    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "Insuficiență renală acută. Durere articulară severă (avertizare FDA 2015). "
        "Pemfigoid bulos (pufoase mari pe piele — boală autoimună rară)."
    )

    p = doc.add_paragraph()
    add_marker(p, "INCERT")
    p.add_run(
        "Insuficiență cardiacă. Un alt medicament din aceeași clasă (saxagliptin) a arătat risc "
        "crescut de spitalizare pentru insuficiență cardiacă în trialul SAVOR-TIMI 53. Trialul "
        "TECOS pe sitagliptin a fost neutru (fără risc crescut). RCP recomandă prudență la "
        "pacienți cu factori de risc cardiac. Pacientul actual are antecedent de stent coronarian "
        "(2012) — această întrebare e valabilă pentru cardiolog/diabetolog."
    )
    doc.add_page_break()

    # === PARTEA II: TRIPLIXAM ===
    add_heading_custom(doc, "PARTEA II — Triplixam 10 mg / 2.5 mg / 5 mg", level=1, color=RGBColor(0x1F, 0x4E, 0x79))
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "Triplixam combină trei medicamente pentru tensiune: "
    )
    p.add_run("perindopril arginine 10 mg + indapamidă 2.5 mg + amlodipină 5 mg").bold = True
    p.add_run(
        ". Produs de Servier. RCP versiunea 06.2021 (surse suplimentare 2023). "
        "Aprobat de ANMDMR pentru hipertensiunea esențială când pacientul e deja controlat "
        "cu combinația perindopril+indapamidă și cu amlodipina separat."
    )

    # --- PERINDOPRIL ---
    add_heading_custom(doc, "3. Componenta PERINDOPRIL (IECA)", level=2)
    p = doc.add_paragraph(
        "Perindoprilul e un „inhibitor al enzimei de conversie a angiotensinei” (IECA / ACE-inhibitor). "
        "Scade tensiunea prin blocarea producției unui hormon (angiotensina II) care strânge vasele "
        "de sânge. Vasele se lărgesc → tensiunea scade."
    )

    add_heading_custom(doc, "Reacția adversă cea mai tipică: TUSEA SEACĂ", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "Frecvență comună (1-10 din 100 pacienți, literatura citează 5-20%). Tusea apare "
        "pentru că perindoprilul crește bradikinina, o substanță care irită căile respiratorii."
    )

    add_callout(
        doc,
        "Exemplu concret — tusea de IECA",
        "„Bunicul tușește seara de 3 săptămâni. Nu are răceală, nu e umezeală. "
        "Tusea e seacă, fără spută, iritantă, deranjantă la culcare.” → "
        "Cel mai probabil e tuse de perindopril. Nu se ignoră — se anunță medicul. "
        "Tusea dispare complet în 1-4 săptămâni după oprirea perindoprilului. "
        "Dacă e problematică, medicul poate schimba cu un „sartan” (losartan, valsartan, "
        "telmisartan etc.) — aceeași clasă farmacologică, dar fără efectul pe bradikinină.",
    )

    add_heading_custom(doc, "Reacția adversă cea mai GRAVĂ: ANGIOEDEMUL", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    run = p.add_run("Raritate: 1-10 din 1000 (uncommon). Dar POTENȚIAL FATAL.")
    run.bold = True
    run.font.color.rgb = RED

    p = doc.add_paragraph("Ce e angioedemul: umflare bruscă a țesuturilor sub piele. "
                          "Cel mai periculos când prinde:")
    add_bullet(doc, [
        "Buzele și limba — dificultate de a vorbi sau înghiți",
        "Laringele și gâtul — dificultate de a respira, sufocare",
        "Intestinul (rar) — durere abdominală severă fără explicație",
    ])

    p = doc.add_paragraph()
    p.add_run("Când apare:").bold = True
    add_bullet(doc, [
        "Frecvent în primele 3 luni de tratament, DAR posibil oricând",
        "Mai frecvent la pacienți de rasă neagră",
        "Risc CRESCUT în combinație cu alte medicamente — VEZI SECȚIUNEA V (interacțiuni)",
    ])

    p = doc.add_paragraph()
    p.add_run("Ce faceți:").bold = True
    p.add_run(" orice umflare bruscă la față/buze/limbă → 112 IMEDIAT. ")
    p.add_run("Nu așteptați „să vedeți dacă trece”.")

    add_heading_custom(doc, "Alte reacții adverse ale perindoprilului", level=3)
    rows = [
        ["Amețeli, durere de cap, furnicături", "Frecvent", "Mai ales la începutul tratamentului."],
        ["Hipotensiune (tensiune prea mică)", "Frecvent", "Mai ales la ridicare bruscă în picioare — risc de cădere la vârstnici."],
        ["Hiperkaliemie (potasiu mare în sânge)", "Mai puțin frecvent", "Reversibilă la oprirea medicamentului."],
        ["Erupții pe piele, mâncărime", "Frecvent", "Medicul decide dacă trebuie oprit."],
        ["Insuficiență renală acută", "Foarte rar", "Risc crescut: stenoză artera renală, deshidratare, AINS (ibuprofen)."],
        ["Neutropenie / agranulocitoză", "Foarte rar", "Risc crescut: insuficiență renală + boli autoimune + imunosupresoare."],
        ["Afectare hepatică", "Foarte rar", "Icter, oprire medicament dacă apare."],
    ]
    add_table(doc, ["Reacție", "Frecvență RCP", "Context / Ce să faceți"], rows)

    # --- INDAPAMIDA ---
    doc.add_paragraph()
    add_heading_custom(doc, "4. Componenta INDAPAMIDĂ (diuretic)", level=2)
    p = doc.add_paragraph(
        "Indapamida e un diuretic „tiazid-like”. Ajută rinichii să elimine mai multă sare "
        "și apă → volumul de sânge scade → tensiunea scade."
    )

    add_heading_custom(doc, "Reacția adversă cea mai importantă: SCĂDEREA POTASIULUI", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    run = p.add_run("Hipokaliemia")
    run.bold = True
    p.add_run(" (potasiu scăzut) e riscul-cheie al oricărui diuretic.")

    p = doc.add_paragraph()
    p.add_run("Simptome:").bold = True
    add_bullet(doc, [
        "Slăbiciune musculară, senzație de „moale”",
        "Crampe, mai ales la gambe",
        "Palpitații neregulate",
        "Oboseală severă, apatie",
    ])

    p = doc.add_paragraph()
    p.add_run("Pericol:").bold = True
    add_bullet(doc, [
        "Aritmii cardiace severe (bătăi neregulate)",
        "„Torsada vârfurilor” — aritmie potențial fatală (mai ales dacă se combină cu alte medicamente care prelungesc QT-ul, ex: anumite antibiotice, antiaritmice)",
        "Sensibilitate crescută la digoxină (dacă e dat)",
    ])

    p = doc.add_paragraph()
    p.add_run("Factori care cresc riscul la acest pacient:").bold = True
    add_bullet(doc, [
        "Vârsta 66 ani (riscul crește vizibil peste 70)",
        "Boală cardiacă ischemică (stent 2012) — inimă mai sensibilă la aritmii",
        "Diaree/vărsături/deshidratare — elimină mai mult potasiu",
        "Diabetul în sine poate masca simptomele de hipokaliemie",
    ])

    p = doc.add_paragraph()
    p.add_run("Cum se previne:").bold = True
    p.add_run(
        " medicul dozează potasiu în sânge la început de tratament (prima săptămână) "
        "și periodic. Dacă valoarea scade sub 3.4 mmol/L, se corectează (supliment de potasiu, "
        "reducere dozei de indapamidă, sau adăugarea unui diuretic „care economisește potasiu”)."
    )

    add_heading_custom(doc, "Alte reacții adverse ale indapamidei", level=3)
    rows = [
        ["Hiponatremie (sodiu scăzut)", "Frecvență necunoscută",
         "Risc la vârstnici. Simptome inițial: somnolență, confuzie. Sever: convulsii. Analiza Na în sânge periodic."],
        ["Hiperuricemie / puseuri de gută", "Frecvent",
         "Crește acidul uric în sânge. Dacă pacientul are istoric de gută, semnalați medicului."],
        ["Reacții alergice pielii, erupții maculo-papulare", "Frecvent", "De obicei ușor; semnalați medicului."],
        ["Fotosensibilitate (piele sensibilă la soare)", "Rar la indapamidă, frecvent la tiazide",
         "Evitați soarele puternic 11:00-16:00 vara; folosiți SPF 50+ pe față."],
        ["Pancreatită", "Foarte rar", "Durere abdominală intensă → urgențe (cum s-a descris la sitagliptin)."],
        ["Stevens-Johnson / necroliza epidermică toxică", "Foarte rar", "Reacții fatale ale pielii; orice erupție severă → urgențe."],
        ["Agranulocitoză, aplazie medulară", "Foarte rar", "Scăderea celulelor albe → risc infecție."],
    ]
    add_table(doc, ["Reacție", "Frecvență RCP", "Context / Ce să faceți"], rows)

    # --- AMLODIPINA ---
    doc.add_paragraph()
    add_heading_custom(doc, "5. Componenta AMLODIPINA (blocant de canal de calciu)", level=2)
    p = doc.add_paragraph(
        "Amlodipina blochează canalele de calciu în pereții vaselor de sânge. "
        "Fără calciu, vasele se relaxează → tensiunea scade. "
        "Nu afectează direct ritmul inimii."
    )

    add_heading_custom(doc, "Reacția adversă cea mai tipică: EDEM LA GLEZNE", level=3)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "Foarte frecvent (peste 1 din 10 pacienți la doze de 10 mg; mai puțin frecvent la 5 mg, "
        "doza din Triplixam). Mecanism: amlodipina dilată capilarele pe care le avem în glezne "
        "→ lichid iese din vase în țesut → umflare."
    )

    add_callout(
        doc,
        "Exemplu concret — edemul de amlodipină",
        "„La sfârșitul zilei, gleznele bunicului sunt umflate, pantofii sunt strâmți. "
        "Dimineața sunt normale.” → E edem de amlodipină, nu insuficiență cardiacă sau rinichi.\n\n"
        "Ce ajută: ridicare picioare seara 20-30 min, ciorapi elastici cu compresie ușoară. "
        "Dacă e invalidant (face greu mersul), medicul poate schimba amlodipina cu un alt "
        "medicament din clasă diferită (ex: sartan).",
    )

    add_heading_custom(doc, "Alte reacții adverse ale amlodipinei", level=3)
    rows = [
        ["Flushing (roșeață a feței, senzație de căldură)", "Frecvent", "Mai ales primele săptămâni; se reduce în timp."],
        ["Palpitații", "Frecvent", "Inima bate mai repede temporar; cedează de obicei."],
        ["Amețeală, somnolență, oboseală", "Frecvent", "Mai ales la începutul tratamentului."],
        ["Durere abdominală, greață", "Frecvent", "Dispare cu timpul."],
        ["Hiperplazie gingivală (îngroșarea gingiilor)", "Foarte rar", "Apare după luni/ani. Igienă orală atentă previne."],
        ["Hiperglicemie (crește glicemia)", "Foarte rar", "Relevant pentru pacientul diabetic — monitorizare glicemie."],
        ["Hepatită", "Foarte rar", "Enzime hepatice crescute; medicul face control."],
    ]
    add_table(doc, ["Reacție", "Frecvență RCP", "Context / Ce să faceți"], rows)
    doc.add_page_break()

    # === PARTEA III: INTERACTIUNI SPECIFICE ===
    add_heading_custom(doc, "PARTEA III — Interacțiuni specifice la acest pacient",
                       level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    # Interaction A
    add_heading_custom(doc, "A. Jamesi (sitagliptin) + Triplixam (perindopril) → risc de angioedem",
                       level=2, color=RED)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Triplixam, secțiunea 4.5 interacțiuni.")

    p = doc.add_paragraph()
    p.add_run("Mecanism:").bold = True
    p.add_run(
        " Sitagliptinul blochează o enzimă (DPP-4) care în mod normal degradează substanța P "
        "(un peptid vasoactiv). Perindoprilul blochează ECA, care degradează bradikinina. "
        "Ambele substanțe acumulate → risc crescut de umflare bruscă (angioedem), în special "
        "la față, buze, limbă."
    )

    p = doc.add_paragraph()
    p.add_run("Risc practic:").bold = True
    p.add_run(" rar, dar documentat. Citat în RCP-ul Triplixam ca interacțiune care „necesită atenție”.")

    p = doc.add_paragraph()
    p.add_run("Ce să urmăriți:").bold = True
    p.add_run(" orice umflare bruscă a feței, buzelor, limbii, senzație de limbă „grea”, "
              "dificultate de a respira sau înghiți.")

    p = doc.add_paragraph()
    p.add_run("Ce să faceți:").bold = True
    p.add_run(" URGENȚE IMEDIAT (112). Nu așteptați.")

    add_callout(
        doc,
        "Notă: combinația NU e contraindicată",
        "Mulți pacienți iau sitagliptin și perindopril împreună fără probleme. "
        "Interacțiunea e o atenționare de urmărire, nu o interdicție. "
        "Medicul a decis această combinație după evaluarea beneficiilor (control diabet + tensiune) "
        "vs. risc.",
    )

    # Interaction B
    doc.add_paragraph()
    add_heading_custom(doc, "B. Indapamidă + contrast iodat la CT → risc de insuficiență renală acută",
                       level=2, color=ORANGE)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Triplixam, secțiunea 4.5.")

    p = doc.add_paragraph()
    p.add_run("Mecanism:").bold = True
    p.add_run(
        " Indapamida elimină apă prin urină. Dacă pacientul e deshidratat când primește "
        "contrast iodat, rinichii primesc mai puțin sânge → contrastul devine mai toxic "
        "pentru rinichi."
    )

    p = doc.add_paragraph()
    p.add_run("De aceea:").bold = True
    add_bullet(doc, [
        "Hidratare activă duminică 19.04 (1.5-2 L apă plată, dacă tolerabilitatea cardiacă permite)",
        "Întrebare explicită radiologului: mențin integral Triplixam sau omit doza din ziua CT?",
        "Nu opriți Triplixam pe cont propriu — decizia aparține radiologului/medicului curant",
    ])

    # Interaction C
    doc.add_paragraph()
    add_heading_custom(doc, "C. Metformin + contrast iodat → risc acidoză lactică",
                       level=2, color=ORANGE)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Janumet + RCP Triplixam (ambele menționează).")

    p = doc.add_paragraph(
        "De aceea: STOP Jamesi sâmbătă 18.04 ora 17:00 (H-48 înainte de CT luni). "
        "Reluare miercuri 22.04 ora 17:00 (H+48 după CT) DOAR după creatinină de control "
        "confirmată normală."
    )

    # Interaction D
    doc.add_paragraph()
    add_heading_custom(doc, "D. Aspenter (aspirină 75 mg) + Triplixam — COMBINAȚIE SIGURĂ",
                       level=2, color=GREEN)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run("Sursă: RCP Triplixam, secțiunea 4.5.")

    p = doc.add_paragraph(
        "RCP menționează atenționare la „acetilsalicilic la doze antiinflamatoare” (> 1-3 g/zi). "
        "La doza cardioprotectoare de 75 mg/zi folosită la acest pacient, interacțiunea clinică "
        "e minimă. Aspenter 75 mg + Triplixam e o combinație standard la pacienți post-stent. "
        "NU necesită modificări înaintea CT-ului. Aspenter se continuă normal."
    )

    doc.add_page_break()

    # === PARTEA IV: CHECKLIST ===
    add_heading_custom(doc, "PARTEA IV — Checklist pentru familie",
                       level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    add_heading_custom(doc, "Ce să urmăriți zilnic", level=2)

    p = doc.add_paragraph()
    p.add_run("La Jamesi:").bold = True
    add_bullet(doc, [
        "Greață severă care nu cedează când ia medicamentul cu masa → anunțați medicul",
        "Diaree persistentă peste 3 zile → anunțați medicul (risc deshidratare → risc acidoza lactică)",
        "Durere puternică în stomac / partea de sus a abdomenului → URGENȚE (pancreatită)",
        "Respirație rapidă + confuzie + dureri musculare severe → URGENȚE (acidoza lactică)",
        "Mâini/picioare amorțite de luni de zile → dozare vitamină B12",
    ])

    p = doc.add_paragraph()
    p.add_run("La Triplixam:").bold = True
    add_bullet(doc, [
        "Umflare la față / buze / limbă → URGENȚE IMEDIAT (angioedem, 112)",
        "Tuse seacă persistentă peste 2 săptămâni fără alt motiv → anunțați medicul",
        "Amețeli severe la ridicare din pat (risc de cădere) → medicul evaluează ajustare doză",
        "Glezne foarte umflate seara → acceptat, dar semnalați dacă e invalidant",
        "Slăbiciune musculară / crampe severe → medicul verifică potasiu și sodiu",
        "Scăderea cantității de urină + oboseală severă → insuficiență renală acută, URGENȚE",
        "Reacție severă a pielii la soare → protejați, semnalați",
    ])

    doc.add_paragraph()
    add_heading_custom(doc, "Când SUNAȚI 112 IMEDIAT", level=2, color=RED)
    add_bullet(doc, [
        "Umflare bruscă la față, buze, limbă, sau gât",
        "Dificultate severă de respirație",
        "Durere toracică severă",
        "Durere abdominală foarte puternică cu vărsături",
        "Pierdere de conștiență sau aproape de leșin",
        "Confuzie severă bruscă + respirație anormală",
        "Reacție severă a pielii (vezicule, exfoliere)",
    ])

    doc.add_paragraph()
    add_heading_custom(doc, "Când sunați medicul curant (nu urgent, dar curând)", level=2)
    add_bullet(doc, [
        "Reacție adversă persistentă care nu e urgentă (tuse, edem gleznă, greață lungă)",
        "Simptome noi care apar după începutul unui medicament",
        "Valori neobișnuite la auto-măsurare (tensiune, glicemie)",
        "Înainte de orice intervenție medicală (extracție dentară, vaccin, chirurgie)",
    ])
    doc.add_page_break()

    # === PARTEA V: LIMITELE DATELOR ===
    add_heading_custom(doc, "PARTEA V — Ce NU am găsit (transparentă)",
                       level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    p = doc.add_paragraph()
    add_marker(p, "NEGASIT")
    p.add_run("Rata exactă a angioedemului la combinația sitagliptin + perindopril. "
              "RCP menționează riscul, dar nu dă un procent. Date provin din raportări postmarketing "
              "izolate.")

    p = doc.add_paragraph()
    add_marker(p, "NEGASIT")
    p.add_run("Un RCP Triplixam 10/2.5/5 mg în limba română, emis de ANMDMR, cu data versiunii vizibile. "
              "Am folosit RCP-ul internațional al Servier (06.2021) și sursa Rwanda FDA (2023) — "
              "conținutul e echivalent.")

    p = doc.add_paragraph()
    add_marker(p, "INCERT")
    p.add_run("Pentru sitagliptin și insuficiență cardiacă: datele între trialurile clinice pe "
              "DPP-4 inhibitori sunt mixte. TECOS (sitagliptin) neutru, SAVOR-TIMI (saxagliptin) "
              "cu semnal de risc. Pentru pacientul actual cu boală cardiacă, e întrebare pentru "
              "cardiolog/diabetolog.")

    p = doc.add_paragraph()
    add_marker(p, "PROBABIL")
    p.add_run(
        "Interacțiunile specifice dintre substanțele active ale celor două medicamente "
        "la acest pacient individual depind de factori pe care un RCP nu-i poate prezice: "
        "vârsta, funcția renală actuală, funcția hepatică, simptome specifice. Evaluarea "
        "finală aparține medicului curant, care cunoaște întregul tablou."
    )

    doc.add_paragraph()
    p = doc.add_paragraph(
        "Aceste aspecte sunt întrebări legitime de pus medicului curant sau farmacistului "
        "la următoarea întâlnire."
    )
    doc.add_page_break()

    # === PARTEA VI: SURSE ===
    add_heading_custom(doc, "PARTEA VI — Surse citate", level=1, color=RGBColor(0x1F, 0x4E, 0x79))

    add_heading_custom(doc, "Pentru Jamesi (sitagliptin/metformin)", level=2)
    add_bullet(doc, [
        "Summary of Product Characteristics (SmPC) Janumet 50mg/1000mg — Electronic Medicines Compendium UK (emc). Sursă: https://www.medicines.org.uk/emc/product/564/smpc — consultată 18.04.2026.",
        "FDA prescribing information Janumet — 2017 label. Sursă: https://www.accessdata.fda.gov/drugsatfda_docs/label/2017/022044s042lbl.pdf",
        "DailyMed JANUMET. Sursă: https://dailymed.nlm.nih.gov/dailymed/drugInfo.cfm?setid=d19c7ed0-ad5c-426e-b2df-722508f97d67",
    ])

    add_heading_custom(doc, "Pentru Triplixam (perindopril/indapamidă/amlodipină)", level=2)
    add_bullet(doc, [
        "SmPC Triplixam (5/1.25/5, 10/2.5/5, 10/2.5/10, 5/1.25/10) — Servier, versiune 06.2021. Sursă: https://myservier-me.com/wp-content/uploads/2022/09/SmPC_Triplixam_Version-06.2021-Bhn.pdf — consultată 18.04.2026.",
        "SmPC Triplixam (Rwanda FDA) — 2023. Sursă: https://rwandafda.gov.rw/wp-content/uploads/2023/07/",
        "PMC8964631: „Perindopril/Indapamide/Amlodipine in Hypertension: A Profile of Its Use”. Sursă: https://pmc.ncbi.nlm.nih.gov/articles/PMC8964631/",
        "Farmacia Tei — Triplixam 10/2.5/5 (prezență în piața RO, referință ANMDMR). Sursă: https://comenzi.farmaciatei.ro/medicamente-cu-reteta/medicamente/triplixam-10mg2-5mg5mg-30-comprimate-filmate-les-laboratoires-servier-p345501",
    ])

    add_heading_custom(doc, "Note de temporalitate", level=2)
    p = doc.add_paragraph()
    add_marker(p, "CERT")
    p.add_run(
        "RCP-urile sunt actualizate periodic de EMA / ANMDMR. Versiunea Servier din 06.2021 "
        "poate avea actualizări pe anii 2024-2025. "
        "Pentru decizii terapeutice importante, se verifică versiunea curentă pe site-ul ANMDMR "
        "(www.anm.ro) sau EMA (www.ema.europa.eu)."
    )
    doc.add_paragraph()

    # === FINAL ===
    add_heading_custom(doc, "Control versiune", level=1)
    add_bullet(doc, [
        "Autor: Claude Code (Opus 4.7, 1M context)",
        "Data generare: 2026-04-18",
        "Versiune raport: 1.0",
        "Sesiunea Claude: remediere audit sub-clauza 7 + cercetare reacții adverse",
        "Marcaje aplicate: [CERT], [PROBABIL], [INCERT], [NEGASIT]",
        "Conformitate: Regula 17 CLAUDE.md (marcaj certitudine info medicală)",
    ])

    add_callout(
        doc,
        "ATENȚIONARE FINALĂ",
        "Informațiile din acest document sunt traduceri și rezumate ale prospectelor oficiale (RCP/SmPC). "
        "NU reprezintă recomandare medicală personalizată. Decizia de a continua, opri sau "
        "modifica un medicament aparține EXCLUSIV medicului curant, care cunoaște toate aspectele "
        "individuale ale pacientului. În caz de dubiu, sunați medicul. În caz de urgență, sunați 112.",
        color=RED,
    )

    return doc


def main():
    doc = build_document()
    output_path = Path(__file__).parent / "2026-04-18_raport_reactii_adverse_jamesi_triplixam.docx"
    doc.save(str(output_path))
    print(f"Document salvat: {output_path}")
    print(f"Dimensiune: {output_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
